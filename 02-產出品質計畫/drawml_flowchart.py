#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DrawML 流程圖產生器（quality-plan-maker 整合版）
============================================

以 Office Open XML DrawML（wpc:wpc canvas）在 Word (.docx) 內繪製
可編輯、向量式流程圖。形狀可點擊、拖曳、縮放、改色，對應
整體品質計畫「全文黑色、圖形無填色」的金質獎格式規定。

與原版 drawml-flowchart.skill 差異
----------------------------------
1. 新增 `quality_plan` 主題：無填色（a:noFill）、黑色框線、黑色文字
2. 新增 `finalize_flowcharts(docx_path, pending)` 後處理器：
   - 掃描文件中的 `<<<DRAWML_FLOWCHART:id>>>` 標記段落
   - 以對應流程圖定義產生的 DrawML 段落取代之
   - 自動補上 wpc/wps 命名空間與 mc:Ignorable
3. 新增「品質計畫常用流程圖」預設函式 `preset_*`：
   - preset_material_approval()  材料/設備選定送審流程圖
   - preset_work_inspection()    施工檢驗流程圖
   - preset_nc_material()        材料自主檢查不合格管制流程圖
   - preset_nc_work()            施工自主檢查不合格管制流程圖
   - preset_corrective()         矯正措施流程圖
   - preset_incoming_doc()       收文傳送及歸檔流程圖
   - preset_outgoing_doc()       發文傳送及歸檔流程圖
   - preset_equipment_test()     設備功能運轉檢測程序流程圖
   - preset_audit_flow()         品質稽核流程圖
   - preset_three_col(...)       三欄式施工作業流程及檢驗圖（改用 python-docx 表格 + DrawML）

使用方式（於 quality_plan_template.py 內）
-----------------------------------------
    from drawml_flowchart import (
        add_flowchart_placeholder, finalize_flowcharts,
        preset_material_approval, preset_nc_material, ...
    )

    # 在每個流程圖的位置：
    add_flowchart_placeholder(doc, 'flow_material_approval',
                              preset_material_approval(title='圖伍-1 材料/設備選定送審流程圖'))

    # 全部章節完成後：
    doc.save(OUT_PATH)
    finalize_flowcharts(OUT_PATH)
    fix_docx(OUT_PATH)
"""

import os, io, zipfile, shutil, copy, re
from collections import defaultdict, deque
from lxml import etree

# ── Namespaces ────────────────────────────────────────────────────────────────
W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WP  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
A   = 'http://schemas.openxmlformats.org/drawingml/2006/main'
WPS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
WPC = 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas'
MC  = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
XML_NS = 'http://www.w3.org/XML/1998/namespace'

NSMAP = {
    'w': W, 'wp': WP, 'a': A, 'wps': WPS, 'wpc': WPC, 'mc': MC,
}

def q(ns_uri, tag):
    return f'{{{ns_uri}}}{tag}'


# ── Units & geometry ─────────────────────────────────────────────────────────
EMU_IN = 914400  # EMU per inch

# Shape dimensions (width, height) in EMU — slightly larger for CJK text
DIMS = {
    'process':    (int(2.1*EMU_IN), int(0.65*EMU_IN)),
    'decision':   (int(1.8*EMU_IN), int(1.05*EMU_IN)),
    'terminator': (int(1.9*EMU_IN), int(0.55*EMU_IN)),
    'data':       (int(2.1*EMU_IN), int(0.65*EMU_IN)),
    'document':   (int(2.1*EMU_IN), int(0.70*EMU_IN)),
    'connector':  (int(0.55*EMU_IN), int(0.55*EMU_IN)),
    'default':    (int(2.1*EMU_IN), int(0.65*EMU_IN)),
}

GEOM = {
    'process':    'flowChartProcess',
    'decision':   'flowChartDecision',
    'terminator': 'flowChartTerminator',
    'data':       'flowChartPunchedCard',
    'document':   'flowChartDocument',
    'connector':  'flowChartConnector',
    'default':    'flowChartProcess',
}

# ── Themes ────────────────────────────────────────────────────────────────────
# 值為 (fill, text, border)；fill='' 代表 a:noFill（透明／無填色）
THEMES = {
    # 【品質計畫專用主題】符合金質獎格式規定：全文黑色、圖形無填色
    'quality_plan': {
        'process':    ('', '000000', '000000'),
        'decision':   ('', '000000', '000000'),
        'terminator': ('', '000000', '000000'),
        'data':       ('', '000000', '000000'),
        'document':   ('', '000000', '000000'),
        'connector':  ('', '000000', '000000'),
        'default':    ('', '000000', '000000'),
        'arrow':      '000000',
        'label_bg':   'FFFFFF',   # 判斷分支標籤（是／否）：白底黑框黑字
        'label_fg':   '000000',
        'label_border': '000000',
    },
    'blue': {
        'process':    ('4472C4', 'FFFFFF', '2F5496'),
        'decision':   ('ED7D31', 'FFFFFF', 'C55A11'),
        'terminator': ('70AD47', 'FFFFFF', '375623'),
        'data':       ('4472C4', 'FFFFFF', '2F5496'),
        'document':   ('4472C4', 'FFFFFF', '2F5496'),
        'connector':  ('7030A0', 'FFFFFF', '560D89'),
        'default':    ('4472C4', 'FFFFFF', '2F5496'),
        'arrow':      '404040',
        'label_bg':   'F2F2F2',
        'label_fg':   '404040',
        'label_border': '404040',
    },
    'mono': {
        'process':    ('404040', 'FFFFFF', '202020'),
        'decision':   ('707070', 'FFFFFF', '505050'),
        'terminator': ('202020', 'FFFFFF', '000000'),
        'data':       ('505050', 'FFFFFF', '303030'),
        'document':   ('404040', 'FFFFFF', '202020'),
        'connector':  ('606060', 'FFFFFF', '404040'),
        'default':    ('404040', 'FFFFFF', '202020'),
        'arrow':      '505050',
        'label_bg':   'F0F0F0',
        'label_fg':   '303030',
        'label_border': '505050',
    },
}

V_GAP   = int(0.50 * EMU_IN)
H_GAP   = int(0.60 * EMU_IN)
MARGIN  = int(0.40 * EMU_IN)
LABEL_H = int(0.32 * EMU_IN)
LABEL_W = int(0.70 * EMU_IN)


# ── Layout (Sugiyama-style layered layout) ───────────────────────────────────

def _assign_layers(node_ids, edges):
    """
    以拓撲排序決定各節點 layer；遇到迴圈（例如「不合格 → 回前步驟」）時，
    反向邊視為退回邊，不參與 layer 計算，但仍會於 build_canvas() 繪製箭頭。
    """
    adj = defaultdict(list)
    indeg = defaultdict(int)
    for e in edges:
        adj[e['source']].append(e['target'])
    # 先做一次 DFS 偵測反向邊（back-edges），排除後再算 layer
    WHITE, GRAY, BLACK = 0, 1, 2
    color = {n: WHITE for n in node_ids}
    back_edges = set()

    def dfs(u):
        color[u] = GRAY
        for v in adj[u]:
            if color[v] == GRAY:
                back_edges.add((u, v))
            elif color[v] == WHITE:
                dfs(v)
        color[u] = BLACK

    import sys as _sys
    _sys.setrecursionlimit(max(1000, len(node_ids) * 10))
    for n in node_ids:
        if color[n] == WHITE:
            dfs(n)

    # 構造無環圖的 indeg 與 adj
    dag_adj = defaultdict(list)
    dag_indeg = defaultdict(int)
    for e in edges:
        if (e['source'], e['target']) in back_edges:
            continue
        dag_adj[e['source']].append(e['target'])
        dag_indeg[e['target']] += 1
    for n in node_ids:
        dag_indeg.setdefault(n, 0)

    # Kahn 拓撲排序 + 層次計算
    layer = {n: 0 for n in node_ids}
    queue = deque([n for n in node_ids if dag_indeg[n] == 0])
    if not queue:
        queue.append(node_ids[0])
    processed = set()
    while queue:
        n = queue.popleft()
        if n in processed:
            continue
        processed.add(n)
        for m in dag_adj[n]:
            if layer[m] < layer[n] + 1:
                layer[m] = layer[n] + 1
            dag_indeg[m] -= 1
            if dag_indeg[m] <= 0 and m not in processed:
                queue.append(m)
    return layer


def compute_layout(nodes, edges, direction='TB'):
    """Return {id: (x, y, w, h)} in EMU."""
    ids = [n['id'] for n in nodes]
    types = {n['id']: n.get('type', 'process') for n in nodes}
    layer_map = _assign_layers(ids, edges)

    layers = defaultdict(list)
    for nid in ids:
        layers[layer_map[nid]].append(nid)

    pos = {}
    if direction == 'TB':
        y = MARGIN
        # Find widest layer to compute canvas width
        max_layer_w = 0
        for li in sorted(layers):
            widths = [DIMS.get(types[n], DIMS['default'])[0] for n in layers[li]]
            total = sum(widths) + H_GAP * (len(layers[li]) - 1)
            max_layer_w = max(max_layer_w, total)
        canvas_cx = MARGIN + max_layer_w // 2

        for li in sorted(layers):
            layer_nodes = layers[li]
            widths  = [DIMS.get(types[n], DIMS['default'])[0] for n in layer_nodes]
            heights = [DIMS.get(types[n], DIMS['default'])[1] for n in layer_nodes]
            total_w = sum(widths) + H_GAP * (len(layer_nodes) - 1)
            x = canvas_cx - total_w // 2
            for i, nid in enumerate(layer_nodes):
                w, h = DIMS.get(types[nid], DIMS['default'])
                pos[nid] = (x, y, w, h)
                x += w + H_GAP
            y += max(heights) + V_GAP
    else:
        x = MARGIN
        for li in sorted(layers):
            layer_nodes = layers[li]
            widths  = [DIMS.get(types[n], DIMS['default'])[0] for n in layer_nodes]
            heights = [DIMS.get(types[n], DIMS['default'])[1] for n in layer_nodes]
            y = MARGIN
            for i, nid in enumerate(layer_nodes):
                w, h = DIMS.get(types[nid], DIMS['default'])
                pos[nid] = (x, y, w, h)
                y += h + V_GAP
            x += max(widths) + H_GAP
    return pos


# ── XML Builders ──────────────────────────────────────────────────────────────

def _sub(parent, ns_uri, tag, **attribs):
    el = etree.SubElement(parent, q(ns_uri, tag))
    for k, v in attribs.items():
        el.set(k, str(v))
    return el


def _sp_lock(parent):
    cNvSpPr = _sub(parent, WPS, 'cNvSpPr')
    locks = _sub(cNvSpPr, A, 'spLocks')
    locks.set('noChangeArrowheads', '1')


def _xfrm(parent, x, y, cx, cy, flipH=False, flipV=False):
    xfrm = _sub(parent, A, 'xfrm')
    if flipH: xfrm.set('flipH', '1')
    if flipV: xfrm.set('flipV', '1')
    off = _sub(xfrm, A, 'off'); off.set('x', str(x)); off.set('y', str(y))
    ext = _sub(xfrm, A, 'ext'); ext.set('cx', str(cx)); ext.set('cy', str(cy))


def _fill(parent, hex_color):
    """若 hex_color 為空或 None，使用 a:noFill；否則 a:solidFill。"""
    if not hex_color:
        return _sub(parent, A, 'noFill')
    sf = _sub(parent, A, 'solidFill')
    clr = _sub(sf, A, 'srgbClr'); clr.set('val', hex_color)
    return sf


def _border_line(parent, hex_color, width_pt=1.0):
    ln = _sub(parent, A, 'ln')
    ln.set('w', str(int(width_pt * 12700)))
    _fill(ln, hex_color)


def build_shape(node, pos, theme, font_cjk=None, font_latin=None, font_pt=10):
    """建立單一節點的 wps:wsp 形狀元素。"""
    x, y, w, h = pos
    ntype = node.get('type', 'default')
    label = node.get('label', node['id'])
    fill, text_color, border = theme.get(ntype, theme['default'])

    wsp = etree.Element(q(WPS, 'wsp'))
    _sp_lock(wsp)

    spPr = _sub(wsp, WPS, 'spPr')
    _xfrm(spPr, x, y, w, h)
    geom = _sub(spPr, A, 'prstGeom')
    geom.set('prst', GEOM.get(ntype, GEOM['default']))
    _sub(geom, A, 'avLst')
    _fill(spPr, fill)
    _border_line(spPr, border, width_pt=1.25)

    # 文字
    txbx = _sub(wsp, WPS, 'txbx')
    content = _sub(txbx, W, 'txbxContent')
    # 逐行處理（允許 \n 換行）
    lines = label.split('\n') if label else ['']
    for li, line in enumerate(lines):
        p = _sub(content, W, 'p')
        pPr = _sub(p, W, 'pPr')
        _sub(pPr, W, 'jc').set(q(W, 'val'), 'center')
        sp = _sub(pPr, W, 'spacing')
        sp.set(q(W, 'before'), '0'); sp.set(q(W, 'after'), '0'); sp.set(q(W, 'line'), '240'); sp.set(q(W, 'lineRule'), 'auto')

        r = _sub(p, W, 'r')
        rPr = _sub(r, W, 'rPr')
        # 設定字型
        if font_cjk or font_latin:
            rFonts = _sub(rPr, W, 'rFonts')
            if font_latin:
                rFonts.set(q(W, 'ascii'), font_latin)
                rFonts.set(q(W, 'hAnsi'), font_latin)
            if font_cjk:
                rFonts.set(q(W, 'eastAsia'), font_cjk)
        _sub(rPr, W, 'color').set(q(W, 'val'), text_color)
        sz_halfpt = str(int(font_pt * 2))
        _sub(rPr, W, 'sz').set(q(W, 'val'), sz_halfpt)
        _sub(rPr, W, 'szCs').set(q(W, 'val'), sz_halfpt)

        t = _sub(r, W, 't')
        t.text = line
        if line and (line[0] == ' ' or line[-1] == ' '):
            t.set(q(XML_NS, 'space'), 'preserve')

    bodyPr = _sub(wsp, WPS, 'bodyPr')
    bodyPr.set('anchor', 'ctr')
    bodyPr.set('lIns', '45720'); bodyPr.set('rIns', '45720')
    bodyPr.set('tIns', '36576'); bodyPr.set('bIns', '36576')
    bodyPr.set('wrap', 'square')
    _sub(bodyPr, A, 'noAutofit')
    return wsp


def build_connector(x1, y1, x2, y2, theme, edge_label=None, font_cjk=None, font_latin=None, font_pt=9):
    """回傳 [connector wsp] 或 [connector wsp, label wsp]。"""
    arrow_color   = theme.get('arrow', '404040')
    label_bg      = theme.get('label_bg', 'FFFFFF')
    label_fg      = theme.get('label_fg', '000000')
    label_border  = theme.get('label_border', label_fg)

    result = []
    lx = min(x1, x2); ly = min(y1, y2)
    cx = max(abs(x2 - x1), 914)
    cy = max(abs(y2 - y1), 914)
    flipH = x2 < x1
    flipV = y2 < y1

    wsp = etree.Element(q(WPS, 'wsp'))
    _sub(wsp, WPS, 'cNvCnPr')
    spPr = _sub(wsp, WPS, 'spPr')
    _xfrm(spPr, lx, ly, cx, cy, flipH=flipH, flipV=flipV)

    geom = _sub(spPr, A, 'prstGeom')
    prst = 'straightConnector1' if (x1 == x2 or y1 == y2) else 'bentConnector3'
    geom.set('prst', prst)
    _sub(geom, A, 'avLst')

    ln = _sub(spPr, A, 'ln')
    ln.set('w', '15875')  # ~1.25pt
    _fill(ln, arrow_color)
    _sub(ln, A, 'headEnd').set('type', 'none')
    tail = _sub(ln, A, 'tailEnd')
    tail.set('type', 'arrow'); tail.set('w', 'med'); tail.set('len', 'med')
    _sub(wsp, WPS, 'bodyPr')
    result.append(wsp)

    if edge_label:
        lw = LABEL_W; lh = LABEL_H
        label_x = (x1 + x2) // 2 - lw // 2
        label_y = (y1 + y2) // 2 - lh // 2

        lwsp = etree.Element(q(WPS, 'wsp'))
        _sp_lock(lwsp)
        lspPr = _sub(lwsp, WPS, 'spPr')
        _xfrm(lspPr, label_x, label_y, lw, lh)
        lgeom = _sub(lspPr, A, 'prstGeom')
        lgeom.set('prst', 'rect'); _sub(lgeom, A, 'avLst')
        # 純色背景（白或灰）＋黑框
        _fill(lspPr, label_bg)
        lln = _sub(lspPr, A, 'ln'); lln.set('w', '6350')
        _fill(lln, label_border)

        ltxbx = _sub(lwsp, WPS, 'txbx')
        lcontent = _sub(ltxbx, W, 'txbxContent')
        lp = _sub(lcontent, W, 'p')
        lpPr = _sub(lp, W, 'pPr')
        _sub(lpPr, W, 'jc').set(q(W, 'val'), 'center')
        lsp = _sub(lpPr, W, 'spacing')
        lsp.set(q(W, 'before'), '0'); lsp.set(q(W, 'after'), '0')
        lr = _sub(lp, W, 'r')
        lrPr = _sub(lr, W, 'rPr')
        if font_cjk or font_latin:
            rFonts = _sub(lrPr, W, 'rFonts')
            if font_latin:
                rFonts.set(q(W, 'ascii'), font_latin); rFonts.set(q(W, 'hAnsi'), font_latin)
            if font_cjk:
                rFonts.set(q(W, 'eastAsia'), font_cjk)
        _sub(lrPr, W, 'color').set(q(W, 'val'), label_fg)
        sz = str(int(font_pt * 2))
        _sub(lrPr, W, 'sz').set(q(W, 'val'), sz)
        _sub(lrPr, W, 'szCs').set(q(W, 'val'), sz)
        lt = _sub(lr, W, 't'); lt.text = edge_label
        lbodyPr = _sub(lwsp, WPS, 'bodyPr')
        lbodyPr.set('anchor', 'ctr')
        lbodyPr.set('lIns', '18288'); lbodyPr.set('rIns', '18288')
        lbodyPr.set('tIns', '9144'); lbodyPr.set('bIns', '9144')
        _sub(lbodyPr, A, 'noAutofit')
        result.append(lwsp)

    return result


def build_canvas(nodes, edges, positions, theme_data, direction='TB',
                 font_cjk=None, font_latin=None, font_pt=10):
    """建立完整的 wpc:wpc canvas 元素，回傳 (canvas, width, height)。"""
    canvas = etree.Element(q(WPC, 'wpc'))
    _sub(canvas, WPC, 'bg')
    _sub(canvas, WPC, 'whole')

    for node in nodes:
        nid = node['id']
        if nid not in positions:
            continue
        sp = build_shape(node, positions[nid], theme_data,
                         font_cjk=font_cjk, font_latin=font_latin, font_pt=font_pt)
        canvas.append(sp)

    for edge in edges:
        src = edge.get('source'); dst = edge.get('target')
        if src not in positions or dst not in positions:
            continue
        sx, sy, sw, sh = positions[src]
        dx, dy, dw, dh = positions[dst]
        if direction == 'TB':
            x1, y1 = sx + sw // 2, sy + sh
            x2, y2 = dx + dw // 2, dy
        else:
            x1, y1 = sx + sw, sy + sh // 2
            x2, y2 = dx, dy + dh // 2
        elabel = edge.get('label', '')
        conn_els = build_connector(
            x1, y1, x2, y2, theme_data,
            edge_label=elabel if elabel else None,
            font_cjk=font_cjk, font_latin=font_latin, font_pt=max(font_pt-1, 8))
        for el in conn_els:
            canvas.append(el)

    max_x = max_y = 0
    for x, y, w, h in positions.values():
        max_x = max(max_x, x + w)
        max_y = max(max_y, y + h)
    canvas_w = max_x + MARGIN
    canvas_h = max_y + MARGIN
    return canvas, canvas_w, canvas_h


def wrap_canvas_in_paragraph(canvas, canvas_w, canvas_h, doc_id=1):
    """將 wpc:wpc 包成 w:p > w:r > w:drawing > wp:inline 結構。"""
    p = etree.Element(q(W, 'p'))
    # 置中
    pPr = _sub(p, W, 'pPr')
    _sub(pPr, W, 'jc').set(q(W, 'val'), 'center')
    r = _sub(p, W, 'r')
    ac = _sub(r, MC, 'AlternateContent')
    choice = _sub(ac, MC, 'Choice'); choice.set('Requires', 'wpc')
    drawing = _sub(choice, W, 'drawing')
    inline = _sub(drawing, WP, 'inline')
    inline.set('distT', '0'); inline.set('distB', '0')
    inline.set('distL', '0'); inline.set('distR', '0')
    extent = _sub(inline, WP, 'extent'); extent.set('cx', str(canvas_w)); extent.set('cy', str(canvas_h))
    ee = _sub(inline, WP, 'effectExtent')
    ee.set('l', '0'); ee.set('t', '0'); ee.set('r', '0'); ee.set('b', '0')
    docPr = _sub(inline, WP, 'docPr')
    docPr.set('id', str(doc_id)); docPr.set('name', f'Flowchart{doc_id}')
    _sub(inline, WP, 'cNvGraphicFramePr')
    graphic = _sub(inline, A, 'graphic')
    gdata = _sub(graphic, A, 'graphicData'); gdata.set('uri', WPC)
    gdata.append(canvas)

    fallback = _sub(ac, MC, 'Fallback')
    pict = _sub(fallback, W, 'pict')
    VML = 'urn:schemas-microsoft-com:vml'
    rect = etree.SubElement(pict, f'{{{VML}}}rect')
    rect.set('style', f'width:{canvas_w//914}pt;height:{canvas_h//914}pt')
    return p


def build_flowchart_paragraph(flow_def, theme_name='quality_plan', direction='TB',
                              font_cjk=None, font_latin=None, font_pt=10, doc_id=1):
    """給定 flowchart 定義（nodes / edges / title 可選），回傳 w:p 元素。"""
    nodes = flow_def.get('nodes', [])
    edges = flow_def.get('edges', [])
    theme = THEMES.get(theme_name, THEMES['quality_plan'])
    positions = compute_layout(nodes, edges, direction=direction)
    canvas, cw, ch = build_canvas(nodes, edges, positions, theme, direction=direction,
                                  font_cjk=font_cjk, font_latin=font_latin, font_pt=font_pt)
    return wrap_canvas_in_paragraph(canvas, cw, ch, doc_id=doc_id)


# ─────────────────────────────────────────────────────────────────────────────
# 後處理：於 saved docx 中把 <<<DRAWML_FLOWCHART:id>>> 段落換成實際 DrawML
# ─────────────────────────────────────────────────────────────────────────────

MARKER_PATTERN = re.compile(r'<<<DRAWML_FLOWCHART:([A-Za-z0-9_\-]+)>>>')

NEEDED_NS = {
    'mc':  MC,
    'wp':  WP,
    'a':   A,
    'wps': WPS,
    'wpc': WPC,
}
MC_IGNORABLE_ADD = ['wpc', 'wps']


def _ensure_namespaces(root):
    """確保 <w:document> 根元素宣告 wpc/wps 命名空間，並將其加入 mc:Ignorable。"""
    current_nsmap = dict(root.nsmap)
    # 若命名空間已齊備，直接回傳原 root
    missing = [uri for uri in NEEDED_NS.values() if uri not in current_nsmap.values()]

    if missing:
        # lxml 無法就地修改 nsmap → 重建 root
        new_nsmap = dict(current_nsmap)
        for prefix, uri in NEEDED_NS.items():
            if uri not in new_nsmap.values():
                # 避免前綴衝突
                p = prefix
                while p in new_nsmap:
                    p = p + '_'
                new_nsmap[p] = uri
        new_root = etree.Element(root.tag, nsmap=new_nsmap)
        for k, v in root.attrib.items():
            new_root.set(k, v)
        for child in list(root):
            new_root.append(child)
        root = new_root

    ig_attr = q(MC, 'Ignorable')
    current = root.get(ig_attr, '')
    parts = set(current.split())
    for tok in MC_IGNORABLE_ADD:
        parts.add(tok)
    root.set(ig_attr, ' '.join(sorted(parts)))
    return root


def _find_marker_paragraphs(root):
    """找出所有含 marker 文字的 <w:p> 節點，回傳 [(marker_id, p_element), ...]。"""
    results = []
    body = root.find(q(W, 'body'))
    if body is None:
        return results
    for p in body.iter(q(W, 'p')):
        # 收集段落中所有 w:t 的文字
        texts = [(t.text or '') for t in p.iter(q(W, 't'))]
        full = ''.join(texts)
        m = MARKER_PATTERN.search(full)
        if m:
            results.append((m.group(1), p))
    return results


def finalize_flowcharts(docx_path, pending=None,
                        theme_name='quality_plan',
                        direction='TB',
                        font_cjk='標楷體',
                        font_latin='Arial',
                        font_pt=10,
                        verbose=True):
    """
    後處理：於 saved docx 中把 <<<DRAWML_FLOWCHART:id>>> 段落替換為實際的 DrawML 流程圖。

    pending: dict[id] → flow_def
             若為 None，則由模組全域 _PENDING 取得（見 register_flowchart()）。
    """
    if pending is None:
        pending = dict(_PENDING)
    if not pending:
        if verbose:
            print('⚠ finalize_flowcharts: 無待處理流程圖，跳過。')
        return

    import tempfile as _tf
    _tmpfd, tmp = _tf.mkstemp(suffix='.tmp', dir='/tmp')
    import os as _os; _os.close(_tmpfd)
    shutil.copy2(docx_path, tmp)
    with zipfile.ZipFile(tmp, 'r') as zin:
        doc_xml = zin.read('word/document.xml')

    root = etree.fromstring(doc_xml)
    root = _ensure_namespaces(root)

    markers = _find_marker_paragraphs(root)
    replaced = 0
    for i, (mid, p_el) in enumerate(markers):
        if mid not in pending:
            if verbose:
                print(f'⚠ finalize: marker "{mid}" 無對應定義，略過。')
            continue
        flow_def = pending[mid]
        # 品管組織架構圖：直接呼叫 build_quality_org_chart()
        if flow_def.get('type') == 'quality_org_chart':
            flow_para = build_quality_org_chart(
                label_tech   = flow_def.get('label_tech',    '主任技師'),
                label_foreman= flow_def.get('label_foreman', '工地主任'),
                label_qc     = flow_def.get('label_qc',      '品管人員'),
                label_safety = flow_def.get('label_safety',  '職業安全衛生\n管理人員'),
                label_sub    = flow_def.get('label_sub',     '各協力廠商'),
                label_eng    = flow_def.get('label_eng',     '現場工程師'),
                label_admin  = flow_def.get('label_admin',   '行政人員'),
                font_cjk     = flow_def.get('font_cjk',     font_cjk),
                font_latin   = flow_def.get('font_latin',   font_latin),
                font_pt      = flow_def.get('font_pt',      font_pt),
                doc_id       = i + 1,
            )
        else:
            flow_para = build_flowchart_paragraph(
                flow_def,
                theme_name=theme_name,
                direction=flow_def.get('direction', direction),
                font_cjk=font_cjk, font_latin=font_latin, font_pt=font_pt,
                doc_id=i + 1,
            )
        parent = p_el.getparent()
        idx = list(parent).index(p_el)
        parent.remove(p_el)
        parent.insert(idx, flow_para)
        replaced += 1

    new_xml = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
    with zipfile.ZipFile(tmp, 'r') as zin, zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == 'word/document.xml':
                zout.writestr(item, new_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
    os.remove(tmp)

    # 清空全域 pending 清單
    _PENDING.clear()

    if verbose:
        print(f'✓ finalize_flowcharts: 替換 {replaced}/{len(markers)} 個流程圖占位符')


# ─────────────────────────────────────────────────────────────────────────────
# 全域 pending 管理 + 占位符插入 helper（供 quality_plan_template.py 使用）
# ─────────────────────────────────────────────────────────────────────────────

_PENDING = {}


def register_flowchart(marker_id, flow_def):
    """把流程圖定義登錄至全域 _PENDING。"""
    _PENDING[marker_id] = flow_def


def add_flowchart_placeholder(doc, marker_id, flow_def):
    """
    在 python-docx 文件目前位置插入一個 DrawML 占位符段落，
    並登錄流程圖定義供後續 finalize_flowcharts() 替換。

    參數:
      doc:        python-docx Document
      marker_id:  唯一識別字串（建議使用 figure 編號，如 'flow_fig_5_1'）
      flow_def:   {'nodes': [...], 'edges': [...], 'direction': 'TB'|'LR'（選）}
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    register_flowchart(marker_id, flow_def)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'<<<DRAWML_FLOWCHART:{marker_id}>>>')
    # 以較小字體避免占太多版面
    from docx.shared import Pt
    run.font.size = Pt(1)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# 品質計畫常用流程圖預設
# ─────────────────────────────────────────────────────────────────────────────

def preset_material_approval():
    """材料/設備選定送審流程圖。"""
    return {
        'nodes': [
            {'id': 'start',     'label': '工程開始',           'type': 'terminator'},
            {'id': 'plan',      'label': '擬定材料/設備清單',  'type': 'process'},
            {'id': 'submit',    'label': '廠商送審（型錄/樣品/試驗報告）', 'type': 'data'},
            {'id': 'review',    'label': '監造單位審查',       'type': 'process'},
            {'id': 'decide',    'label': '審查是否核可？',     'type': 'decision'},
            {'id': 'revise',    'label': '退回廠商補正',       'type': 'process'},
            {'id': 'approve',   'label': '核可進場',           'type': 'process'},
            {'id': 'record',    'label': '登錄核可清冊歸檔',   'type': 'document'},
            {'id': 'end',       'label': '可供施工使用',       'type': 'terminator'},
        ],
        'edges': [
            {'source': 'start',   'target': 'plan'},
            {'source': 'plan',    'target': 'submit'},
            {'source': 'submit',  'target': 'review'},
            {'source': 'review',  'target': 'decide'},
            {'source': 'decide',  'target': 'revise',  'label': '否'},
            {'source': 'revise',  'target': 'submit'},
            {'source': 'decide',  'target': 'approve', 'label': '是'},
            {'source': 'approve', 'target': 'record'},
            {'source': 'record',  'target': 'end'},
        ],
    }


def preset_work_inspection():
    """施工檢驗流程圖。"""
    return {
        'nodes': [
            {'id': 'start',    'label': '施工作業開始',         'type': 'terminator'},
            {'id': 'self_ck',  'label': '承攬廠商自主檢查',     'type': 'process'},
            {'id': 'self_ok',  'label': '自主檢查合格？',       'type': 'decision'},
            {'id': 'fix1',     'label': '不合格改善',           'type': 'process'},
            {'id': 'notify',   'label': '通知監造查驗',         'type': 'data'},
            {'id': 'audit',    'label': '監造抽驗/停留點查驗',  'type': 'process'},
            {'id': 'audit_ok', 'label': '查驗合格？',           'type': 'decision'},
            {'id': 'fix2',     'label': '依意見改善並複驗',     'type': 'process'},
            {'id': 'record',   'label': '填具檢驗紀錄表',       'type': 'document'},
            {'id': 'next',     'label': '進入下一工項',         'type': 'terminator'},
        ],
        'edges': [
            {'source': 'start',    'target': 'self_ck'},
            {'source': 'self_ck',  'target': 'self_ok'},
            {'source': 'self_ok',  'target': 'fix1',   'label': '否'},
            {'source': 'fix1',     'target': 'self_ck'},
            {'source': 'self_ok',  'target': 'notify', 'label': '是'},
            {'source': 'notify',   'target': 'audit'},
            {'source': 'audit',    'target': 'audit_ok'},
            {'source': 'audit_ok', 'target': 'fix2',   'label': '否'},
            {'source': 'fix2',     'target': 'audit'},
            {'source': 'audit_ok', 'target': 'record', 'label': '是'},
            {'source': 'record',   'target': 'next'},
        ],
    }


def preset_nc_material():
    """材料自主檢查不合格管制流程圖。"""
    return {
        'nodes': [
            {'id': 'arrive',  'label': '材料進場',          'type': 'terminator'},
            {'id': 'check',   'label': '自主檢查',          'type': 'process'},
            {'id': 'judge',   'label': '是否合格？',        'type': 'decision'},
            {'id': 'nc',      'label': '開立不合格通知',    'type': 'document'},
            {'id': 'segr',    'label': '隔離標示並退場',    'type': 'process'},
            {'id': 'track',   'label': '填具不合格管制表追蹤', 'type': 'document'},
            {'id': 'accept',  'label': '准予使用於工程',    'type': 'process'},
            {'id': 'end',     'label': '歸檔結案',          'type': 'terminator'},
        ],
        'edges': [
            {'source': 'arrive', 'target': 'check'},
            {'source': 'check',  'target': 'judge'},
            {'source': 'judge',  'target': 'nc',     'label': '否'},
            {'source': 'nc',     'target': 'segr'},
            {'source': 'segr',   'target': 'track'},
            {'source': 'judge',  'target': 'accept', 'label': '是'},
            {'source': 'accept', 'target': 'end'},
            {'source': 'track',  'target': 'end'},
        ],
    }


def preset_nc_work():
    """施工自主檢查不合格管制流程圖。"""
    return {
        'nodes': [
            {'id': 'step',    'label': '施工作業',          'type': 'terminator'},
            {'id': 'check',   'label': '自主檢查',          'type': 'process'},
            {'id': 'judge',   'label': '是否合格？',        'type': 'decision'},
            {'id': 'nc',      'label': '開立不合格通知',    'type': 'document'},
            {'id': 'stop',    'label': '停止後續作業',      'type': 'process'},
            {'id': 'rework',  'label': '改善/重作',         'type': 'process'},
            {'id': 'recheck', 'label': '複驗',              'type': 'process'},
            {'id': 'pass',    'label': '進入下一工項',      'type': 'process'},
            {'id': 'record',  'label': '填具管制表歸檔',    'type': 'document'},
            {'id': 'end',     'label': '結案',              'type': 'terminator'},
        ],
        'edges': [
            {'source': 'step',    'target': 'check'},
            {'source': 'check',   'target': 'judge'},
            {'source': 'judge',   'target': 'nc',      'label': '否'},
            {'source': 'nc',      'target': 'stop'},
            {'source': 'stop',    'target': 'rework'},
            {'source': 'rework',  'target': 'recheck'},
            {'source': 'recheck', 'target': 'judge'},
            {'source': 'judge',   'target': 'pass',    'label': '是'},
            {'source': 'pass',    'target': 'record'},
            {'source': 'record',  'target': 'end'},
        ],
    }


def preset_corrective():
    """矯正措施流程圖。"""
    return {
        'nodes': [
            {'id': 'found',   'label': '發現不符合事項',    'type': 'terminator'},
            {'id': 'analyze', 'label': '原因分析',          'type': 'process'},
            {'id': 'plan',    'label': '擬定矯正措施計畫',  'type': 'process'},
            {'id': 'impl',    'label': '實施矯正',          'type': 'process'},
            {'id': 'verify',  'label': '追蹤查證是否有效？', 'type': 'decision'},
            {'id': 'redo',    'label': '重新檢討擬案',      'type': 'process'},
            {'id': 'close',   'label': '結案歸檔',          'type': 'document'},
            {'id': 'end',     'label': '預防重複發生',      'type': 'terminator'},
        ],
        'edges': [
            {'source': 'found',   'target': 'analyze'},
            {'source': 'analyze', 'target': 'plan'},
            {'source': 'plan',    'target': 'impl'},
            {'source': 'impl',    'target': 'verify'},
            {'source': 'verify',  'target': 'redo',    'label': '否'},
            {'source': 'redo',    'target': 'plan'},
            {'source': 'verify',  'target': 'close',   'label': '是'},
            {'source': 'close',   'target': 'end'},
        ],
    }


def preset_incoming_doc():
    """收文傳送及歸檔流程圖。"""
    return {
        'nodes': [
            {'id': 'recv',   'label': '收文',               'type': 'terminator'},
            {'id': 'reg',    'label': '登錄收文管制表',      'type': 'document'},
            {'id': 'dist',   'label': '分文（主辦/會辦）',   'type': 'process'},
            {'id': 'proc',   'label': '承辦處理',            'type': 'process'},
            {'id': 'cc',     'label': '副本抄送相關人員',    'type': 'process'},
            {'id': 'arch',   'label': '歸檔',                'type': 'terminator'},
        ],
        'edges': [
            {'source': 'recv', 'target': 'reg'},
            {'source': 'reg',  'target': 'dist'},
            {'source': 'dist', 'target': 'proc'},
            {'source': 'proc', 'target': 'cc'},
            {'source': 'cc',   'target': 'arch'},
        ],
    }


def preset_outgoing_doc():
    """發文傳送及歸檔流程圖。"""
    return {
        'nodes': [
            {'id': 'draft',  'label': '擬稿',                     'type': 'terminator'},
            {'id': 'rev',    'label': '核稿（工地主任→主任技師）', 'type': 'process'},
            {'id': 'send',   'label': '發文',                     'type': 'process'},
            {'id': 'deliv',  'label': '送達對方',                 'type': 'process'},
            {'id': 'cc',     'label': '副本存查',                 'type': 'process'},
            {'id': 'arch',   'label': '歸檔',                     'type': 'terminator'},
        ],
        'edges': [
            {'source': 'draft', 'target': 'rev'},
            {'source': 'rev',   'target': 'send'},
            {'source': 'send',  'target': 'deliv'},
            {'source': 'deliv', 'target': 'cc'},
            {'source': 'cc',    'target': 'arch'},
        ],
    }


def preset_equipment_test():
    """設備功能運轉檢測程序流程圖（9 章制景觀專用）。"""
    return {
        'nodes': [
            {'id': 'ready',  'label': '設備安裝完成',        'type': 'terminator'},
            {'id': 'single', 'label': '單機試運轉',          'type': 'process'},
            {'id': 'total',  'label': '整體功能運轉',        'type': 'process'},
            {'id': 'test',   'label': '功能檢測是否達標？',  'type': 'decision'},
            {'id': 'adjust', 'label': '調整/修正後複測',      'type': 'process'},
            {'id': 'record', 'label': '填具檢測紀錄表',      'type': 'document'},
            {'id': 'accept', 'label': '驗收交付',            'type': 'terminator'},
        ],
        'edges': [
            {'source': 'ready',  'target': 'single'},
            {'source': 'single', 'target': 'total'},
            {'source': 'total',  'target': 'test'},
            {'source': 'test',   'target': 'adjust', 'label': '否'},
            {'source': 'adjust', 'target': 'total'},
            {'source': 'test',   'target': 'record', 'label': '是'},
            {'source': 'record', 'target': 'accept'},
        ],
    }


def preset_audit_flow():
    """品質稽核流程圖。"""
    return {
        'nodes': [
            {'id': 'plan',    'label': '年度稽核計畫',           'type': 'terminator'},
            {'id': 'notify',  'label': '通知受稽核單位',         'type': 'data'},
            {'id': 'execute', 'label': '執行稽核',               'type': 'process'},
            {'id': 'finding', 'label': '填寫稽核發現單',         'type': 'document'},
            {'id': 'car',     'label': '發現缺失？',             'type': 'decision'},
            {'id': 'corr',    'label': '要求矯正並追蹤',         'type': 'process'},
            {'id': 'close',   'label': '結案歸檔',               'type': 'document'},
            {'id': 'end',     'label': '提報管理階層',           'type': 'terminator'},
        ],
        'edges': [
            {'source': 'plan',    'target': 'notify'},
            {'source': 'notify',  'target': 'execute'},
            {'source': 'execute', 'target': 'finding'},
            {'source': 'finding', 'target': 'car'},
            {'source': 'car',     'target': 'corr',  'label': '是'},
            {'source': 'corr',    'target': 'close'},
            {'source': 'car',     'target': 'close', 'label': '否'},
            {'source': 'close',   'target': 'end'},
        ],
    }


# ─────────────────────────────────────────────────────────────────────────────
# 三欄式施工作業流程及檢驗圖（DrawML 混合 python-docx 表格實作）
# ─────────────────────────────────────────────────────────────────────────────

def add_three_col_flow(doc, title, items,
                       marker_prefix='threecol',
                       font_cjk='標楷體', font_latin='Arial', font_pt=10):
    """
    三欄式施工作業流程及檢驗圖（欄1 作業流程圖｜欄2 檢驗要點｜欄3 相關記錄/文件）。

    items: list of dicts，每項對應一列
      [
        {'flow': '步驟1文字', 'check': '檢驗要點...', 'doc': '相關文件...',
         'mark': '', 'type': 'process'},
        ...
      ]
      - mark: '※'（自主查驗停留點）/ '＊'（監造停留點）/ '◎'（安衛查驗點）
      - type: process/decision/terminator/data/document（預設 process）

    做法：
      - 建立一個 python-docx 3 欄 N 列表格
      - 欄1 的每列放入一個「小型單節點 DrawML 形狀」（仍可點擊編輯）
      - 欄2、欄3 為純文字
      - 列與列之間的流程箭頭由儲存格下邊框 + 中央箭頭符號示意
    """
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement

    n_rows = 1 + len(items)
    table = doc.add_table(rows=n_rows, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 欄寬
    col_widths_cm = [6.0, 5.5, 5.0]
    for ci, cw in enumerate(col_widths_cm):
        for r in table.rows:
            r.cells[ci].width = Cm(cw)

    # 欄標題（淺灰底、黑字、黑框）
    headers = ['作業流程圖', '檢驗要點', '相關記錄/文件']
    hdr = table.rows[0]
    for ci, htxt in enumerate(headers):
        cell = hdr.cells[ci]
        # 淺灰底
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(_qn('w:val'), 'clear'); shd.set(_qn('w:color'), 'auto'); shd.set(_qn('w:fill'), 'F2F2F2')
        tcPr.append(shd)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(htxt)
        run.bold = True
        run.font.name = font_latin
        rPr = run._element.get_or_add_rPr()
        rPr.rFonts.set(_qn('w:eastAsia'), font_cjk)
        rPr.rFonts.set(_qn('w:ascii'), font_latin); rPr.rFonts.set(_qn('w:hAnsi'), font_latin)
        run.font.size = Pt(font_pt + 1)

    # 每列
    for ri, item in enumerate(items):
        row = table.rows[ri + 1]
        flow_txt  = item.get('flow', '')
        check_txt = item.get('check', '')
        doc_txt   = item.get('doc', '')
        mark      = item.get('mark', '')
        ntype     = item.get('type', 'process')

        # 欄1：DrawML 單節點（透過 register 讓 finalize 處理）
        marker_id = f'{marker_prefix}_r{ri}'
        flow_def = {
            'nodes': [{'id': 'n1', 'label': (mark + flow_txt) if mark else flow_txt, 'type': ntype}],
            'edges': [],
        }
        register_flowchart(marker_id, flow_def)
        cell1 = row.cells[0]
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell1.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'<<<DRAWML_FLOWCHART:{marker_id}>>>')
        run.font.size = Pt(1)

        # 欄2/欄3：純文字
        for ci, txt in [(1, check_txt), (2, doc_txt)]:
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.name = font_latin
            rPr = run._element.get_or_add_rPr()
            rPr.rFonts.set(_qn('w:eastAsia'), font_cjk)
            rPr.rFonts.set(_qn('w:ascii'), font_latin); rPr.rFonts.set(_qn('w:hAnsi'), font_latin)
            run.font.size = Pt(font_pt)

    return table


if __name__ == '__main__':
    # 小型自測：產生 7 張常用流程圖
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--out', default='/tmp/drawml_demo.docx')
    args = parser.parse_args()

    from docx import Document as DocxDoc
    doc = DocxDoc()
    doc.add_heading('DrawML 流程圖 demo（quality_plan 主題）', level=0)

    titles_and_presets = [
        ('材料/設備選定送審流程圖', preset_material_approval()),
        ('施工檢驗流程圖',         preset_work_inspection()),
        ('材料不合格管制流程圖',   preset_nc_material()),
        ('施工不合格管制流程圖',   preset_nc_work()),
        ('矯正措施流程圖',         preset_corrective()),
        ('收文傳送及歸檔流程圖',   preset_incoming_doc()),
        ('發文傳送及歸檔流程圖',   preset_outgoing_doc()),
        ('設備功能運轉檢測程序流程圖', preset_equipment_test()),
        ('品質稽核流程圖',          preset_audit_flow()),
    ]
    for i, (title, flow) in enumerate(titles_and_presets):
        doc.add_paragraph().add_run(title).bold = True
        add_flowchart_placeholder(doc, f'demo_{i}', flow)
        doc.add_paragraph('')

    doc.save(args.out)
    finalize_flowcharts(args.out)
    print(f'Demo saved: {args.out}')


# ─────────────────────────────────────────────────────────────────────────────
# 品質管理組織架構圖（品質計畫專用，v8.5）
# ─────────────────────────────────────────────────────────────────────────────

def build_quality_org_chart(
    label_tech='主任技師\n（待補）',
    label_foreman='工地主任\n（待補）',
    label_qc='品管人員\n（待補）',
    label_safety='職業安全衛生\n管理人員\n（待補）',
    label_sub='各協力廠商',
    label_eng='現場工程師',
    label_admin='行政人員',
    font_cjk='標楷體',
    font_latin='Times New Roman',
    font_pt=10,
    doc_id=999
):
    """品質管理組織架構圖（符合品質計畫金質獎規格）。

    結構：
      主任技師 ─實線─► 工地主任 ─實線─► {品管人員, [現場工程師|行政人員], 職安衛}
      主任技師 ─虛線──────────────────► 品管人員  (品質政策執行)
      品管人員 ◄─點劃線─► 現場工程師           (稽核、查證)
      職安衛   ◄─點劃線─► 行政人員             (稽核、查證)
      現場工程師群組 ─實線─► 各協力廠商

    回傳 w:p 元素（可直接插入 python-docx 的 body）。
    """
    S = 360000  # 1 cm = 360000 EMU

    CW = int(14.5 * S)
    CH = int(10.5 * S)

    BW   = int(3.0 * S)
    BH   = int(1.0 * S)
    BH3  = int(1.5 * S)
    BSUB = int(2.5 * S)

    CX_MAIN = int(7.25 * S)
    CX_QC   = int(1.5  * S)
    CX_ENG  = int(5.5  * S)
    CX_ADMN = int(9.0  * S)
    CX_SAFE = int(12.75 * S)

    Y0 = int(0.2 * S)
    Y1 = int(2.0 * S)
    Y2 = int(3.9 * S)
    Y3 = int(6.5 * S)

    GX  = int(4.0 * S)
    GY  = Y2 - int(0.2 * S)
    GW  = int(6.5 * S)
    GH  = BH + int(0.4 * S)
    GCX = GX + GW // 2

    LX = int(0.1 * S)
    LY = int(7.7 * S)
    LW = int(5.5 * S)
    LH = int(2.3 * S)

    BLACK = '000000'
    GRAY  = 'D9D9D9'

    def bx(cx, w=BW): return cx - w // 2

    def _tbox(x, y, w, h, text, bg=None, bdr=BLACK, bdr_pt=1.25):
        wsp = etree.Element(q(WPS, 'wsp'))
        _sp_lock(wsp)
        spPr = _sub(wsp, WPS, 'spPr')
        _xfrm(spPr, x, y, w, h)
        geom = _sub(spPr, A, 'prstGeom'); geom.set('prst', 'rect'); _sub(geom, A, 'avLst')
        _fill(spPr, bg)
        if bdr is not None and bdr_pt > 0:
            ln = _sub(spPr, A, 'ln'); ln.set('w', str(int(bdr_pt * 12700))); _fill(ln, bdr)
        else:
            ln = _sub(spPr, A, 'ln'); ln.set('w', '0'); _sub(ln, A, 'noFill')
        txbx = _sub(wsp, WPS, 'txbx')
        cnt  = _sub(txbx, W, 'txbxContent')
        for line in (text or '').split('\n'):
            p   = _sub(cnt, W, 'p')
            pPr = _sub(p, W, 'pPr')
            _sub(pPr, W, 'jc').set(q(W, 'val'), 'center')
            sp2 = _sub(pPr, W, 'spacing')
            sp2.set(q(W, 'before'), '0'); sp2.set(q(W, 'after'), '0')
            r   = _sub(p, W, 'r')
            rPr = _sub(r, W, 'rPr')
            if font_cjk or font_latin:
                rf = _sub(rPr, W, 'rFonts')
                if font_latin: rf.set(q(W, 'ascii'), font_latin); rf.set(q(W, 'hAnsi'), font_latin)
                if font_cjk:   rf.set(q(W, 'eastAsia'), font_cjk)
            _sub(rPr, W, 'color').set(q(W, 'val'), BLACK)
            sz = str(int(font_pt * 2))
            _sub(rPr, W, 'sz').set(q(W, 'val'), sz); _sub(rPr, W, 'szCs').set(q(W, 'val'), sz)
            t = _sub(r, W, 't'); t.text = line
        bpr = _sub(wsp, WPS, 'bodyPr'); bpr.set('anchor', 'ctr')
        bpr.set('lIns', '45720'); bpr.set('rIns', '45720')
        bpr.set('tIns', '36576'); bpr.set('bIns', '36576'); bpr.set('wrap', 'square')
        _sub(bpr, A, 'noAutofit')
        return wsp

    def _conn(x1, y1, x2, y2, style='solid'):
        lx = min(x1, x2); ly = min(y1, y2)
        cx = max(abs(x2 - x1), 914)
        cy = max(abs(y2 - y1), 914)
        wsp = etree.Element(q(WPS, 'wsp'))
        _sub(wsp, WPS, 'cNvCnPr')
        spPr = _sub(wsp, WPS, 'spPr')
        _xfrm(spPr, lx, ly, cx, cy, flipH=(x2 < x1), flipV=(y2 < y1))
        geom = _sub(spPr, A, 'prstGeom')
        prst = 'straightConnector1' if (x1 == x2 or y1 == y2) else 'bentConnector3'
        geom.set('prst', prst); _sub(geom, A, 'avLst')
        ln = _sub(spPr, A, 'ln'); ln.set('w', '15875')
        _fill(ln, BLACK)
        if style in ('dashed', 'dashdot'):
            _sub(ln, A, 'prstDash').set('val', 'lgDash' if style == 'dashed' else 'lgDashDot')
        if style == 'dashdot':
            hd = _sub(ln, A, 'headEnd'); hd.set('type', 'arrow'); hd.set('w', 'med'); hd.set('len', 'med')
            tl = _sub(ln, A, 'tailEnd'); tl.set('type', 'arrow'); tl.set('w', 'med'); tl.set('len', 'med')
        else:
            _sub(ln, A, 'headEnd').set('type', 'none')
            _sub(ln, A, 'tailEnd').set('type', 'none')
        _sub(wsp, WPS, 'bodyPr')
        return wsp

    canvas = etree.Element(q(WPC, 'wpc'))
    _sub(canvas, WPC, 'bg')
    _sub(canvas, WPC, 'whole')

    def add(el): canvas.append(el)

    # 0. 灰色群組背景
    add(_tbox(GX, GY, GW, GH, '', bg=GRAY, bdr=None, bdr_pt=0))

    # 1. 節點方塊
    add(_tbox(bx(CX_MAIN),       Y0, BW,   BH,  label_tech))
    add(_tbox(bx(CX_MAIN),       Y1, BW,   BH,  label_foreman))
    add(_tbox(bx(CX_QC),         Y2, BW,   BH,  label_qc))
    add(_tbox(bx(CX_ENG, BSUB),  Y2, BSUB, BH,  label_eng))
    add(_tbox(bx(CX_ADMN, BSUB), Y2, BSUB, BH,  label_admin))
    add(_tbox(bx(CX_SAFE, BW),   Y2, BW,   BH3, label_safety))
    add(_tbox(bx(CX_MAIN),       Y3, BW,   BH,  label_sub))

    # 2. 實線（隸屬）
    add(_conn(CX_MAIN, Y0 + BH,   CX_MAIN, Y1,   'solid'))
    add(_conn(CX_MAIN, Y1 + BH,   CX_QC,   Y2,   'solid'))
    add(_conn(CX_MAIN, Y1 + BH,   GCX,     GY,   'solid'))
    add(_conn(CX_MAIN, Y1 + BH,   CX_SAFE, Y2,   'solid'))
    add(_conn(GCX,     GY + GH,   CX_MAIN, Y3,   'solid'))

    # 3. 虛線（品質政策執行）
    add(_conn(bx(CX_MAIN), Y0 + BH // 2, CX_QC, Y2, 'dashed'))

    # 4. 點劃線雙向箭頭（稽核、查證）
    add(_conn(bx(CX_QC) + BW,        Y2 + BH // 2,
              bx(CX_ENG, BSUB),      Y2 + BH // 2, 'dashdot'))
    add(_conn(bx(CX_ADMN, BSUB) + BSUB, Y2 + BH // 2,
              bx(CX_SAFE, BW),           Y2 + BH // 2, 'dashdot'))

    # 5. 圖例
    add(_tbox(LX, LY, LW, LH, '', bg=None, bdr=BLACK, bdr_pt=0.75))
    add(_tbox(LX + int(0.2*S), LY + int(0.1*S),
              int(1.5*S), int(0.5*S), '圖例：', bg=None, bdr=None, bdr_pt=0))

    LLINE_X1 = LX + int(0.3 * S)
    LLINE_X2 = LX + int(1.8 * S)
    LTXT_X   = LX + int(2.0 * S)
    LTXT_W   = int(3.3 * S)
    LTXT_H   = int(0.5 * S)
    for y_off, style, lbl in [
        (int(0.70 * S), 'solid',   '隸屬'),
        (int(1.25 * S), 'dashed',  '品質政策執行'),
        (int(1.80 * S), 'dashdot', '稽核、查證'),
    ]:
        YL = LY + y_off
        add(_conn(LLINE_X1, YL, LLINE_X2, YL, style))
        add(_tbox(LTXT_X, YL - int(0.25*S), LTXT_W, LTXT_H, lbl,
                  bg=None, bdr=None, bdr_pt=0))

    return wrap_canvas_in_paragraph(canvas, CW, CH, doc_id=doc_id)


# ─────────────────────────────────────────────────────────────────────────────
# preset_quality_org_chart — 品管組織架構圖快捷封裝
# ─────────────────────────────────────────────────────────────────────────────

def preset_quality_org_chart(
    label_tech   = '主任技師\n（待補）',
    label_foreman= '工地主任\n（待補）',
    label_qc     = '品管人員\n（待補）',
    label_safety = '職業安全衛生\n管理人員\n（待補）',
    label_sub    = '各協力廠商',
    label_eng    = '現場工程師',
    label_admin  = '行政人員',
    font_cjk  = '標楷體',
    font_latin= 'Arial',
    font_pt   = 10,
):
    """回傳品管組織架構圖的 flow_def dict（type='quality_org_chart'）。
    由 finalize_flowcharts() 偵測後呼叫 build_quality_org_chart()，
    不含主辦機關 / 設計單位 / 監造單位。
    """
    return {
        'type':          'quality_org_chart',
        'label_tech':    label_tech,
        'label_foreman': label_foreman,
        'label_qc':      label_qc,
        'label_safety':  label_safety,
        'label_sub':     label_sub,
        'label_eng':     label_eng,
        'label_admin':   label_admin,
        'font_cjk':      font_cjk,
        'font_latin':    font_latin,
        'font_pt':       font_pt,
    }
