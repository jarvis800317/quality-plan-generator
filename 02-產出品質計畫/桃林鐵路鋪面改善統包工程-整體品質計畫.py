#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
整體品質計畫產生器 — 道路/鐵路鋪面改善工程（6章制）
工程：桃林鐵路鋪面改善統包工程（成功路至健行路）
產出：DOCX 整體品質計畫
"""

import os, io, sys, zipfile, re, copy
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
from fontTools.ttLib import TTFont as _TTFont

from drawml_flowchart import (
    add_flowchart_placeholder, add_three_col_flow,
    finalize_flowcharts,
    preset_material_approval, preset_work_inspection,
    preset_nc_material, preset_nc_work,
    preset_corrective, preset_incoming_doc, preset_outgoing_doc,
    preset_equipment_test, preset_audit_flow,
    preset_quality_org_chart,
)

# ═══════════════════════════════════════════════════════
# ⭐ 工程設定區（桃林鐵路鋪面改善統包工程）
# ═══════════════════════════════════════════════════════

ENG_NAME    = "桃林鐵路鋪面改善統包工程（成功路至健行路）"
ENG_OWNER   = "桃園市政府養謢工程處"
ENG_DESIGN  = "拓緯工程顧問有限公司（設計技師 洪仲德）"
ENG_SUPER   = "拓緯工程顧問有限公司（監造技師 洪仲德）"
ENG_CONTR   = "（承攬廠商名稱）"
ENG_SITE    = "桃園市桃林鐵路沿線（成功路至健行路）"
ENG_DAYS    = "210"                     # 日曆天
ENG_AMOUNT  = "新臺幣46,014,931元整"
ENG_FOREMAN = "（工地主任姓名）"
ENG_QA      = "（品管人員姓名）（結業證書字號：）"
ENG_TECH    = "（主任技師姓名）（技師證號：）"

# ── DOCX 字型設定（依監造計畫分析：標楷體 + Arial）──
DOCX_FONT_CJK   = '標楷體'
DOCX_FONT_LATIN = 'Arial'

# ── DOCX 字型大小（pt）──
DOCX_SZ_BODY        = 12
DOCX_SZ_H1          = 20
DOCX_SZ_H2          = 16
DOCX_SZ_H3          = 14
DOCX_SZ_TOC         = 12
DOCX_SZ_TABLE       = 11
DOCX_SZ_COVER_MAIN  = 32
DOCX_SZ_COVER_PROJ  = 24
DOCX_SZ_COVER_INFO  = 16
DOCX_SZ_NOTE        = 10

# ── 章節旗標（6章制：1,000萬～5,000萬工程）──
CHAPTER_FLAGS = {
    'scope':      True,   # 壹　計畫範圍
    'mgmt':       True,   # 貳　管理權責及分工
    'method':     False,  # 參　施工要領（❌未達5,000萬不強制）
    'standard':   True,   # 肆　品質管理標準
    'inspection': True,   # 伍　材料與設備及施工檢驗程序
    'equipment':  False,  # 陸　設備功能運轉檢測程序及標準（❌本工程無機電設備）
    'selfcheck':  True,   # 柒　自主檢查表
    'nc':         False,  # 捌　不合格品之管制（❌未達5,000萬不強制）
    'ca':         False,  # 玖　矫正與預防措施（❌未達5,000萬不強制）
    'audit':      False,  # 拾　內部品質稽核（❌未達5,000萬不強制）
    'doc':        True,   # 拾壹　文件紀錄管理系統
}

# ── 輸出路徑（動態計算）──
_WORKSPACE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(_WORKSPACE, f"{ENG_NAME}-整體品質計畫.docx")

# ── PIL 字型（固定路徑）──
FONT_CJK   = "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf"
FONT_ASCII = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"

# ═══════════════════════════════════════════════════════
# 輔助函式
# ═══════════════════════════════════════════════════════

def set_run_font(run, size_pt, bold=False, color=(0,0,0)):
    run.font.name = DOCX_FONT_LATIN
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    rFonts.set(qn('w:ascii'),    DOCX_FONT_LATIN)
    rFonts.set(qn('w:hAnsi'),    DOCX_FONT_LATIN)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.page_break_before = True
    for run in p.runs:
        run.font.name = DOCX_FONT_LATIN
        run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    return p

def body(doc, text, indent=0, bold=False):
    p = doc.add_paragraph()
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, DOCX_SZ_BODY, bold=bold)
    return p

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    illegal_tags = {'noWrap', 'tcMar', 'vAlign', 'hideMark'}
    first_illegal = next(
        (child for child in tcPr
         if child.tag.split('}')[-1] in illegal_tags), None)
    if first_illegal is not None:
        tcPr.insertBefore(shd, first_illegal)
    else:
        tcPr.append(shd)

def make_table(doc, headers, rows, col_widths, header_bg='F2F2F2', font_size=None):
    if font_size is None:
        font_size = DOCX_SZ_TABLE
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row = table.rows[0]
    for ci, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[ci]
        cell.width = Cm(w)
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, font_size, bold=True)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, (cell_txt, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[ci]
            cell.width = Cm(w)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_txt))
            set_run_font(run, font_size)
    return table

def add_table_caption(doc, chapter, seq, name):
    caption_text = f"表{chapter}-{seq} {name}"
    p = doc.add_paragraph(caption_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    set_run_font(run, DOCX_SZ_TABLE, bold=True)
    return p

def add_figure_caption(doc, chapter, seq, name):
    caption_text = f"圖{chapter}-{seq} {name}"
    p = doc.add_paragraph(caption_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    set_run_font(run, DOCX_SZ_BODY)
    return p

def add_img(doc, img_bytes, width_cm=14):
    if isinstance(img_bytes, tuple):
        if img_bytes[0] == '__DRAWML_FLOW__':
            flow_def = img_bytes[1]
            marker_id = f"flow_{id(flow_def)}"
            add_flowchart_placeholder(doc, marker_id, flow_def)
            return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_bytes, width=Cm(width_cm))
    return p

def _add_right_tab(paragraph, pos_cm=15.0):
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:tabs')):
        pPr.remove(old)
    tabs_el = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'),    'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'),    str(int(pos_cm * 567)))
    tabs_el.append(tab)
    pPr.append(tabs_el)

def make_toc(doc, entries, right_pos_cm=15.0):
    for entry in entries:
        p = doc.add_paragraph()
        _add_right_tab(p, right_pos_cm)
        if entry['level'] == 2:
            p.paragraph_format.left_indent = Cm(1.0)
        run_title = p.add_run(entry['title'])
        set_run_font(run_title, DOCX_SZ_TOC, bold=(entry['level']==1))
        run_page = p.add_run(f'\t{entry["page"]}')
        set_run_font(run_page, DOCX_SZ_TOC, bold=(entry['level']==1))
    return doc

def _add_page_num_field(paragraph):
    run = paragraph.add_run()
    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    run._r.extend([fc_begin, instr, fc_sep, fc_end])
    return paragraph

def _set_pg_num_type(sectPr, fmt, start=1):
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'),  fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

def setup_section_footer(section, fmt):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        for r in p.runs:
            r.text = ''
    if fmt == 'none':
        return
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_num_field(p)
    _set_pg_num_type(section._sectPr, fmt, start=1)

def insert_section_break(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')
    pgSz.set(qn('w:h'), '16838')
    sectPr.append(pgSz)
    pgType = OxmlElement('w:type')
    pgType.set(qn('w:val'), 'nextPage')
    sectPr.append(pgType)
    pPr.append(sectPr)
    return p

def fix_docx(out_path):
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path, 'r') as zi, \
         zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zo:
        for item in zi.infolist():
            data = zi.read(item.filename)
            if item.filename == 'word/settings.xml':
                data = re.sub(b'<w:zoom[^/]*/>', b'<w:zoom w:percent="100"/>', data)
            zo.writestr(item, data)
    os.replace(tmp, out_path)
    print(f"✅ DOCX 已儲存並修正：{out_path}")

def cover_para(doc, text, size, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)
    return p

# ═══════════════════════════════════════════════════════
# 主要施工工項清單（桃林鐵路）
# ═══════════════════════════════════════════════════════

WORK_ITEMS = [
    ("假設工程", "施工圍籬、交通維持設施、告示牌設置"),
    ("拆除工程", "既有舊化木棧道拆除、既有設施清除"),
    ("放樣整地工程", "基準點設置、地面整平、清除及掘除"),
    ("開挖及基礎工程", "橋面拓寬段開挖、鋼筋基礎施作"),
    ("混凝土工程", "210kgf/cm²混凝土澆置、振搗、养护"),
    ("鋼線網鋪設工程", "銲接鋼線網 D=6.00mm 15×15cm鋪設"),
    ("鋼軌除鏽防滑工程", "既有鋼軌除鏽及防滑塗層施工"),
    ("鋪面標線工程", "成型標線及熱處理聚酯反光標線施作"),
    ("太陽能星光燈安裝工程", "新建太陽能星光燈基礎及燈具安裝"),
    ("自主檢查及完工整理", "自主檢查、缺失改善、完工清理"),
]

# ═══════════════════════════════════════════════════════
# 工程主要施工項目及數量表（表1-1）
# ═══════════════════════════════════════════════════════

MAIN_QUANTITIES = [
    ("1", "210kgf/cm² 混凝土", "m³", "1,789.56"),
    ("2", "銲接鋼線網 D=6.00mm, 15×15cm", "m²", "8,128.80"),
    ("3", "鋼軌除鏽及防滑塗層", "m", "12,946.40"),
    ("4", "成型標線及熱處理聚酯反光標線", "組/m²", "36/106.2"),
    ("5", "鋼筋 SD280（含彎紮組立）", "kg", "113.0"),
]

# ═══════════════════════════════════════════════════════
# 品質管理標準表內容（混凝土工程）
# ═══════════════════════════════════════════════════════

def get_standard_rows_concrete():
    return [
        ("放樣", "位置、高程", "依圖說±2cm", "施工前", "水準儀、捲尺", "每次放樣", "放樣紀錄", "重新放樣"),
        ("混凝土進場", "坍度、溫度", "坍度15±3.5cm，溫度13~32°C", "進場時", "坍度試驗、溫度計", "每車", "試驗記錄", "退料"),
        ("鋼筋組立", "號數、間距、保護層", "依圖說±10mm", "澆置前", "捲尺量測", "每次", "鋼筋檢查表", "校正後重驗"),
        ("模板組立", "尺寸、垂直度", "高程±5mm", "組立後", "水準儀、垂球", "每次", "模板檢查表", "修正後重驗"),
        ("混凝土澆置", "澆置連續性、振搗", "間距≤45cm，5~10秒/處", "澆置中", "目視、計時", "全程", "施工日誌", "補足振搗"),
        ("試體取樣", "試體數量", "每20~100m³取1組", "澆置同時", "試體模", "每車", "試驗報告", "追加取樣"),
        ("混凝土养护", "养护方式及時間", "濕潤养护≥7天", "拆模後", "目視", "每日", "养护紀錄", "延長养护"),
    ]

def get_standard_rows_steel():
    return [
        ("鋼筋進場", "規格、數量", "符合CNS 560", "進場時", "目視、量測", "每批", "材料報驗單", "退貨"),
        ("鋼筋加工", "切斷尺寸、彎曲角度", "依圖說±5mm", "加工前", "捲尺、樣板", "每件", "加工紀錄", "重新加工"),
    ]

def get_standard_rows_pavement():
    return [
        ("鋼線網鋪設", "鋪設平整、搭接長度", "搭接≥20cm", "鋪設中", "目視、捲尺", "每次", "施工紀錄", "重新鋪設"),
        ("標線施作", "位置、反光性", "依圖說，抗滑值BPN 50以上", "施作後", "抗滑值測定", "每處", "試驗報告", "重新施作"),
    ]

# ═══════════════════════════════════════════════════════
# 自主檢查表底部固定格式
# ═══════════════════════════════════════════════════════

SELFCHECK_FOOTER = [
    "缺失複查結果：",
    "□已完成改善（檢附改善前中後照片）",
    "□未完成改善，填具「缺失改善追蹤表」追蹤改善",
    "",
    "複查日期：　　年　　月　　日　　複查人員職稱：　　　　　簽名：",
    "",
    "備註：",
    "1.☆表示檢驗停留點之檢驗項目。",
    "2.合格「○」，不合格「╳」，不需檢查「／」。",
    "3.檢查標準應具體明確或量化尺寸。",
    "4.嚴重缺失未及時改善，應填具「不合格品管制總表」追蹤改善。",
    "5.本表由工地現場施工人員實地檢查後覈實記載簽認。",
    "",
    "工地主任（工地負責人）簽名：＿＿＿＿　　現場施工人員（檢查人員）簽名：＿＿＿＿",
]

# ═══════════════════════════════════════════════════════
# 主程式
# ═══════════════════════════════════════════════════════

def build_doc():
    doc = Document()

    # ── 全域頁面設定（A4）──
    for sec in doc.sections:
        sec.page_width    = Cm(21)
        sec.page_height   = Cm(29.7)
        sec.left_margin   = Cm(3)
        sec.right_margin  = Cm(2.5)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2)

    # ── 樣式初始化 ──
    normal = doc.styles['Normal']
    normal.font.name = DOCX_FONT_LATIN
    normal.font.size = Pt(DOCX_SZ_BODY)
    _nr = normal.element.get_or_add_rPr()
    _nf = _nr.get_or_add_rFonts()
    _nf.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    _nf.set(qn('w:ascii'),    DOCX_FONT_LATIN)
    _nf.set(qn('w:hAnsi'),    DOCX_FONT_LATIN)

    HEADING_COLORS = {1:(0,0,0), 2:(0,0,0), 3:(0,0,0)}
    for lvl, sz, bold_flag in [(1, DOCX_SZ_H1, True),(2, DOCX_SZ_H2, True),(3, DOCX_SZ_H3, True)]:
        s = doc.styles[f'Heading {lvl}']
        s.font.name = DOCX_FONT_LATIN
        s.font.size = Pt(sz)
        s.font.bold = bold_flag
        s.font.color.rgb = RGBColor(*HEADING_COLORS[lvl])
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after  = Pt(6)
        if lvl == 1:
            s.paragraph_format.page_break_before = True
        _hr = s.element.get_or_add_rPr()
        _hf = _hr.get_or_add_rFonts()
        _hf.set(qn('w:eastAsia'), DOCX_FONT_CJK)
        _hf.set(qn('w:ascii'),    DOCX_FONT_LATIN)
        _hf.set(qn('w:hAnsi'),    DOCX_FONT_LATIN)

    # ════════════════════════════════════
    # 封面
    # ════════════════════════════════════
    for _ in range(6): doc.add_paragraph()
    cover_para(doc, "整體品質計畫", DOCX_SZ_COVER_MAIN, bold=True, space_before=0, space_after=18)
    cover_para(doc, ENG_NAME, DOCX_SZ_COVER_PROJ, bold=True, space_before=0, space_after=24)
    cover_para(doc, f"主辦機關：{ENG_OWNER}", DOCX_SZ_COVER_INFO, space_before=6, space_after=6)
    cover_para(doc, f"設計單位：{ENG_DESIGN}", DOCX_SZ_COVER_INFO, space_before=0, space_after=6)
    cover_para(doc, f"監造單位：{ENG_SUPER}", DOCX_SZ_COVER_INFO, space_before=0, space_after=6)
    cover_para(doc, f"承攬廠商：{ENG_CONTR}", DOCX_SZ_COVER_INFO, space_before=0, space_after=6)
    for _ in range(3): doc.add_paragraph()
    cover_para(doc, "中華民國　　　年　　月", DOCX_SZ_COVER_INFO, space_before=0, space_after=0)
    insert_section_break(doc)

    # ════════════════════════════════════
    # 目錄（僅產出已啟用的章節）
    # ════════════════════════════════════
    p_toc_title = doc.add_paragraph("目    錄")
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_toc_title.runs[0]
    set_run_font(run, DOCX_SZ_H2, bold=True)

    toc_entries = []
    ch = 1  # 動態章號

    if CHAPTER_FLAGS['scope']:
        toc_entries += [
            {'level':1,'title':'壹、計畫範圍',               'page':1},
            {'level':2,'title':'  一、依據',                 'page':1},
            {'level':2,'title':'  二、工程概要',             'page':1},
            {'level':2,'title':'  三、工程主要施工項目及數量', 'page':2},
            {'level':2,'title':'  四、適用對象',              'page':2},
            {'level':2,'title':'  五、名詞定義',              'page':2},
        ]
        ch += 1

    if CHAPTER_FLAGS['mgmt']:
        toc_entries += [
            {'level':1,'title':'貳、管理權責及分工',         'page':3},
            {'level':2,'title':'  一、品管組織架構',         'page':3},
            {'level':2,'title':'  二、工作職掌',             'page':3},
            {'level':2,'title':'  三、人員配置',              'page':4},
            {'level':2,'title':'  四、應用表單',              'page':4},
        ]
        ch += 1

    if CHAPTER_FLAGS['standard']:
        toc_entries += [
            {'level':1,'title':'肆、品質管理標準',           'page':5},
            {'level':2,'title':'  一、品質管理標準訂定',      'page':5},
            {'level':2,'title':'  二、各工項品質管理標準表',   'page':5},
            {'level':2,'title':'  三、應用表單',              'page':6},
        ]
        ch += 1

    if CHAPTER_FLAGS['inspection']:
        toc_entries += [
            {'level':1,'title':'伍、材料與設備及施工檢驗程序', 'page':7},
            {'level':2,'title':'  一、材料與設備檢驗程序',    'page':7},
            {'level':2,'title':'  二、施工檢驗程序',          'page':8},
            {'level':2,'title':'  三、應用表單',              'page':9},
        ]
        ch += 1

    if CHAPTER_FLAGS['selfcheck']:
        toc_entries += [
            {'level':1,'title':'柒、自主檢查表',             'page':10},
            {'level':2,'title':'  一、自主檢查表之訂定',       'page':10},
            {'level':2,'title':'  二、自主檢查表之執行',       'page':10},
            {'level':2,'title':'  三、各工項自主檢查表',       'page':11},
            {'level':2,'title':'  四、應用表單',               'page':13},
        ]
        ch += 1

    if CHAPTER_FLAGS['doc']:
        toc_entries += [
            {'level':1,'title':'拾壹、文件紀錄管理系統',       'page':14},
            {'level':2,'title':'  一、文件管理系統',          'page':14},
            {'level':2,'title':'  二、紀錄管理作業程序',      'page':15},
            {'level':2,'title':'  三、文件紀錄移轉及存檔',     'page':15},
        ]

    make_toc(doc, toc_entries)
    insert_section_break(doc)

    # ════════════════════════════════════
    # 第一章 計畫範圍
    # ════════════════════════════════════
    if CHAPTER_FLAGS['scope']:
        heading(doc, "第一章　計畫範圍", 1)

        heading(doc, "一、依據", 2)
        body(doc, "　　本品質計畫撰寫之依據，係以工程契約（含施工規範及設計圖說）、技師法、建築法、建築師法、營造業法、公共工程專業技師簽證規則、職業安全衛生法、職業安全衛生設施規則、營造安全衛生設施標準、環境保護相關規定、公共工程施工綱要規範、公共工程施工階段契約約定權責分工表（無委託專案管理廠商）、本公司品質系統作業規定及監造計畫編擬。")

        heading(doc, "二、工程概要", 2)
        body(doc, f"　　本案為桃林鐵路鋪面改善統包工程，位於桃園市桃林鐵路沿線（成功路至健行路），主要工作內容為拆除既有舊化木棧道、新建人行步道及自行車道，並針對跨橋段（約3k+800處）進行橋面拓寬，同時新增太陽能星光燈等設施。")

        doc.add_paragraph()
        info_data = [
            ("工程名稱", ENG_NAME),
            ("主辦機關", ENG_OWNER),
            ("設計單位", "拓緯工程顧問有限公司"),
            ("監造單位", "拓緯工程顧問有限公司"),
            ("承攬廠商", ENG_CONTR),
            ("工程地點", ENG_SITE),
            ("契約工期", f"{ENG_DAYS}日曆天"),
            ("契約金額", ENG_AMOUNT),
        ]
        for label, val in info_data:
            body(doc, f"{label}：{val}")

        heading(doc, "三、工程主要施工項目及數量", 2)
        add_table_caption(doc, 1, 1, "工程主要施工項目及數量表")
        make_table(doc,
            headers=["項次", "項目", "單位", "數量"],
            rows=MAIN_QUANTITIES,
            col_widths=[2.0, 9.5, 2.0, 2.5])

        heading(doc, "四、適用對象", 2)
        body(doc, "　　本品質計畫之適用對象為本公司（承攬廠商）辦理本工程之所有施工及管理人員、協力廠商及材料供應廠商等。凡從事本工程之施工、材料採購及品質管制相關作業之人員，均應遵循本品質計畫之規定。")

        heading(doc, "五、名詞定義", 2)
        definitions = [
            ("本工程", ENG_NAME),
            ("主辦機關", f"{ENG_OWNER}（甲方）"),
            ("設計單位", "拓緯工程顧問有限公司（設計技師 洪仲德）"),
            ("監造單位", "拓緯工程顧問有限公司（監造技師 洪仲德）"),
            ("承攬廠商、本公司", f"{ENG_CONTR}（乙方）"),
            ("工地主任", "受聘於本公司，擔任本工程之工地事務及施工管理之人員。"),
            ("品管人員", "本公司專責辦理品質計畫與執行查證工作之人員，須取得公共工程品質管理訓練課程結業證書。"),
            ("主任技師", "受聘於本公司之執業技師，負責施工技術指導及施工安全。"),
            ("協力廠商", "材料及施工分包廠商。"),
            ("施工計畫", "承包商依據本工程特性所製作之施工組織、工期、人力、材料、機具、職業安全衛生及環保等之整體計畫。"),
            ("品質管制", "為確保工程品質能符合規定要求，所建立之系統性作業程序與管理措施。"),
            ("檢驗停留點", "凡屬隱蔽部分於掩蓋前、永久性工程建造前或材料使用前，關係施工品質之控制點，均應由承包商提出申請，並由監造單位代表會同作各種查驗。"),
            ("自主檢查", "由承攬廠商按特定規則所進行之檢驗，以確認施工成果符合規定。"),
            ("不符合", "施工結果或材料設備未符合契約規定之情形。"),
            ("矫正措施", "針對已發現之不符合所採取之改善行動。"),
            ("預防措施", "針對潛在不符合所採取之預防性行動。"),
            ("內部品質稽核", "本公司上級主管帶隊至工地進行之品質稽核。"),
        ]
        for i, (term, definition) in enumerate(definitions, 1):
            roman = ["(一)","(二)","(三)","(四)","(五)","(六)","(七)","(八)","(九)","(十)",
                     "(十一)","(十二)","(十三)","(十四)","(十五)","(十六)","(十七)"][i-1] if i <= 17 else f"({i})"
            body(doc, f"{roman}{term}：{definition}")

    # ════════════════════════════════════
    # 第二章 管理權責及分工
    # ════════════════════════════════════
    if CHAPTER_FLAGS['mgmt']:
        heading(doc, "第二章　管理權責及分工", 1)
        heading(doc, "一、品管組織架構", 2)
        body(doc, "　　本公司依據工程特性及規模，於工地設立品管組織，由工地負責人督導，品管人員執行日常品質管制工作，各專業現場工程師及協力廠商配合執行。各級人員均依據本品質計畫及相關程序文件，確實執行各項品質管理作業。")

        # 品管組織架構圖（使用 DrawML）
        add_flowchart_placeholder(doc, 'org_chart', preset_quality_org_chart(
            label_tech    = f'主任技師\n（{ENG_TECH}）',
            label_foreman = f'工地主任\n（{ENG_FOREMAN}）',
            label_qc      = f'品管人員\n（{ENG_QA}）',
            label_safety  = '職業安全\n衛生管理',
            label_sub     = '各協力廠商',
            label_eng     = '現場工程師',
            label_admin   = '行政人員',
        ))
        add_figure_caption(doc, 2, 1, "品管組織架構圖")

        heading(doc, "二、工作職掌", 2)
        add_table_caption(doc, 2, 1, "工作職掌表")
        make_table(doc,
            headers=["職稱", "工作職掌"],
            rows=[
                ("工地主任", "督導品質管制系統之執行，負責對外溝通協調，確認各項資源到位。"),
                ("主任技師", "督導施工技術，確保施工方法符合設計規範，處理技術問題。"),
                ("品管人員", "依品質計畫執行各項品質管制作業，填報品質紀錄，追踪缺失改善。"),
                ("現場工程師", "協助工地主任執行施工管理，督導施工人員依圖說施作，填報施工日誌。"),
                ("職業安全衛生管理", "執行工地安全衛生管理業務，督導施工人員遵守安全衛生規定。"),
                ("協力廠商", "依合約及品質要求執行施工，並配合本公司之品質管制作業。"),
            ],
            col_widths=[4.0, 12.0])

        heading(doc, "三、人員配置", 2)
        body(doc, f"　　本工程依契約規定及工程特性，配置品管組織人員如下（正式開工前由本公司檢具相關資料提送監造單位轉主辦機關報核）：")
        add_table_caption(doc, 2, 2, "品管組織人員配置表")
        make_table(doc,
            headers=["職稱", "人數", "資格要求", "主要職責"],
            rows=[
                ("工地主任", "1", "相關工程經驗", "工地行政及施工管理"),
                ("主任技師", "1", "土木/結構相關技師證照", "施工技術指導"),
                ("品管人員", "1", "公共工程品質管理訓練結業", "品質管制執行"),
                ("現場工程師", "1~2", "相關科系背景", "施工管理及品質紀錄"),
            ],
            col_widths=[3.5, 1.5, 4.0, 7.0])

        heading(doc, "四、應用表單", 2)
        body(doc, "　　本章相關應用表單包括：品管組織架構圖（圖2-1）、工作職掌表（表2-1）、品管組織人員配置表（表2-2）等。")

    # ════════════════════════════════════
    # 第四章 品質管理標準
    # ════════════════════════════════════
    if CHAPTER_FLAGS['standard']:
        heading(doc, "第四章　品質管理標準", 1)
        heading(doc, "一、品質管理標準訂定", 2)
        body(doc, "　　本公司依據工程契約、施工規範及設計圖說，訂定本工程各施工項目之品質管理標準，作為施工及檢驗之依據。各項管理標準均明確訂定管理項目、管理標準、檢查時機、檢查方法、檢查頻率及不符合時之處置方法。")

        heading(doc, "二、各工項品質管理標準表", 2)

        # 混凝土工程
        body(doc, "（一）混凝土工程", bold=True)
        add_table_caption(doc, 4, 1, "混凝土工程施工品質管理標準表")
        make_table(doc,
            headers=["施工流程", "管理項目", "管理標準", "檢查時機", "檢查方法", "檢查頻率", "管理紀錄", "不符合處置"],
            rows=get_standard_rows_concrete(),
            col_widths=[2.5, 2.5, 3.0, 1.5, 2.0, 1.5, 2.0, 2.5])

        # 鋼筋工程
        body(doc, "（二）鋼筋工程", bold=True)
        add_table_caption(doc, 4, 2, "鋼筋工程施工品質管理標準表")
        make_table(doc,
            headers=["施工流程", "管理項目", "管理標準", "檢查時機", "檢查方法", "檢查頻率", "管理紀錄", "不符合處置"],
            rows=get_standard_rows_steel(),
            col_widths=[2.5, 2.5, 3.0, 1.5, 2.0, 1.5, 2.0, 2.5])

        # 鋪面工程
        body(doc, "（三）鋪面及標線工程", bold=True)
        add_table_caption(doc, 4, 3, "鋪面及標線工程施工品質管理標準表")
        make_table(doc,
            headers=["施工流程", "管理項目", "管理標準", "檢查時機", "檢查方法", "檢查頻率", "管理紀錄", "不符合處置"],
            rows=get_standard_rows_pavement(),
            col_widths=[2.5, 2.5, 3.0, 1.5, 2.0, 1.5, 2.0, 2.5])

        heading(doc, "三、應用表單", 2)
        body(doc, "　　本章相關應用表單包括：各工項品質管理標準表（表4-1至表4-3）等。")

    # ════════════════════════════════════
    # 第五章 材料與設備及施工檢驗程序
    # ════════════════════════════════════
    if CHAPTER_FLAGS['inspection']:
        heading(doc, "第五章　材料與設備及施工檢驗程序", 1)
        heading(doc, "一、材料與設備檢驗程序", 2)
        body(doc, "　　承包商應於材料設備進場前，依契約規定及「材料設備送審管制總表」（表5-1）所列項目，提出相關資料向監造單位申請審查，審查通過後方得進場使用。")

        # 材料設備送審流程圖
        add_flowchart_placeholder(doc, 'material_approval', preset_material_approval())
        add_figure_caption(doc, 5, 1, "材料/設備選定送審流程圖")

        body(doc, "（一）材料設備送審作業", bold=True)
        body(doc, "1.承包商應於開工前，依據契約圖說規定，檢討各項材料及設備之規範需求。")
        body(doc, "2.依「材料設備送審管制總表」所列項目，備妥型錄、規格說明、產地證明、相關試驗報告等資料，於規定時限前提送監造單位審查。")
        body(doc, "3.監造單位收到申請後，應於5天內完成審查，並填具「材料設備相關文件審查紀錄表」（表5-2）通知承包商。")
        body(doc, "4.審查合格之材料設備样品，應於工地保存一份作為比對依據。")

        body(doc, "（二）材料設備進場檢驗", bold=True)
        body(doc, "1.材料設備進場時，承包商應填具「材料設備進場報驗單」，會同監造單位人員進行規格及外觀檢驗。")
        body(doc, "2.須抽樣試驗之材料，承包商應填具書面申請單，向監造單位申請會同取樣，並於24小時內完成取樣後送交試驗機構。")
        body(doc, "3.試驗機構應具備TAF認證（或符合契約規定），試驗報告須加蓋騎縫章及試驗員簽章。")
        body(doc, "4.檢驗合格後，方可於工程中使用；不合格者應立即退貨或改善後重新申請檢驗。")

        # 材料及設備品質管制總表
        add_table_caption(doc, 5, 1, "材料及設備品質管理標準表")
        make_table(doc,
            headers=["項次", "材料/設備名稱", "品質標準", "檢驗方式", "頻率", "依據"],
            rows=[
                ("1", "210kgf/cm²混凝土", "抗壓強度符合規範；坍度15±3.5cm；溫度13~32°C", "抽樣試驗", "每20~100m³取1組試體", "CNS 1230"),
                ("2", "銲接鋼線網", "符合CNS 6919；線徑D=6.00mm, 15×15cm", "抽樣試驗", "每進場批次", "CNS 6919"),
                ("3", "鋼筋 SD280", "符合CNS 560；SD280", "抽樣試驗", "每批", "CNS 560"),
                ("4", "防滑標線材料", "抗滑值BPN 50以上；符合CNS 1333", "抽樣試驗", "每處", "CNS 1333"),
            ],
            col_widths=[1.0, 4.0, 5.5, 2.0, 2.5, 1.5])

        heading(doc, "二、施工檢驗程序", 2)
        body(doc, "　　各項施工依品質管理標準所訂之檢驗時機及方法，由品管人員會同工地主任執行檢驗，檢驗結果詳實記錄於「施工檢驗紀錄表」。")

        # 施工抽查流程圖
        add_flowchart_placeholder(doc, 'work_inspection', preset_work_inspection())
        add_figure_caption(doc, 5, 2, "施工抽查流程圖")

        body(doc, "（一）一般施工檢驗", bold=True)
        body(doc, "1.各工項施工前，須先確認材料規格、施工圖說及品質管理標準。")
        body(doc, "2.施工中依品質管理標準之檢查時機及頻率執行檢驗，並詳實記錄。")
        body(doc, "3.發現不符合項目時，應立即開立「施工抽查缺失改善通知單」，限期改善後再行施工。")

        body(doc, "（二）檢驗停留點檢驗", bold=True)
        body(doc, "1.混凝土澆置前、鋼筋組立完成後、模板組立完成後等檢驗停留點，須會同監造單位人員到場查驗。")
        body(doc, "2.承包商應於預定檢驗時機24小時前通知監造單位。")
        body(doc, "3.檢驗合格後方得繼續下一步驟施工。")

        heading(doc, "三、應用表單", 2)
        body(doc, "　　本章相關應用表單包括：材料設備送審管制總表（表5-1）、材料設備相關文件審查紀錄表（表5-2）、材料設備進場報驗單（表5-3）、材料設備品質抽驗紀錄表（表5-4）等。")

    # ════════════════════════════════════
    # 第七章 自主檢查表
    # ════════════════════════════════════
    if CHAPTER_FLAGS['selfcheck']:
        heading(doc, "第七章　自主檢查表", 1)
        heading(doc, "一、自主檢查表之訂定", 2)
        body(doc, "　　本公司依據契約規定及施工品質管理標準，為本工程各主要工項訂定自主檢查表。自主檢查表明定檢查項目、檢查標準及不合格時之處置方法，由施工人員於施工過程中切實執行檢查並記錄，以確保施工品質。")

        heading(doc, "二、自主檢查表之執行", 2)
        body(doc, "　　各工項自主檢查表由施工人員（現場工程師或領班）於施工前、施工中及施工後依規定時機執行檢查，檢查結果由工地主任及品管人員複核。發現不合格項目時，應即時改善並追蹤至完全改善為止。")

        heading(doc, "三、各工項自主檢查表", 2)

        # 混凝土工程自主檢查表
        body(doc, "（一）混凝土工程自主檢查表（柒-1）", bold=True)
        add_table_caption(doc, 7, 1, "混凝土工程施工自主檢查表")
        make_table(doc,
            headers=["檢查項目", "檢查標準", "檢查結果", "改善追蹤"],
            rows=[
                ("假設工程、交通維持", "依核定交通維持計畫設置圍籬及告示牌", "", ""),
                ("基準點/導線點設置", "位置準確，設置於不受施工影響處", "", ""),
                ("鋼筋進場檢驗", "規格、數量符合圖說，無銹蝕", "", ""),
                ("鋼筋加工", "尺寸及彎曲角度符合圖說", "", ""),
                ("鋼筋組立", "間距及保護層符合圖說", "", ""),
                ("模板組立", "尺寸及垂直度符合圖說", "", ""),
                ("混凝土坍度試驗", "坍度15±3.5cm，溫度13~32°C", "", ""),
                ("混凝土澆置", "連續澆置，振搗確實", "", ""),
                ("試體取樣", "每20~100m³取1組試體", "", ""),
                ("养护", "濕潤养护≥7天", "", ""),
            ],
            col_widths=[4.5, 6.5, 2.0, 3.0])
        for line in SELFCHECK_FOOTER:
            body(doc, line)

        # 鋪面工程自主檢查表
        body(doc, "（二）鋪面工程自主檢查表（柒-2）", bold=True)
        add_table_caption(doc, 7, 2, "鋪面工程施工自主檢查表")
        make_table(doc,
            headers=["檢查項目", "檢查標準", "檢查結果", "改善追蹤"],
            rows=[
                ("鋼線網進場檢驗", "規格D=6.00mm, 15×15cm，符合CNS 6919", "", ""),
                ("鋼線網鋪設", "鋪設平整，搭接長度≥20cm", "", ""),
                ("標線材料進場", "反光性能及抗滑值符合規範", "", ""),
                ("標線施作", "位置正確，線形美觀，反光及抗滑性能符合要求", "", ""),
            ],
            col_widths=[4.5, 6.5, 2.0, 3.0])
        for line in SELFCHECK_FOOTER:
            body(doc, line)

        heading(doc, "四、應用表單", 2)
        body(doc, "　　本章相關應用表單包括：混凝土工程施工自主檢查表（表7-1）、鋪面工程施工自主檢查表（表7-2）等。")

    # ════════════════════════════════════
    # 第十一章 文件紀錄管理系統
    # ════════════════════════════════════
    if CHAPTER_FLAGS['doc']:
        heading(doc, "第十一章　文件紀錄管理系統", 1)
        heading(doc, "一、文件管理系統", 2)
        body(doc, "　　本公司建立完整之文件管理系統，對本工程往來之公文、施工圖說、品質計畫、施工計畫、試驗報告、檢驗紀錄等文件予以分類、編碼、收發及歸檔管理，確保文件之有效性及可追溯性。")

        # 收文流程圖
        add_flowchart_placeholder(doc, 'incoming_doc', preset_incoming_doc())
        add_figure_caption(doc, 11, 1, "收文傳送及歸檔流程圖")

        # 發文流程圖
        add_flowchart_placeholder(doc, 'outgoing_doc', preset_outgoing_doc())
        add_figure_caption(doc, 11, 2, "發文傳送及歸檔流程圖")

        body(doc, "（一）文件編碼原則", bold=True)
        body(doc, "　　本工程文件依類別採用不同代碼：Q-品質文件、M-材料文件、C-施工文件、A-行政文件、D-設計文件。版本編號採用「第A版」、「第B版」…之方式標註。")

        body(doc, "（二）文件修訂管制", bold=True)
        body(doc, "　　品質計畫及施工計畫等文件如有修訂需要，須填寫「文件修訂申請單」，經核定後方得發行新版次，並將舊版回收作廢。")

        heading(doc, "二、紀錄管理作業程序", 2)
        body(doc, "　　各項品質紀錄依來源分類管理，包括：材料設備試驗報告、施工檢驗紀錄、自主檢查表、督察紀錄、施工日誌等。所有紀錄須填寫完整、不得塗改，並於規定期限內完成分類歸檔。")

        add_table_caption(doc, 11, 1, "文件分類代碼表")
        make_table(doc,
            headers=["代碼", "類別", "說明"],
            rows=[
                ("Q", "品質文件", "品質計畫、稽核報告、矫正預防記錄"),
                ("M", "材料文件", "材料設備送審資料、試驗報告、產品型錄"),
                ("C", "施工文件", "施工日誌、施工照片、自主檢查表"),
                ("A", "行政文件", "來函、發文、會議紀錄、估驗計價文件"),
                ("D", "設計文件", "設計圖說、施工圖、相關變更設計"),
            ],
            col_widths=[2.0, 4.0, 10.0])

        heading(doc, "三、文件紀錄移轉及存檔", 2)
        body(doc, "　　工程完成後，本公司將所有品質文件及工程紀錄整理彙編後，依規定移轉予主辦機關。相關紀錄之保存期限依工程契約規定及「公共工程施工品質管理作業要點」辦理，本工程品質紀錄保存年限為驗收後十年。")

    # ════════════════════════════════════
    # 頁尾設定（三段式）
    # ════════════════════════════════════
    sections = list(doc.sections)
    # 封面節（無頁尾）
    setup_section_footer(sections[0], 'none')
    # 目錄節（羅馬數字）- 與封面同節
    setup_section_footer(sections[0], 'lowerRoman')
    # 內文節（阿拉伯數字）
    if len(sections) > 1:
        setup_section_footer(sections[-1], 'decimal')
    else:
        # 只有一節時手動處理
        pass

    doc.save(OUT_PATH)
    print(f"✅ DOCX 初步儲存：{OUT_PATH}")

    # DrawML 流程圖處理
    try:
        finalize_flowcharts(OUT_PATH, theme_name='quality_plan', direction='TB')
    except Exception as e:
        print(f"⚠️ 流程圖處理略過：{e}")

    # 修正 python-docx bug
    fix_docx(OUT_PATH)
    print(f"✅ 品質計畫 DOCX 產生完成：{OUT_PATH}")

if __name__ == '__main__':
    build_doc()
