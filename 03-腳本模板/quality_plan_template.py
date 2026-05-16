#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
整體品質計畫產生器 — 景觀工程（11章制）
工程：桃園市第47期桃園區中平市地重劃工程
產出：DOCX 整體品質計畫
"""

import os, io, sys, zipfile, re, copy
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
from fontTools.ttLib import TTFont as _TTFont

from drawml_flowchart import (
    add_flowchart_placeholder, add_three_col_flow,
    finalize_flowcharts,
    preset_material_approval, preset_work_inspection,
    preset_nc_material, preset_nc_work,
    preset_corrective, preset_incoming_doc, preset_outgoing_doc,
    preset_equipment_test, preset_audit_flow,
    preset_quality_org_chart,
)

# ═══════════════════════════════════════════════════════
# ⭐ 工程設定區（每次新工程必須更新以下所有欄位）
# ═══════════════════════════════════════════════════════

ENG_NAME    = "桃園市第47期桃園區中平市地重劃工程"
ENG_OWNER   = "桃園市政府地政局"
ENG_DESIGN  = "杜風工程服務股份有限公司（設計技師 趙厚任）"
ENG_SUPER   = "杜風工程服務股份有限公司（監造技師 高堉霖）"
ENG_CONTR   = "（承攬廠商名稱）"       # 後補
ENG_SITE    = "桃園市桃園區"
ENG_DAYS    = "540"                     # 日曆天
ENG_AMOUNT  = "新臺幣107,947,000元整"
ENG_FOREMAN = "（工地主任姓名）"        # 後補
ENG_QA      = "（品管人員姓名）（結業證書字號：）"   # 後補
ENG_TECH    = "（主任技師姓名）（技師證號：）"        # 後補

# ── DOCX 字型設定（依監造計畫分析：DFKai-SB + Arial）──
DOCX_FONT_CJK   = '標楷體'
DOCX_FONT_LATIN = 'Arial'

# ── DOCX 字型大小（pt）──
DOCX_SZ_BODY        = 12
DOCX_SZ_H1          = 20
DOCX_SZ_H2          = 16
DOCX_SZ_H3          = 14
DOCX_SZ_TOC         = 12
DOCX_SZ_TABLE       = 11
DOCX_SZ_COVER_MAIN  = 32
DOCX_SZ_COVER_PROJ  = 24
DOCX_SZ_COVER_INFO  = 16
DOCX_SZ_NOTE        = 10

# ── 章節旗標（全選 11 章）──
CHAPTER_FLAGS = {
    'scope':      True,   # 第一章　計畫範圍
    'mgmt':       True,   # 第二章　管理責任及權責分工
    'method':     True,   # 第三章　施工要領
    'standard':   True,   # 第四章　品質管理標準
    'inspection': True,   # 第五章　材料及施工檢驗程序
    'equipment':  True,   # 第六章　設備功能運轉測試
    'selfcheck':  True,   # 第七章　自主檢查表
    'nc':         True,   # 第八章　不合格品之管制
    'ca':         True,   # 第九章　矯正與預防措施
    'audit':      True,   # 第十章　內部品質稽核
    'doc':        True,   # 第十一章 文件記錄管理系統
}

# ── 輸出路徑（動態計算）──
_WORKSPACE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH = os.path.join(_WORKSPACE, "03_產出品質計畫", f"{ENG_NAME}-整體品質計畫.docx")

# ── PIL 字型（sandbox 固定路徑）──
FONT_CJK   = "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf"
FONT_ASCII = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"

# ═══════════════════════════════════════════════════════
# 輔助函式
# ═══════════════════════════════════════════════════════

def set_run_font(run, size_pt, bold=False, color=(0,0,0)):
    run.font.name = DOCX_FONT_LATIN
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    rFonts.set(qn('w:ascii'),    DOCX_FONT_LATIN)
    rFonts.set(qn('w:hAnsi'),    DOCX_FONT_LATIN)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.page_break_before = True
    for run in p.runs:
        run.font.name = DOCX_FONT_LATIN
        run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    return p

def body(doc, text, indent=0, bold=False):
    p = doc.add_paragraph()
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, DOCX_SZ_BODY, bold=bold)
    return p

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    illegal_tags = {'noWrap', 'tcMar', 'vAlign', 'hideMark'}
    first_illegal = next(
        (child for child in tcPr
         if child.tag.split('}')[-1] in illegal_tags), None)
    if first_illegal is not None:
        tcPr.insertBefore(shd, first_illegal)
    else:
        tcPr.append(shd)

def make_table(doc, headers, rows, col_widths, header_bg='F2F2F2', font_size=None):
    if font_size is None:
        font_size = DOCX_SZ_TABLE
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row = table.rows[0]
    for ci, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[ci]
        cell.width = Cm(w)
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, font_size, bold=True)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, (cell_txt, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[ci]
            cell.width = Cm(w)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_txt))
            set_run_font(run, font_size)
    return table

def add_table_caption(doc, chapter, seq, name):
    caption_text = f"表{chapter}-{seq} {name}"
    p = doc.add_paragraph(caption_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    set_run_font(run, DOCX_SZ_TABLE, bold=True)
    return p

def add_figure_caption(doc, chapter, seq, name):
    caption_text = f"圖{chapter}-{seq} {name}"
    p = doc.add_paragraph(caption_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    set_run_font(run, DOCX_SZ_BODY)
    return p

def add_img(doc, img_bytes, width_cm=14):
    if isinstance(img_bytes, tuple):
        if img_bytes[0] == '__DRAWML_FLOW__':
            flow_def = img_bytes[1]
            marker_id = f"flow_{id(flow_def)}"
            add_flowchart_placeholder(doc, marker_id, flow_def)
            return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_bytes, width=Cm(width_cm))
    return p

def _add_right_tab(paragraph, pos_cm=15.0):
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:tabs')):
        pPr.remove(old)
    tabs_el = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'),    'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'),    str(int(pos_cm * 567)))
    tabs_el.append(tab)
    pPr.append(tabs_el)

def make_toc(doc, entries, right_pos_cm=15.0):
    for entry in entries:
        p = doc.add_paragraph()
        _add_right_tab(p, right_pos_cm)
        if entry['level'] == 2:
            p.paragraph_format.left_indent = Cm(1.0)
        run_title = p.add_run(entry['title'])
        set_run_font(run_title, DOCX_SZ_TOC, bold=(entry['level']==1))
        run_page = p.add_run(f'\t{entry["page"]}')
        set_run_font(run_page, DOCX_SZ_TOC, bold=(entry['level']==1))
    return doc

def _add_page_num_field(paragraph):
    run = paragraph.add_run()
    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    run._r.extend([fc_begin, instr, fc_sep, fc_end])
    return paragraph

def _set_pg_num_type(sectPr, fmt, start=1):
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'),  fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

def setup_section_footer(section, fmt):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        for r in p.runs:
            r.text = ''
    if fmt == 'none':
        return
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_num_field(p)
    _set_pg_num_type(section._sectPr, fmt, start=1)

def insert_section_break(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')
    pgSz.set(qn('w:h'), '16838')
    sectPr.append(pgSz)
    pgType = OxmlElement('w:type')
    pgType.set(qn('w:val'), 'nextPage')
    sectPr.append(pgType)
    pPr.append(sectPr)
    return p

def fix_docx(out_path):
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path, 'r') as zi, \
         zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zo:
        for item in zi.infolist():
            data = zi.read(item.filename)
            if item.filename == 'word/settings.xml':
                data = re.sub(b'<w:zoom[^/]*/>', b'<w:zoom w:percent="100"/>', data)
            zo.writestr(item, data)
    os.replace(tmp, out_path)
    print(f"✅ DOCX 已儲存並修正：{out_path}")

def cover_para(doc, text, size, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)
    return p

# ═══════════════════════════════════════════════════════
# 主要施工工項清單
# ═══════════════════════════════════════════════════════

WORK_ITEMS = [
    ("交通臨時安全措施", "交通維持、圍籬標誌設置"),
    ("土方管制工程", "基地開挖、填築、廢方處理"),
    ("測量工程", "放樣控制樁、縱橫斷面測量"),
    ("鋼筋工程", "鋼筋加工、組立、綁紮"),
    ("模板工程", "模板組立、支撐、拆模"),
    ("混凝土工程", "預拌混凝土澆置、振搗、養護"),
    ("瀝青混凝土鋪面工程", "路基整平、瀝青混凝土鋪設"),
    ("預鑄路緣石工程", "基礎施作、緣石安裝"),
    ("鋪面磚（高壓）工程", "底層整平、高壓混凝土磚鋪設"),
    ("磨石子工程", "底層施作、磨石子澆置、研磨"),
    ("抿石子工程", "底層施作、抿石子施作"),
    ("彈性無縫鋪面工程", "底層整平、彈性材料鋪設"),
    ("地坪工程", "地坪底層、面層施作"),
    ("磁磚工程", "底層處理、磁磚鋪貼"),
    ("木作（塑木）工程", "基礎施作、塑木構件組裝"),
    ("公園廁所工程", "基礎施作、牆體結構、設備安裝"),
    ("兒童遊戲場設備工程", "基礎施作、設備安裝、試驗"),
    ("鍍鋅鋼網圍籬工程", "基座確認、圍籬組立"),
    ("公園排水工程", "排水管溝開挖、管線埋設、陰井施作"),
    ("植栽工程", "現地整理、植物種植、養護補植"),
    ("景觀照明工程", "管線配管、基座施作、燈具安裝"),
    ("電力配線工程", "電管配設、電線穿線、接線測試"),
    ("器具設備工程", "設備進場、安裝、功能測試"),
    ("假設工程", "施工圍籬、臨時設施設置"),
    ("汛期工地防災準備", "汛期前檢查、防災設施佈置"),
]

# ═══════════════════════════════════════════════════════
# 工程概述數量表（主要工項）
# ═══════════════════════════════════════════════════════

MAIN_QUANTITIES = [
    ("壹.一", "整地拆除工程", "", ""),
    ("壹.一.1", "清除及掘除", "M2", "15,148.00"),
    ("壹.一.2", "地坪打除（未含運費）", "M2", "11,545.00"),
    ("壹.一.3", "機械拆除，鋼筋混凝土（未含運費）", "M3", "951.00"),
    ("壹.一.16", "廢棄物處理，營建廢棄物清除", "M3", "3,260.00"),
    ("壹.二", "景觀工程", "", ""),
    ("壹.二.1", "景觀設施工程（停車場、鋪面、遊具等）", "式", "1"),
    ("壹.二.2", "木作工程（塑木花架、座椅等）", "式", "1"),
    ("壹.二.3", "遊戲場設施工程", "式", "1"),
    ("壹.二.4", "公園排水工程", "式", "1"),
    ("壹.三", "植栽工程", "式", "1"),
    ("壹.四", "景觀照明工程", "式", "1"),
]

# ═══════════════════════════════════════════════════════
# 品質管理標準表內容（核心工項）
# ═══════════════════════════════════════════════════════

def get_standard_rows_concrete():
    return [
        ("放樣", "位置、尺寸", "依圖說±5cm", "施工前", "目視量測", "每次", "放樣紀錄表", "重新放樣"),
        ("模板組立", "尺寸、垂直度", "高程±5mm", "組立後", "目視量測", "每次", "模板檢查表", "修正後重驗"),
        ("鋼筋組立", "號數、間距、保護層", "±10mm", "澆置前", "量測", "每次", "鋼筋檢查表", "補正後重驗"),
        ("混凝土坍度", "坍度值", "15±4 cm", "進場時", "坍度試驗", "每車", "坍度試驗記錄", "退料換批"),
        ("混凝土溫度", "入模溫度", "13°C～32°C", "進場時", "溫度計量測", "每車", "試驗記錄", "退料換批"),
        ("氯離子含量", "含量", "≦0.15 kg/m³", "進場時", "氯離子測定儀", "每50m³", "試驗報告", "退料換批"),
        ("澆置間隔", "澆置連續性", "≦45 min", "澆置中", "計時", "全程", "施工日誌", "停止澆置處理"),
        ("振動器", "振搗間距", "≦45cm，5～10秒", "澆置中", "目視", "全程", "施工日誌", "重新振搗"),
        ("養護", "養護時間", "≧7天（持續濕潤）", "拆模後", "目視", "每日", "施工日誌", "延長養護"),
        ("混凝土強度", "抗壓強度f'c", "依圖說規格", "養護後", "試體試驗", "每50m³", "試驗報告", "補強或拆除"),
    ]

def get_standard_rows_planting():
    return [
        ("植物進場", "種類、規格", "依契約圖說", "進場時", "目視量測", "每批", "進場驗收表", "退料換批"),
        ("土壤準備", "土壤質量", "依契約規範", "種植前", "目視", "全區", "整地紀錄", "換土改良"),
        ("植穴開挖", "深度、尺寸", "依圖說規格±5%", "種植前", "量測", "每株", "種植紀錄", "重新開挖"),
        ("種植", "位置、深度、間距", "依圖說規定", "種植時", "量測目視", "每株", "種植紀錄", "重新種植"),
        ("固定支撐", "支柱穩固", "依契約規格", "種植後", "目視拉力", "每株", "種植紀錄", "重新固定"),
        ("灌水養護", "澆水量、頻率", "依規範", "種植後", "目視", "定期", "養護日誌", "增加澆水"),
        ("成活率", "成活株數", "≧95%", "驗收前", "目視計數", "全株", "驗收記錄", "補植"),
    ]

def get_standard_rows_lighting():
    return [
        ("電管配設", "管徑、深度", "依圖說規格", "埋設前", "量測目視", "全程", "施工紀錄", "重新配管"),
        ("基座施作", "尺寸、高程", "±5mm", "施作後", "量測", "每座", "基座紀錄", "修正重驗"),
        ("燈具安裝", "型號、位置", "依圖說", "安裝後", "目視", "每具", "安裝紀錄", "更換或調整"),
        ("接地電阻", "接地阻抗", "≦10Ω", "完成後", "接地電阻計", "每迴路", "測試記錄", "重做接地"),
        ("絕緣電阻", "絕緣性能", "≧1MΩ", "完成後", "絕緣電阻計", "每迴路", "測試記錄", "修復後重測"),
        ("通電測試", "燈具亮度、功能", "全數正常亮燈", "通電後", "目視測試", "全數", "測試記錄", "更換燈具"),
    ]

# ═══════════════════════════════════════════════════════
# 自主檢查表底部固定格式
# ═══════════════════════════════════════════════════════

SELFCHECK_FOOTER = [
    "缺失複查結果：",
    "□已完成改善",
    "□未完成改善，填至「不合格管制總表」第○項進行追蹤改善",
    "",
    "複查日期：　　年　　月　　日　　複查人員職稱：　　　　　簽名：",
    "",
    "備註：",
    "1.檢查標準及實際檢查情形應具體明確或量化尺寸。",
    "2.合格「○」，不合格「╳」，不需檢查「／」。",
    "3.嚴重缺失未及時完成改善，應填具「不合格品管制總表」追蹤改善。",
    "4.本表由工地現場工程師或領班實地檢查後覈實記載簽認。",
    "",
    "現場工程師：＿＿＿＿＿　　工地負責人（工地主任）：＿＿＿＿＿",
]


# ═══════════════════════════════════════════════════════
# 主程式
# ═══════════════════════════════════════════════════════

def build_doc():
    doc = Document()

    # ── 全域頁面設定（A4）──
    for sec in doc.sections:
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin   = Cm(3)
        sec.right_margin  = Cm(2.5)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2)

    # ── 樣式初始化 ──
    normal = doc.styles['Normal']
    normal.font.name = DOCX_FONT_LATIN
    normal.font.size = Pt(DOCX_SZ_BODY)
    _nr = normal.element.get_or_add_rPr()
    _nf = _nr.get_or_add_rFonts()
    _nf.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    _nf.set(qn('w:ascii'),    DOCX_FONT_LATIN)
    _nf.set(qn('w:hAnsi'),    DOCX_FONT_LATIN)

    HEADING_COLORS = {1:(0,0,0), 2:(0,0,0), 3:(0,0,0)}
    for lvl, sz, bold_flag in [(1, DOCX_SZ_H1, True),(2, DOCX_SZ_H2, True),(3, DOCX_SZ_H3, True)]:
        s = doc.styles[f'Heading {lvl}']
        s.font.name = DOCX_FONT_LATIN
        s.font.size = Pt(sz)
        s.font.bold = bold_flag
        s.font.color.rgb = RGBColor(*HEADING_COLORS[lvl])
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after  = Pt(6)
        if lvl == 1:
            s.paragraph_format.page_break_before = True
        _hr = s.element.get_or_add_rPr()
        _hf = _hr.get_or_add_rFonts()
        _hf.set(qn('w:eastAsia'), DOCX_FONT_CJK)
        _hf.set(qn('w:ascii'),    DOCX_FONT_LATIN)
        _hf.set(qn('w:hAnsi'),    DOCX_FONT_LATIN)

    # ════════════════════════════════════
    # 封面
    # ════════════════════════════════════
    for _ in range(6): doc.add_paragraph()
    cover_para(doc, "整體品質計畫", DOCX_SZ_COVER_MAIN, bold=True, space_before=0, space_after=18)
    cover_para(doc, ENG_NAME, DOCX_SZ_COVER_PROJ, bold=True, space_before=0, space_after=24)
    cover_para(doc, f"主辦機關：{ENG_OWNER}", DOCX_SZ_COVER_INFO, space_before=6, space_after=6)
    cover_para(doc, f"設計單位：{ENG_DESIGN}", DOCX_SZ_COVER_INFO, space_before=0, space_after=6)
    cover_para(doc, f"監造單位：{ENG_SUPER}", DOCX_SZ_COVER_INFO, space_before=0, space_after=6)
    cover_para(doc, f"承攬廠商：{ENG_CONTR}", DOCX_SZ_COVER_INFO, space_before=0, space_after=6)
    for _ in range(3): doc.add_paragraph()
    cover_para(doc, "中華民國　　　年　　月", DOCX_SZ_COVER_INFO, space_before=0, space_after=0)
    insert_section_break(doc)

    # ════════════════════════════════════
    # 目錄
    # ════════════════════════════════════
    p_toc_title = doc.add_paragraph("目    錄")
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_toc_title.runs[0]
    set_run_font(run, DOCX_SZ_H2, bold=True)

    toc_entries = [
        {'level':1,'title':'第一章　計畫範圍','page':1},
        {'level':2,'title':'  一、依據','page':1},
        {'level':2,'title':'  二、工程概要','page':1},
        {'level':2,'title':'  三、適用對象','page':8},
        {'level':2,'title':'  四、名詞定義','page':9},
        {'level':1,'title':'第二章　管理責任及權責分工','page':11},
        {'level':2,'title':'  一、品管組織','page':11},
        {'level':2,'title':'  二、工作職掌','page':12},
        {'level':2,'title':'  三、人員輪值及請休假事項','page':14},
        {'level':2,'title':'  四、管理審查','page':14},
        {'level':2,'title':'  五、應用表單','page':15},
        {'level':1,'title':'第三章　施工要領','page':16},
        {'level':2,'title':'  一、分項施工計畫提送時程與管制','page':16},
        {'level':2,'title':'  二、施工要領訂定','page':17},
        {'level':1,'title':'第四章　品質管理標準','page':39},
        {'level':2,'title':'  一、品質管理標準訂定','page':39},
        {'level':1,'title':'第五章　材料及施工檢驗程序','page':82},
        {'level':2,'title':'  一、材料設備檢驗程序','page':82},
        {'level':2,'title':'  二、施工檢驗程序','page':84},
        {'level':2,'title':'  三、應用表單','page':112},
        {'level':1,'title':'第六章　設備功能運轉測試程序及標準','page':132},
        {'level':2,'title':'  一、設備功能運轉測試檢測程序','page':132},
        {'level':2,'title':'  二、設備功能運轉測試檢驗標準','page':135},
        {'level':2,'title':'  三、應用表單','page':136},
        {'level':1,'title':'第七章　自主檢查表','page':138},
        {'level':2,'title':'  一、自主檢查表訂定','page':138},
        {'level':2,'title':'  二、自主檢查表執行','page':139},
        {'level':2,'title':'  三、應用表單','page':140},
        {'level':1,'title':'第八章　不合格品之管制','page':169},
        {'level':2,'title':'  一、不合格材料及設備之管制','page':169},
        {'level':2,'title':'  二、施工不合格品質之管制','page':171},
        {'level':2,'title':'  三、應用表單','page':173},
        {'level':1,'title':'第九章　矯正與預防措施','page':177},
        {'level':2,'title':'  一、矯正措施','page':177},
        {'level':2,'title':'  二、預防措施','page':179},
        {'level':2,'title':'  三、應用表單','page':181},
        {'level':1,'title':'第十章　內部品質稽核','page':184},
        {'level':2,'title':'  一、品質稽核權責','page':184},
        {'level':2,'title':'  二、品質稽核範圍','page':185},
        {'level':2,'title':'  三、品質稽核頻率','page':185},
        {'level':2,'title':'  四、品質稽核流程','page':186},
        {'level':2,'title':'  五、應用表單','page':187},
        {'level':1,'title':'第十一章　文件記錄管理系統','page':192},
        {'level':2,'title':'  一、文件管理系統','page':192},
        {'level':2,'title':'  二、紀錄管理作業程序','page':195},
        {'level':2,'title':'  三、文件紀錄移轉及存檔','page':197},
    ]
    make_toc(doc, toc_entries)
    insert_section_break(doc)

    # ════════════════════════════════════
    # 第一章 計畫範圍
    # ════════════════════════════════════
    heading(doc, "第一章　計畫範圍", 1)

    body(doc, "前言", bold=True)
    body(doc, f"　　{ENG_CONTR}（以下簡稱本公司）承攬{ENG_NAME}，為達成契約所規定之工程品質目標，如特性、規格、成本、有效性、壽命週期等，依據「公共工程施工品質管理制度及作業要點」建立施工品質管制系統，設立品質管理組織，訂定整體品質計畫，據以推動施實。執行施工及材料設備之檢查、內部品質稽核、統計分析、矯正與預防措施，並對結果進行檢討及改善，過程留存文件紀錄，以達成全面提升工程品質之目標。")

    heading(doc, "一、依據", 2)
    body(doc, "　　本品質計畫撰寫之依據，係以工程契約（含施工規範及設計圖說）、技師法、建築法、建築師法、營造業法、公共工程專業技師簽證規則、職業安全衛生法、職業安全衛生設施規則、營造安全衛生設施標準、環境保護相關規定、加強公共工程勞工安全衛生管理作業要點、公共工程施工綱要規範、公共工程施工階段契約約定權責分工表（無委託專案管理）、本公司品質系統作業規定編擬。")

    heading(doc, "二、工程概要", 2)
    body(doc, f"　　本案基地位屬「縱貫公路桃園內壢間都市計畫」，區內公共設施用地已劃設逾40年遲未開闢，桃園市政府以公辦重劃方式取得仍有需要之公園用地之基地進行現況清整後歸還於民。")
    doc.add_paragraph()

    info_data = [
        ("工程名稱", ENG_NAME),
        ("主辦機關", ENG_OWNER),
        ("設計單位", "杜風工程服務股份有限公司"),
        ("監造單位", "杜風工程服務股份有限公司"),
        ("承攬廠商", ENG_CONTR),
        ("工程地點", ENG_SITE),
        ("契約工期", f"{ENG_DAYS}日曆天"),
        ("契約金額", ENG_AMOUNT),
    ]
    for label, val in info_data:
        body(doc, f"{label}：{val}")

    heading(doc, "三、工程主要施工項目及數量", 2)
    add_table_caption(doc, 1, 1, "工程主要施工項目及數量")
    make_table(doc,
        headers=["項次","項目及說明","單位","數量"],
        rows=MAIN_QUANTITIES,
        col_widths=[2.0, 10.0, 2.0, 2.5])

    heading(doc, "四、適用對象", 2)
    body(doc, "　　本品質計畫之適用對象為本公司（承攬廠商）辦理本工程之所有施工作業人員、協力廠商及相關人員。凡從事本工程任何施工、材料採購及品質管制相關作業之人員，均應遵循本品質計畫之規定辦理。")

    heading(doc, "五、名詞定義", 2)
    definitions = [
        ("本工程", ENG_NAME),
        ("主辦機關", f"{ENG_OWNER}（甲方）"),
        ("設計廠商", "杜風工程服務股份有限公司"),
        ("監造單位", "杜風工程服務股份有限公司（由桃園市政府地政局書面通知，由杜風工程服務股份有限公司派駐監造相關人員任之。）"),
        ("承攬廠商、本公司", f"{ENG_CONTR}（乙方）"),
        ("協力廠商", "為材料及施工分包廠商。"),
        ("工地主任", "係指受聘於本公司，擔任本工程之工地事務及施工管理之人員。"),
        ("專任工程人員", "係指受聘於本公司之技師或建築師，擔任本工程之施工技術指導及施工安全之人員。"),
        ("品管人員", "為本公司專責辦理品質計畫與執行查證工作，並且配合材料實驗室之各項品質檢驗，直接對本公司負責。"),
        ("現場工程師", "本公司派駐現場協助工地負責人執行各項工程進行之人員。"),
        ("程序", "執行一項工程所規定的方法。"),
        ("檢驗", "一種活動工程，例如量測、檢查、試驗或用樣板比測某一實體之一種或數種特性，將所得結果與規定要求相比較，以確定各個特性是否符合。"),
        ("自主檢驗", "由承攬廠商按特定的規則所做之檢驗。"),
        ("確認", "由檢查與提出客觀證據以證實某一特定預期用途之特別要求已被達成。"),
        ("品質", "一個實體的特性總和，此種特性具有滿足明訂與潛在的需求之能力。"),
        ("一般條款", "為用以規定承包商履行工程契約所應遵守之行為、履行之方式、業主與承包商間之責任與義務及業主、工程司、承包商之三者關係之規範，並為契約文件之一部分。"),
        ("施工規範", "說明有關執行本契約工作之施工規定。"),
        ("契約圖說", "說明有關執行本工程之施工及材料之尺寸、規格、注意事項。"),
        ("施工製造圖", "執行本契約圖說，另畫製施工詳細圖。"),
        ("檢驗程序", "施工、材料進場施作檢查流程及規定。"),
        ("施工缺失", "施工、材料施作時尺寸、規格不符契約及契約圖說之規定。"),
        ("內部品質稽核", "係本公司上級主管帶隊至各工地之品質稽核。"),
        ("矯正措施", "本工程之缺失改善之措施。"),
        ("檢驗停留點", "廠商應於品質計畫之材料及施工檢驗程序，明訂各項施工作業（含假設工程）及材料設備檢驗自主檢查之查驗點；自主檢查之查驗點應涵蓋監造單位明定之檢驗停留點。"),
        ("安衛查驗點", "各作業施工前，就施工過程設定安全衛生查驗點，並於施工中、驗收或使用前，分別執行查驗作業。"),
        ("管理標準", "依據契約、圖說、技術規範與特定條款，以及相關法規及標準等制定之標準，以進場材料及施工管理為依據。"),
    ]
    for i, (term, definition) in enumerate(definitions, 1):
        roman = ["(一)","(二)","(三)","(四)","(五)","(六)","(七)","(八)","(九)","(十)",
                 "(十一)","(十二)","(十三)","(十四)","(十五)","(十六)","(十七)","(十八)",
                 "(十九)","(二十)","(二十一)","(二十二)","(二十三)","(二十四)","(二十五)",
                 "(二十六)","(二十七)","(二十八)"][i-1] if i <= 28 else f"({i})"
        body(doc, f"{roman}{term}：{definition}")

    # ════════════════════════════════════
    # 第二章 管理責任及權責分工
    # ════════════════════════════════════
    heading(doc, "第二章　管理責任及權責分工", 1)
    heading(doc, "一、品管組織", 2)

    # DrawML 品管組織架構圖（不含主辦機關/監造單位，正確層次及線型）
    add_flowchart_placeholder(doc, 'org_chart', preset_quality_org_chart(
        label_tech    = f'主任技師\n（{ENG_TECH}）',
        label_foreman = f'工地主任\n（{ENG_FOREMAN}）',
        label_qc      = f'品管人員\n（{ENG_QA}）',
        label_safety  = '職業安全衛生\n管理人員',
        label_sub     = '各協力廠商',
        label_eng     = '現場工程師',
        label_admin   = '行政人員',
        font_cjk      = DOCX_FONT_CJK,
        font_latin    = DOCX_FONT_LATIN,
        font_pt       = DOCX_SZ_TABLE,
    ))
    add_figure_caption(doc, 2, 1, "品管組織架構圖")

    heading(doc, "二、工作職掌", 2)
    body(doc, "　　依據工程契約及「公共工程施工品質管理作業要點」之規定，指派經訓練合格且具工程實務經驗之人員從事工程品質管理。又依「營造業法」及「職業安全衛生管理辦法」之規定，指派本工程除專任工程人員外，配置工地主任一人、品管人員一人、職業安全衛生業務主管（營造業）一人；工作職掌如下表2-1：")
    add_table_caption(doc, 2, 1, "工作職掌表")
    make_table(doc,
        headers=["職稱","工作職掌及依據"],
        rows=[
            ("管理階層","1.統籌契約管理。\n2.分包商、供應商施工品質成效之稽查。"),
            ("專任工程人員\n（主任技師）","1.督查品管人員及現場施工人員，落實執行品質計畫，並填寫督察紀錄。\n2.依營造業法第35條規定，辦理相關工作，如督察按圖施工、解決施工技術問題，並於查驗文件簽名或蓋章。\n3.依本府工程施工查核小組設置要點規定於工程查核時，到場說明。"),
            ("工地負責人\n（工地主任）","(一)工地管理事項：工地範圍內之部署及配置；工人、材料、機具、設備、門禁及施工裝備之管理。\n(二)工程推動事項：開工之準備；材料、機具、設備檢（試）驗之申請、協調。\n(三)工地環境維護事項：施工場地排水設施之維護；工地圍籬之設置及維護。"),
            ("品管人員","1.依據工程契約、設計圖說、規範、相關技術法規及參考品質計畫製作綱要等，訂定品質計畫，據以推動實施。\n2.執行內部品質稽核，如稽核自主檢查表之檢查項目、檢查結果是否詳實記錄等。\n3.品管統計分析、矯正與預防措施之提出及追蹤改善。\n4.品質文件、紀錄之管理。"),
            ("職業安全衛生\n管理人員","1.執行工地安全衛生管理與環境保護事項。\n2.工地安全衛生緊急狀況之處置。\n3.編製安全衛生管理相關事宜。"),
            ("現場工程師","1.施工順序與施工界面之協調。\n2.各項工程材料、設備之檢驗與試驗。\n3.協辦施工計畫之編製。"),
            ("行政人員","1.文件繕打、文件管控、收發文、會計等一般行政工作。\n2.其他臨時交辦事項。"),
        ],
        col_widths=[3.5, 13.0])

    heading(doc, "三、人員輪值及請休假事項", 2)
    body(doc, "　　原則上，工地人員依照勞基法施行細則第23條之1規定辦理。遇有趕工必要時，工地人員會輪派留守。若工地人員請休假無法駐守工地時，會事先覓妥本公司合格人員代理。")

    heading(doc, "四、管理審查", 2)
    body(doc, "　　依每季審查品質管理系統以確保其持續適切、正確及有效性。審查內容包括前次審查紀錄執行情形、內部品質稽核、矯正措施、預防措施、外部驗證及工程評鑑結果。")

    heading(doc, "五、應用表單", 2)
    add_table_caption(doc, 2, 2, "公共工程施工中營造業專任工程人員督察紀錄表")
    make_table(doc,
        headers=["欄位","說明"],
        rows=[
            ("工程名稱", ENG_NAME),
            ("工程主辦機關", ENG_OWNER),
            ("承攬廠商", ENG_CONTR),
            ("填表日期", "　　年　　月　　日　　時"),
            ("工程進度概述", "預定進度（%）：\n實際進度（%）："),
            ("督察按圖施工\n（營造業法第35條第3款）", "督察項目包含：放樣工程、假設工程（含施工架）、基礎工程、模板工程、混凝土工程、鋼筋工程、景觀設施工程、植栽工程、照明工程等"),
            ("督察簽章", "【專任工程人員：主任技師】\n工地主任（工地負責人）："),
        ],
        col_widths=[4.0, 12.5])

    # ════════════════════════════════════
    # 第三章 施工要領
    # ════════════════════════════════════
    heading(doc, "第三章　施工要領", 1)
    heading(doc, "一、分項施工計畫提送時程與管制", 2)

    add_table_caption(doc, 3, 1, "各分項計畫管制表")
    subplan_rows = [
        ("1","步道鋪面工程（含底結構層及路緣石等，需提送排磚計畫）","施工前",""),
        ("2","籃球場工程","施工前",""),
        ("3","遊戲場彈性地坪工程","施工前",""),
        ("4","遊具設備安裝工程","施工前",""),
        ("5","塑木花架分項工程（含施工圖等）","施工前",""),
        ("6","公園廁所工程（含施工圖等）","施工前",""),
        ("7","公園排水工程","施工前",""),
        ("8","植栽工程","施工前",""),
        ("9","景觀照明工程","施工前",""),
        ("10","停車場工程及圍籬工程（含施工圖等）","施工前",""),
    ]
    make_table(doc,
        headers=["項次","提送項目","預定送審日期","備註"],
        rows=subplan_rows,
        col_widths=[1.5, 10.0, 3.0, 2.0])

    heading(doc, "二、分項施工計畫綱要", 2)
    body(doc, "　　依工程需求，提送分項施工計畫，其內容包含：")
    for item in [
        "(一)工項概述：包括工項概要、內容及數量等。",
        "(二)人員組織：包括施工人員組織、協力廠商人力安排。",
        "(三)施工方法與步驟：包括施作順序及介面整合。",
        "(四)施工機具：包括使用機具及供電設備。",
        "(五)使用材料：包括材料規格、材料數量、儲放及搬運。",
        "(六)預定作業進度：包括施工規劃進度檢討、預定作業進度表。",
        "(七)分項品質計畫：包括施工要領、品質管理標準、材料及施工檢驗程序及自主檢查表。",
        "(八)分項作業安全衛生管理與設施設置措施。",
    ]:
        body(doc, item, indent=1.0)

    heading(doc, "三、施工要領訂定", 2)
    add_table_caption(doc, 3, 2, "各分項工程施工要領一覽表")
    item_list_rows = [(str(i+1), name, f"表三-{i+3}") for i, (name, _) in enumerate(WORK_ITEMS)]
    make_table(doc,
        headers=["項次","工項名稱","編號"],
        rows=item_list_rows,
        col_widths=[1.5, 12.0, 3.0])

    # 各工項施工要領表
    method_headers = ["施工步驟","使用材料","施工機具","注意事項"]
    method_col = [3.5, 3.5, 3.5, 6.0]

    work_method_data = {
        "交通臨時安全措施": [
            ("1.施工前準備","安全圍籬、警告標誌、指示牌、閃光紅燈等","吊卡車","依施工圖說。支架設立位置依指示或是施工圖說。"),
            ("2.確認設置位置","噴漆","","依圖說規定確認點位位置及尺寸。"),
            ("3.安裝完成","","","完成面完整、穩固；路面整潔。"),
        ],
        "土方管制工程": [
            ("1.施工前準備","","","依施工圖說。"),
            ("2.確認設置位置","噴漆","","依圖說規定確認土方位置及尺寸；雜草、樹木應先清除。"),
            ("3.開挖","","挖土機、運土車","自上而下依序開挖；應與廢土分開。"),
            ("4.完成運離","","","路面整潔。"),
        ],
        "測量工程": [
            ("1.施工前準備","","","擬定測量計劃；選擇適當的測量方式及工具。"),
            ("2.控制樁測量","鋼釘","經緯儀、水準儀、皮尺","與其他施工單位之導線點確認一致。"),
            ("3.縱橫斷面測量","木樁、鋼釘","經緯儀、水準儀、皮尺","測定主要結構體橫斷面線。"),
            ("4.結構物放樣","木樁","經緯儀","位置、斜率及高程須符合圖說規定；放樣誤差須符合圖說要求。"),
        ],
        "鋼筋工程": [
            ("1.施工前準備","","鋼捲尺","依據設計圖，施工說明核算數量繪製結構體每根鋼筋剪裁長度。"),
            ("2.材料進場","鋼筋","裝卸用貨車、吊車","放射線檢驗報告、出場報告、取樣送驗；材料須墊高儲放，加蓋帆布。"),
            ("3.鋼筋加工","鋼筋","鋼筋加工機具","裁切長度；排列層次；彎勾和錨定長度。"),
            ("4.鋼筋組立","鋼筋、鐵絲、墊塊","吊車","鐵絲綁紮以防止鋼筋位移；保護層墊置妥善。"),
            ("5.組立後檢查","","皮尺","號數及支數檢查；間距與保護層須符合圖說規範。"),
        ],
        "模板工程": [
            ("1.施工前準備","","鋼捲尺","決定工法，模板應力分析及繪製施工圖送審核可據以施工。"),
            ("2.模板進場","","裝卸用貨車、吊車","注意板面結合間隙；角材尺寸、平直度及間距要符合規定。"),
            ("3.放樣","","經緯儀、水準儀、皮尺","依施工放樣圖說；放樣尺寸高程及坡度須符合圖說規範。"),
            ("4.模板及支撐組立","鐵絲、墊塊","吊車","構造物斷面尺寸須符合規定；支撐須為穩固狀態；須有脫模劑。"),
            ("5.拆模","","L型撬棒","各部份模板依規定期限拆除；清理模板及塗油。"),
        ],
        "混凝土工程": [
            ("1.施工前準備","","","模板尺寸、高程、支撐及鋼筋組立須檢查；保護層應符合圖說規定。"),
            ("2.拌合廠生產混凝土","","拌合設備","檢視配比設計資料；供應量及運輸路況之規劃，避免供料中斷。"),
            ("3.到場材料檢查","混凝土","預拌車、坍度測定儀、氯離子測定儀","應進行坍度（15±4cm）、氯離子（≦0.15kg/m³）試驗；抗壓試體須製作。"),
            ("4.混凝土澆置","混凝土","預拌車、洩槽、壓送車、振動棒","澆置順序須妥善規劃；澆置高度超過1.5m須使用漏斗；振動棒間隔45cm振搗5～10秒。"),
            ("5.養護","","","混凝土表面須保持濕潤；視水泥強度調整養護時間（≧7天）。"),
            ("6.拆模後修飾","水泥砂漿","砂磨機","鐵線及繫結器須修剪或拆除；保護層及蜂巢修補須用原配比修飾。"),
        ],
        "瀝青混凝土鋪面工程": [
            ("1.施工前準備","","","依據設計圖，施工說明核算數量；材料設備進場規格符合送審資料。"),
            ("2.新舊底層相交處理","壓條","皮槌","新舊底層相交處以壓條分割；完成面須平整乾淨。"),
            ("3.面層施作","彩繪止滑塗料","高壓水槍、養生膠帶、噴槍、鏝刀","表面使用高壓水清洗；底劑以0.05L/m²均勻噴灑；第一層面層噴塗厚度約1mm。"),
            ("4.標線","熱處理聚酯標線材","熱拌爐、標線車","路面溫度10～40°C；噴出材料溫度180～200°C；標線寬度、厚度應符合圖說規定。"),
            ("5.完成","","","表面平整無積水；表面清潔；設施穩固。"),
        ],
        "預鑄路緣石工程": [
            ("1.施工前準備","","鋼捲尺","依據設計圖，施工說明核算數量；基準點與高程點確認。"),
            ("2.底層基礎施作","","平路機、挖土機、振動輾壓機","路面雜物須清除乾淨；路面須整平壓實。"),
            ("3.緣石施作","緣石、收邊材","拖板車、木槌、橡皮榔頭","以木槌或橡皮榔頭輕擊表面使磚緊密接合；剛鋪設完成應設置圍籬防止人車進入。"),
            ("4.完成","","","施作完成後，應立即清理鋪面表面；完成面正直無歪斜、穩固無鬆動。"),
        ],
        "鋪面磚（高壓）工程": [
            ("1.施工前準備","","鋼捲尺","依據設計圖核算數量；材料設備確認無受損。"),
            ("2.底層整平","級配粒料","平路機、振動輾壓機","底層整平壓實；符合圖說高程。"),
            ("3.磚鋪設","高壓混凝土磚","橡皮榔頭","排列方向依圖說；磚縫均勻；高差≦2mm。"),
            ("4.填縫","填縫材","","填縫密實；表面清潔。"),
            ("5.完成","","","平整度符合規範；完成面整潔。"),
        ],
        "磨石子工程": [
            ("1.施工前準備","","","依圖說確認配色及配比。"),
            ("2.底層施作","水泥砂漿","","底層厚度及高程符合圖說。"),
            ("3.磨石子澆置","水泥、石粒、顏料","鏝刀","配比正確；澆置厚度符合規定；表面分格條安設確實。"),
            ("4.研磨","","磨石機","初磨、中磨、細磨、打蠟依序完成；無崩角缺陷。"),
            ("5.完成","","","表面光滑、色彩均勻；分格縫整齊。"),
        ],
        "抿石子工程": [
            ("1.施工前準備","","","依圖說確認石粒顏色、粒徑及配比。"),
            ("2.底層施作","水泥砂漿","","底層厚度及高程符合圖說。"),
            ("3.抿石子施作","水泥、石粒","鏝刀、海綿","配比正確；石粒均勻分布；厚度符合規定。"),
            ("4.洗面","","高壓水槍","洗面時機適當；石粒外露均勻；無掉落。"),
            ("5.完成","","","表面均勻、石粒顯露均勻；無脫落現象。"),
        ],
        "彈性無縫鋪面工程": [
            ("1.施工前準備","","","確認底層平整度及含水率。"),
            ("2.底層整平","","","底層須整平，凹陷處填補修平。"),
            ("3.彈性材料鋪設","彈性材料","鏝刀","材料配比正確；鋪設均勻；厚度t=4.5～9cm依圖說。"),
            ("4.表面層施作","面層材料","","表面顏色均勻；厚度符合規定。"),
            ("5.完成","","","表面平整無積水；顏色均勻無色差。"),
        ],
        "地坪工程": [
            ("1.施工前準備","","","確認底層平整度及基準高程。"),
            ("2.底層整平","水泥砂漿","","底層厚度、高程符合圖說。"),
            ("3.面層施作","混凝土或砂漿","鏝刀","表面高程、坡度符合圖說。"),
            ("4.養護","","","保濕養護≧7天。"),
            ("5.完成","","","表面平整，高差≦3mm/2m；坡度符合排水要求。"),
        ],
        "磁磚工程": [
            ("1.施工前準備","","","依圖說確認磁磚規格、顏色。"),
            ("2.底層處理","水泥砂漿","","底層平整無裂縫；塗佈界面劑。"),
            ("3.磁磚鋪貼","磁磚、水泥砂漿","橡皮槌","磁磚縫隙均勻；敲擊無空鼓聲。"),
            ("4.填縫","填縫劑","","填縫密實；表面清潔。"),
            ("5.完成","","","無空鼓、無崩角；縫線整齊；表面清潔。"),
        ],
        "木作（塑木）工程": [
            ("1.施工前準備","","鋼捲尺","依據設計圖、型錄，施工說明核算數量。"),
            ("2.基礎施作","鋼筋、混凝土、螺栓","運送車","基礎面雜物清除；確認結構尺寸±5%。"),
            ("3.塑木構件安裝","塑木構件、鋼構件","電動螺絲起子、吊車","注意設置方向、高度；螺栓鎖固確實。"),
            ("4.完成面處理","","","無歪斜鬆動；完成面清潔。"),
        ],
        "公園廁所工程": [
            ("1.施工前準備","","","依圖說確認廁所型式及尺寸。"),
            ("2.基礎施作","混凝土、鋼筋","","基礎高程、尺寸符合圖說。"),
            ("3.牆體結構施作","預鑄構件或RC","吊車","結構垂直度±5mm；尺寸符合圖說。"),
            ("4.設備安裝","衛生設備、水電","","管線位置正確；設備安裝穩固。"),
            ("5.完成","","","功能正常；外觀整潔。"),
        ],
        "兒童遊戲場設備工程": [
            ("1.施工前準備","","鋼捲尺","依據設計圖、型錄，施工說明核算數量；材料設備確認無受損。"),
            ("2.基礎施作","鋼筋、混凝土、螺栓","運送車","基礎面雜物清除乾淨且整平壓實；確認結構尺寸±5%；預埋件穩固。"),
            ("3.主結構施作","兒童遊戲場設備","貨車、電動螺絲起子","基礎深度與範圍符合圖說；注意設置方向、高度與垂直度；完成面正直無歪斜、穩固無鬆動。"),
            ("4.完成","","","立即清理鋪面表面；完成面乾淨整潔；辦理遊戲場檢驗。"),
        ],
        "鍍鋅鋼網圍籬工程": [
            ("1.施工前準備","鍍鋅鋼網圍籬","吊卡車","形式、材料應依圖說；支架設立依施工圖說。"),
            ("2.基座確認","","水平儀、噴漆","依圖說規定確認點位位置及尺寸。"),
            ("3.組立完成","","","完成後完整、穩固。"),
        ],
        "公園排水工程": [
            ("1.施工前準備","","","依圖說確認管材規格及管徑。"),
            ("2.管溝開挖","","挖土機","開挖深度、寬度符合圖說；避免破壞鄰近管線。"),
            ("3.管線埋設","HDPE透水管、PVC管","","管線坡度符合圖說；接頭密合。"),
            ("4.陰井施作","預製混凝土陰井","吊車","陰井位置、高程符合圖說；底板平整。"),
            ("5.回填","","振動夯","分層回填壓實；壓實度≧95%。"),
            ("6.完成","","","排水功能測試正常；地面復原整潔。"),
        ],
        "植栽工程": [
            ("1.施工前準備","","","樹種須確認並檢視出廠報告；土壤肥料等材料要求須符合圖說規範。"),
            ("2.現地整理","肥料、鐵件","挖土機","排水坡度整理須符合圖說；現地依圖說規範整平壓實。"),
            ("3.植物搬運進場","契約內規定之植物","裝卸用貨車","運輸保護措施；裝卸保護措施。"),
            ("4.植物種植","契約內規定之植物","裝卸用貨車","種植的位置、尺寸、深度及密度需符合契約圖說規定。"),
            ("5.養護與補植","肥料、防病蟲害之藥品","灑水設備、割草機","澆水養護並施追肥；病蟲害防治；清除雜草；修剪。"),
        ],
        "景觀照明工程": [
            ("1.施工前準備","燈具設備、管材","","依據設計圖，施工說明核算數量；材料設備確認無受損。"),
            ("2.放樣","","噴漆","依圖說規定確認各燈具位置。"),
            ("3.開挖","","挖土機","確認開挖深度符合規範；溝底不得有大土塊。"),
            ("4.基座施作","基栓","搬運車","確認施作完成高程；確認接地線排涉及檢測。"),
            ("5.配管","PVC電管材","","管道連線前要進行外觀檢查；各管連線塗膠之前管口內外管壁均應清潔乾燥；完成配管後埋設警示帶。"),
            ("6.管溝回填","","挖土機","回填高度符合圖說規範；如溝槽土質較差需外運細土回填。"),
            ("7.燈具安裝","燈具設備","","安裝完成後確認穩固；通電測試，是否有燈具故障未亮。"),
        ],
        "電力配線工程": [
            ("1.施工前準備","","","依圖說確認配線規格及路徑。"),
            ("2.電管配設","PVC電管","","電管固定確實；彎曲半徑符合規定。"),
            ("3.電線穿線","電線電纜","穿線器","電線截面積符合設計；拉力適當不得損傷絕緣層。"),
            ("4.接線測試","","絕緣電阻計","絕緣電阻≧1MΩ；接線正確。"),
            ("5.完成","","","功能測試正常；標示完整。"),
        ],
        "器具設備工程": [
            ("1.施工前準備","","","依圖說確認設備型號及安裝位置。"),
            ("2.設備進場","器具設備","","規格型號確認；外觀完整無損傷。"),
            ("3.設備安裝","","吊車、電動工具","安裝位置符合圖說；固定穩固。"),
            ("4.功能測試","","","功能正常；無異常聲響。"),
            ("5.完成","","","完成測試記錄；保護措施完善。"),
        ],
        "假設工程": [
            ("1.施工前準備","","","依圖說及施工計畫確認假設工程項目。"),
            ("2.圍籬設置","安全圍籬、警告標誌","","圍籬穩固；標誌清晰可見。"),
            ("3.臨時設施","","","符合安全規定；與施工動線配合。"),
            ("4.維護","","","定期檢查維護；確保安全。"),
            ("5.拆除","","","工程完成後依序拆除；清理復原。"),
        ],
        "汛期工地防災準備": [
            ("1.汛前檢查","","","汛前完成工地排水設施檢查；確認排水暢通。"),
            ("2.防災設施佈置","沙包、抽水機、防水布","","沙包堆疊穩固；抽水機備妥可用。"),
            ("3.人員告知","","","告知全體工地人員防災應變程序；緊急聯絡名單備妥。"),
            ("4.汛期監控","","","定時查看天氣預報；雨量超標立即啟動應變。"),
            ("5.災後復原","","","災後立即評估損害；依序修復復工。"),
        ],
    }

    for i, (name, desc) in enumerate(WORK_ITEMS):
        add_table_caption(doc, 3, i+3, f"{name}施工要領")
        rows = work_method_data.get(name, [
            ("1.施工前準備","","","依施工圖說及規範辦理。"),
            ("2.施工中","","","依規範及圖說要求施作。"),
            ("3.完成","","","完成面符合設計要求。"),
        ])
        make_table(doc, headers=method_headers, rows=rows, col_widths=method_col)

    # ════════════════════════════════════
    # 第四章 品質管理標準
    # ════════════════════════════════════
    heading(doc, "第四章　品質管理標準", 1)
    heading(doc, "一、品質管理標準訂定", 2)
    body(doc, "　　本工程各施工工項依契約、圖說及施工規範，訂定品質管理標準如下各表所示。每張表底部標注：＊為監造停留點　※為自主檢查停留點")

    std_headers = ["施工流程","管理項目","檢查標準","檢查時機","檢查方法","檢查頻率","管理紀錄","不符合標準之處置方法"]
    std_cols    = [2.5, 2.5, 3.0, 2.0, 2.5, 2.0, 2.5, 2.5]

    add_table_caption(doc, 4, 1, "混凝土工程品質管理標準表")
    make_table(doc, headers=std_headers, rows=get_standard_rows_concrete(), col_widths=std_cols)
    body(doc, "＊為監造停留點　※為自主檢查停留點")

    add_table_caption(doc, 4, 2, "植栽工程品質管理標準表")
    make_table(doc, headers=std_headers, rows=get_standard_rows_planting(), col_widths=std_cols)
    body(doc, "＊為監造停留點　※為自主檢查停留點")

    add_table_caption(doc, 4, 3, "景觀照明工程品質管理標準表")
    make_table(doc, headers=std_headers, rows=get_standard_rows_lighting(), col_widths=std_cols)
    body(doc, "＊為監造停留點　※為自主檢查停留點")

    body(doc, "（其餘各工項品質管理標準表依各分項施工計畫另訂之，本整體品質計畫附主要工項代表性標準表供參考。）")

    # ════════════════════════════════════
    # 第五章 材料及施工檢驗程序
    # ════════════════════════════════════
    heading(doc, "第五章　材料及施工檢驗程序", 1)
    heading(doc, "一、材料設備檢驗程序", 2)
    body(doc, "　　本公司應於各項材料設備進場前，依契約規定辦理材料設備選定送審，並於監造單位核可後方可進場使用。材料設備選定送審流程如圖5-1所示：")
    add_flowchart_placeholder(doc, 'fig_5_1', preset_material_approval())
    add_figure_caption(doc, 5, 1, "材料/設備選定送審流程圖")

    heading(doc, "二、施工檢驗程序", 2)
    body(doc, "　　各項施工作業應依下列施工檢驗流程進行自主檢查，並於達到檢驗停留點時，通知監造單位檢驗確認後方可繼續施工。")
    add_flowchart_placeholder(doc, 'fig_5_2', preset_work_inspection())
    add_figure_caption(doc, 5, 2, "施工檢驗流程圖")

    heading(doc, "三、各工項施工作業流程及檢驗圖", 2)

    three_col_items_concrete = [
        {'flow': '模板組立', 'check': '尺寸高程±5mm、垂直度', 'doc': '模板檢查表', 'mark': '※', 'type': 'process'},
        {'flow': '鋼筋組立', 'check': '間距、號數、保護層±10mm', 'doc': '鋼筋檢查表', 'mark': '※', 'type': 'process'},
        {'flow': '預拌混凝土進場', 'check': '坍度15±4cm、溫度13～32°C', 'doc': '坍度試驗記錄', 'mark': '＊', 'type': 'process'},
        {'flow': '氯離子檢測', 'check': '≦0.15 kg/m³', 'doc': '氯離子試驗記錄', 'mark': '＊', 'type': 'process'},
        {'flow': '澆置混凝土', 'check': '振搗間距≦45cm，5～10秒', 'doc': '施工日誌', 'mark': '※', 'type': 'process'},
        {'flow': '養護', 'check': '持續濕潤≧7天', 'doc': '養護紀錄', 'type': 'process'},
        {'flow': '試體試驗', 'check': '抗壓強度符合圖說', 'doc': '試驗報告', 'mark': '＊', 'type': 'process'},
    ]
    add_three_col_flow(doc, '圖5-3 混凝土工程施工作業流程及檢驗圖', three_col_items_concrete,
                       marker_prefix='conc', font_cjk=DOCX_FONT_CJK, font_latin=DOCX_FONT_LATIN, font_pt=DOCX_SZ_TABLE)

    three_col_items_plant = [
        {'flow': '植物進場驗收', 'check': '種類規格符合圖說、健康無病蟲害', 'doc': '進場驗收表', 'mark': '＊', 'type': 'process'},
        {'flow': '現地整理', 'check': '排水坡度、整平壓實', 'doc': '整地紀錄', 'type': 'process'},
        {'flow': '植穴開挖', 'check': '深度尺寸±5%', 'doc': '種植紀錄', 'type': 'process'},
        {'flow': '種植', 'check': '位置間距深度符合圖說', 'doc': '種植紀錄', 'mark': '※', 'type': 'process'},
        {'flow': '固定支撐', 'check': '支柱穩固符合規格', 'doc': '種植紀錄', 'type': 'process'},
        {'flow': '澆水養護', 'check': '依規範頻率澆水施肥', 'doc': '養護日誌', 'type': 'process'},
        {'flow': '成活率確認', 'check': '成活率≧95%', 'doc': '驗收記錄', 'mark': '＊', 'type': 'process'},
    ]
    add_three_col_flow(doc, '圖5-4 植栽工程施工作業流程及檢驗圖', three_col_items_plant,
                       marker_prefix='plant', font_cjk=DOCX_FONT_CJK, font_latin=DOCX_FONT_LATIN, font_pt=DOCX_SZ_TABLE)

    three_col_items_light = [
        {'flow': '燈具設備進場驗收', 'check': '型號規格符合送審，外觀完整', 'doc': '進場驗收表', 'mark': '＊', 'type': 'process'},
        {'flow': '放樣確認', 'check': '燈具位置符合圖說', 'doc': '放樣紀錄', 'type': 'process'},
        {'flow': '電管配設及基座施作', 'check': '管徑、深度、基座高程符合圖說', 'doc': '施工紀錄', 'mark': '※', 'type': 'process'},
        {'flow': '燈具安裝', 'check': '安裝穩固、方向正確', 'doc': '安裝紀錄', 'type': 'process'},
        {'flow': '絕緣電阻測試', 'check': '≧1MΩ', 'doc': '測試記錄', 'mark': '＊', 'type': 'process'},
        {'flow': '接地電阻測試', 'check': '≦10Ω', 'doc': '測試記錄', 'mark': '＊', 'type': 'process'},
        {'flow': '通電功能測試', 'check': '全數燈具正常亮燈', 'doc': '測試記錄', 'mark': '＊', 'type': 'process'},
    ]
    add_three_col_flow(doc, '圖5-5 景觀照明工程施工作業流程及檢驗圖', three_col_items_light,
                       marker_prefix='light', font_cjk=DOCX_FONT_CJK, font_latin=DOCX_FONT_LATIN, font_pt=DOCX_SZ_TABLE)

    heading(doc, "四、應用表單", 2)
    add_table_caption(doc, 5, 1, "材料設備送審管制總表")
    make_table(doc,
        headers=["項次","材料/設備名稱","規格","送審日期","核定日期","備註"],
        rows=[
            ("1","高壓混凝土磚","依圖說規格","施工前","",""),
            ("2","預鑄路緣石","依圖說規格","施工前","",""),
            ("3","塑木構件","依圖說規格","施工前","",""),
            ("4","景觀燈具","依圖說型號","施工前","",""),
            ("5","兒童遊具設備","依圖說型號","施工前","",""),
            ("6","植栽材料","依圖說規格、CNS標準","施工前","",""),
            ("7","彈性無縫鋪面材料","依圖說規格","施工前","",""),
            ("8","HDPE透水網管","依圖說管徑","施工前","",""),
        ],
        col_widths=[1.5, 4.0, 4.0, 2.5, 2.5, 2.0])

    add_table_caption(doc, 5, 2, "材料設備檢（試）驗管制總表")
    make_table(doc,
        headers=["項次","試驗項目","試驗標準","試驗頻率","試驗機構","試驗結果"],
        rows=[
            ("1","混凝土抗壓強度試驗","fc'≧依設計","每50m³","公正第三機構",""),
            ("2","混凝土坍度試驗","15±4 cm","每車","現場",""),
            ("3","混凝土氯離子含量試驗","≦0.15 kg/m³","每50m³","公正第三機構",""),
            ("4","鋼筋拉力試驗","依CNS標準","每批","公正第三機構",""),
            ("5","瀝青混凝土芯樣試驗","依規範","每工作天","公正第三機構",""),
            ("6","填方壓實度試驗","≧95%","依規範頻率","公正第三機構",""),
            ("7","植栽健康狀況確認","依圖說規格","每批","現場",""),
        ],
        col_widths=[1.5, 3.5, 3.5, 2.5, 3.0, 2.5])

    # ════════════════════════════════════
    # 第六章 設備功能運轉測試
    # ════════════════════════════════════
    heading(doc, "第六章　設備功能運轉測試程序及標準", 1)
    heading(doc, "一、設備功能運轉測試檢測程序", 2)
    body(doc, "　　本工程景觀照明設備及公園排水設備等，應依下列程序辦理設備功能運轉測試：")
    add_flowchart_placeholder(doc, 'fig_6_1', preset_equipment_test())
    add_figure_caption(doc, 6, 1, "設備功能運轉檢測流程圖")

    body(doc, "設備功能運轉測試分為以下三個階段：")
    body(doc, "（一）單機測試：各設備單獨通電，確認個別設備功能正常。")
    body(doc, "（二）系統測試：相關設備連接後，確認系統功能正常運作。")
    body(doc, "（三）整體功能測試：全部設備連接完成後，確認整體功能運作正常，並模擬實際使用情境。")

    heading(doc, "二、設備功能運轉測試檢驗標準", 2)
    add_table_caption(doc, 6, 1, "機電設備功能運轉檢測標準一覽表")
    make_table(doc,
        headers=["設備名稱","測試項目","測試標準","測試頻率","記錄表單"],
        rows=[
            ("景觀照明燈具","通電亮燈測試","全數正常亮燈","安裝完成後全數","設備功能測試紀錄表"),
            ("景觀照明燈具","絕緣電阻測試","≧1MΩ","每迴路","絕緣電阻測試記錄"),
            ("景觀照明燈具","接地電阻測試","≦10Ω","每迴路","接地電阻測試記錄"),
            ("景觀照明燈具","光度測試","依設計規格","抽驗","光度測試記錄"),
            ("公園排水系統","排水功能測試","排水順暢無積水","完成後","排水功能測試紀錄"),
            ("PE儲水槽","容量及防水測試","依圖說規格","每組","防水測試記錄"),
        ],
        col_widths=[3.0, 3.5, 3.5, 2.5, 4.0])

    heading(doc, "三、應用表單", 2)
    add_table_caption(doc, 6, 2, "設備功能運轉測試記錄表")
    make_table(doc,
        headers=["設備名稱","規格型號","測試日期","測試人員","測試結果","備註"],
        rows=[
            ("景觀照明燈具（A區）","依設計圖說","","","",""),
            ("景觀照明燈具（B-1區）","依設計圖說","","","",""),
            ("景觀照明燈具（B-2/B-3區）","依設計圖說","","","",""),
            ("公園排水系統","依設計圖說","","","",""),
        ],
        col_widths=[3.5, 3.5, 2.5, 2.5, 3.0, 1.5])

    # ════════════════════════════════════
    # 第七章 自主檢查表
    # ════════════════════════════════════
    heading(doc, "第七章　自主檢查表", 1)
    heading(doc, "一、自主檢查表訂定", 2)
    body(doc, "　　本公司依據工程契約、設計圖說及施工規範，針對本工程各主要施工工項，訂定自主檢查表如下。各自主檢查表應由工地現場工程師或領班於實際施工時逐項確認並記錄，不得事後補填。")
    body(doc, "自主檢查停留點（※）：本公司現場人員實施自主查驗，確認合格後方可繼續施工。")
    body(doc, "監造停留點（＊）：需通知監造單位到場確認，監造人員簽認後方可繼續施工。")

    heading(doc, "二、自主檢查表執行", 2)
    body(doc, "　　各工程自主檢查表由現場工程師依施工進度，於每工項施工完成後逐項填寫，確認各項目均合格後簽名，如發現缺失立即填寫不合格品管制表並追蹤改善。")

    heading(doc, "三、應用表單", 2)

    # 主要自主檢查表
    selfcheck_items = [
        ("整地拆除工程", 7, 1, [
            ("拆除範圍確認","依圖說範圍","","",""),
            ("廢棄物分類","依廢棄物清除規定","","",""),
            ("地坪打除厚度","依圖說規格","","",""),
            ("土方開挖深度","±5cm","","",""),
            ("廢方處理","依規定場所運棄","","",""),
        ]),
        ("混凝土工程", 7, 2, [
            ("模板組立尺寸","±5mm","","",""),
            ("鋼筋號數及間距","依圖說規定","","",""),
            ("鋼筋保護層","與水土接觸75mm；不接觸50mm","","",""),
            ("混凝土坍度","15±4 cm","","",""),
            ("混凝土溫度","13°C～32°C","","",""),
            ("氯離子含量","≦0.15 kg/m³","","",""),
            ("振搗間距","≦45cm，5～10秒","","",""),
            ("養護（持續濕潤）","≧7天","","",""),
            ("試體製作","每50m³至少1組","","",""),
        ]),
        ("植栽工程", 7, 3, [
            ("植物種類規格","依契約圖說","","",""),
            ("植物健康狀態","無病蟲害","","",""),
            ("植穴深度","依圖說規定","","",""),
            ("種植位置","依圖說位置±10cm","","",""),
            ("覆土深度","依圖說規定","","",""),
            ("支撐樁設置","穩固符合規格","","",""),
            ("澆水養護","定期澆水依規範","","",""),
        ]),
        ("景觀照明工程", 7, 4, [
            ("電管配設深度","依圖說規格","","",""),
            ("基座高程","±5mm","","",""),
            ("燈具型號規格","依設計圖說型號","","",""),
            ("燈具安裝穩固","無搖晃鬆動","","",""),
            ("絕緣電阻","≧1MΩ","","",""),
            ("接地電阻","≦10Ω","","",""),
            ("通電亮燈測試","全數正常","","",""),
        ]),
        ("公園排水工程", 7, 5, [
            ("管溝開挖深度","依圖說規格","","",""),
            ("管線坡度","依圖說規定","","",""),
            ("管接頭密合","無滲漏","","",""),
            ("陰井位置高程","±5mm","","",""),
            ("回填壓實度","≧95%","","",""),
            ("排水功能測試","排水順暢無積水","","",""),
        ]),
    ]

    for item_name, chap, seq, check_items in selfcheck_items:
        add_table_caption(doc, chap, seq, f"{item_name}施工自主檢查表")
        make_table(doc,
            headers=["檢查項目","設計圖說、規範之檢查標準","實際檢查情形","檢查結果","備註"],
            rows=check_items,
            col_widths=[4.0, 5.5, 3.5, 2.0, 1.5])
        for line in SELFCHECK_FOOTER:
            p = doc.add_paragraph(line)
            run = p.runs[0] if p.runs else p.add_run(line)
            set_run_font(run if p.runs else p.runs[0], DOCX_SZ_NOTE)

    # ════════════════════════════════════
    # 第八章 不合格品之管制
    # ════════════════════════════════════
    heading(doc, "第八章　不合格品之管制", 1)
    heading(doc, "一、不合格材料及設備之管制", 2)
    body(doc, "　　材料設備進場後，如發現不符合契約規定之不合格品，應立即標示隔離，不得使用於工程上，並依下列流程處理：")
    add_flowchart_placeholder(doc, 'fig_8_1', preset_nc_material())
    add_figure_caption(doc, 8, 1, "材料自主檢查不合格管制作業流程")

    heading(doc, "二、施工不合格品質之管制", 2)
    body(doc, "　　施工過程中發現不符合品質要求之缺失，應依下列流程處置：")
    add_flowchart_placeholder(doc, 'fig_8_2', preset_nc_work())
    add_figure_caption(doc, 8, 2, "施工自主檢查不合格管制作業流程")

    heading(doc, "三、應用表單", 2)
    add_table_caption(doc, 8, 1, "不合格管制一覽表")
    make_table(doc,
        headers=["項次","缺失項目","發現日期","處理方式","改善完成日期","複查結果","承辦人"],
        rows=[("","","","","","","")]*3,
        col_widths=[1.5, 3.5, 2.0, 3.0, 2.5, 2.5, 1.5])

    add_table_caption(doc, 8, 2, "不合格品改正紀錄表")
    make_table(doc,
        headers=["工程名稱","缺失編號","缺失描述","改正措施","完成日期","複查結果"],
        rows=[(ENG_NAME,"","","","","")],
        col_widths=[4.0, 2.0, 4.0, 3.5, 2.0, 1.0])

    add_table_caption(doc, 8, 3, "施工缺失改善照片表")
    make_table(doc,
        headers=["項次","缺失部位","改善前照片說明","改善後照片說明","備註"],
        rows=[("","","（貼照片）","（貼照片）","")]*2,
        col_widths=[1.5, 3.0, 4.5, 4.5, 3.0])

    # ════════════════════════════════════
    # 第九章 矯正與預防措施
    # ════════════════════════════════════
    heading(doc, "第九章　矯正與預防措施", 1)
    heading(doc, "一、矯正措施", 2)
    body(doc, "　　凡工程施工中發現有品質缺失（不合格品），本公司即應進行矯正措施，追查問題根本原因，防止同類缺失再度發生。矯正措施流程如下圖所示：")
    add_flowchart_placeholder(doc, 'fig_9_1', preset_corrective())
    add_figure_caption(doc, 9, 1, "矯正措施流程圖")

    body(doc, "矯正措施步驟：")
    for step in [
        "（一）發現缺失：由品管人員或現場工程師發現並記錄缺失。",
        "（二）缺失分析：分析缺失原因，確認是否為系統性問題。",
        "（三）擬定措施：針對根本原因擬定矯正措施。",
        "（四）執行矯正：依照擬定措施執行改善作業。",
        "（五）複查驗證：由品管人員複查確認已完成改善，且同類缺失未再發生。",
        "（六）結案歸檔：填具矯正措施記錄表，存檔備查。",
    ]:
        body(doc, step)

    heading(doc, "二、預防措施", 2)
    body(doc, "　　本公司應定期召開品質管理會議，分析施工數據、趨勢，針對潛在問題（尚未成為缺失但有惡化跡象）提前採取預防措施，防患於未然。")
    body(doc, "預防措施步驟：")
    for step in [
        "（一）資料蒐集：蒐集施工品質相關數據及紀錄。",
        "（二）趨勢分析：分析數據，識別潛在的問題趨勢。",
        "（三）擬定預防措施：針對潛在問題擬定預防行動。",
        "（四）執行預防：落實預防措施，確保潛在問題不發生。",
        "（五）追蹤確認：追蹤預防措施執行效果，確認有效。",
    ]:
        body(doc, step)

    heading(doc, "三、應用表單", 2)
    add_table_caption(doc, 9, 1, "異常矯正與預防處理紀錄表")
    make_table(doc,
        headers=["編號","異常描述","原因分析","矯正/預防措施","責任人","完成日期","結果驗證"],
        rows=[("","","","","","","")]*3,
        col_widths=[1.5, 3.0, 3.0, 3.5, 2.0, 2.0, 1.5])

    # ════════════════════════════════════
    # 第十章 內部品質稽核
    # ════════════════════════════════════
    heading(doc, "第十章　內部品質稽核", 1)
    heading(doc, "一、品質稽核權責", 2)
    body(doc, "　　本工程內部品質稽核由本公司管理階層（主任技師）帶隊，指派品管人員執行，對工地品質管理系統、施工品質及文件紀錄進行稽核。稽核結果應書面報告，由工地主任確認，並追蹤缺失改善。")

    heading(doc, "二、品質稽核範圍", 2)
    body(doc, "　　品質稽核範圍包括：")
    for item in [
        "（一）品質計畫執行情形：確認各章節規定是否確實執行。",
        "（二）自主檢查表填寫：確認各工項自主檢查表是否確實填寫，檢查結果是否詳實記錄。",
        "（三）材料設備管制：確認材料設備進場記錄、試驗報告是否齊全。",
        "（四）不合格品管制：確認不合格品是否依程序處理及追蹤改善。",
        "（五）文件紀錄管理：確認各類文件紀錄是否妥善保存歸檔。",
    ]:
        body(doc, item)

    heading(doc, "三、品質稽核頻率", 2)
    body(doc, "　　內部品質稽核頻率依工程進度每季至少辦理一次，或於發生重大缺失時即時辦理，並於稽核後30日內完成改善追蹤。")

    heading(doc, "四、品質稽核流程", 2)
    add_flowchart_placeholder(doc, 'fig_10_1', preset_audit_flow())
    add_figure_caption(doc, 10, 1, "品質稽核流程圖")

    heading(doc, "五、應用表單", 2)
    add_table_caption(doc, 10, 1, "內部品質稽核紀錄表")
    make_table(doc,
        headers=["稽核項目","稽核標準","稽核結果","缺失描述","改善期限","改善結果"],
        rows=[
            ("自主檢查表執行","依品質計畫規定","","","",""),
            ("材料試驗記錄","依品質計畫規定","","","",""),
            ("不合格品管制","依品質計畫規定","","","",""),
            ("文件紀錄管理","依品質計畫規定","","","",""),
            ("施工要領執行","依品質計畫規定","","","",""),
        ],
        col_widths=[3.0, 3.0, 2.5, 3.0, 2.0, 3.0])

    # ════════════════════════════════════
    # 第十一章 文件記錄管理系統
    # ════════════════════════════════════
    heading(doc, "第十一章　文件記錄管理系統", 1)
    heading(doc, "一、文件管理系統", 2)
    body(doc, "　　本公司建立文件管理系統，確保所有品質相關文件均有適當之編號、版次管制、發行與廢止程序。所有有效文件均需列入文件清冊，並確保現場使用之文件均為有效版本。")

    body(doc, "文件分類代碼如下：")
    add_table_caption(doc, 11, 1, "文件分類代碼表")
    make_table(doc,
        headers=["代碼","文件類別","說明"],
        rows=[
            ("QP","品質計畫","整體品質計畫及各章節附件"),
            ("EP","工程管理","督察紀錄、會議記錄、函件"),
            ("IN","施工紀錄","施工日誌、自主檢查表"),
            ("TR","試驗報告","材料試驗、強度試驗報告"),
            ("NC","不合格管制","不合格品紀錄、矯正措施記錄"),
            ("AU","稽核紀錄","內部稽核記錄"),
            ("DC","設計文件","圖說、規範"),
        ],
        col_widths=[2.0, 4.0, 10.5])

    heading(doc, "二、收發文管理", 2)
    body(doc, "　　本工程所有往來文件應依下列流程辦理收發文管理：")
    add_flowchart_placeholder(doc, 'fig_11_1', preset_incoming_doc())
    add_figure_caption(doc, 11, 1, "收文傳送及歸檔流程圖")

    add_flowchart_placeholder(doc, 'fig_11_2', preset_outgoing_doc())
    add_figure_caption(doc, 11, 2, "發文傳送及歸檔流程圖")

    heading(doc, "三、紀錄管理作業程序", 2)
    body(doc, "　　品質紀錄應妥善保存，以作為品質符合性及有效性之客觀證據。紀錄管理要求如下：")
    for item in [
        "（一）所有品質紀錄應清晰、完整，標示日期及簽名。",
        "（二）紀錄應儲存於適當之環境，防止損壞、遺失或遭受未授權之存取。",
        "（三）品質紀錄保存期限：自工程竣工驗收後保存至少5年，或依契約規定執行。",
        "（四）涉及爭議或訴訟之紀錄，應永久保存至案件終結。",
    ]:
        body(doc, item)

    heading(doc, "四、文件紀錄移轉及存檔", 2)
    body(doc, "　　工程完工後，所有品質紀錄應整理歸檔，依下列規定辦理移轉：")
    for item in [
        "（一）紙本文件整理裝訂，依文件分類代碼及編號順序排列。",
        "（二）電子文件備份至隨身碟或光碟，連同紙本一併提交監造單位。",
        "（三）移轉清冊：列明所有移轉文件清單，由雙方簽認。",
    ]:
        body(doc, item)

    add_table_caption(doc, 11, 2, "文件紀錄移轉清冊")
    make_table(doc,
        headers=["文件代碼","文件名稱","份數","冊數","備註"],
        rows=[
            ("QP","整體品質計畫","2","1",""),
            ("IN","施工日誌","1","依工期",""),
            ("IN","各工項自主檢查表","1","1",""),
            ("TR","材料試驗報告","1","1",""),
            ("NC","不合格品管制記錄","1","1",""),
            ("EP","督察紀錄","1","1",""),
            ("AU","內部稽核記錄","1","1",""),
        ],
        col_widths=[2.5, 5.0, 2.0, 2.0, 5.0])

    # ════════════════════════════════════
    # 儲存與後處理
    # ════════════════════════════════════
    doc.save(OUT_PATH)
    print(f"📄 DOCX 初步儲存：{OUT_PATH}")

    # 設定三段式頁尾
    sections = doc.sections
    if len(sections) >= 3:
        setup_section_footer(sections[0], 'none')
        setup_section_footer(sections[1], 'lowerRoman')
        setup_section_footer(sections[2], 'decimal')
    elif len(sections) == 2:
        setup_section_footer(sections[0], 'none')
        setup_section_footer(sections[1], 'decimal')
    doc.save(OUT_PATH)

    # DrawML 流程圖替換
    # 因掛載目錄 Permission 限制，先複製到 /tmp 再替換，最後複製回來
    import shutil
    tmp_work = "/tmp/qp_finalize.docx"
    print("🔄 替換 DrawML 流程圖中...")
    shutil.copy2(OUT_PATH, tmp_work)
    finalize_flowcharts(tmp_work,
                        theme_name='quality_plan',
                        font_cjk=DOCX_FONT_CJK,
                        font_latin=DOCX_FONT_LATIN,
                        font_pt=max(DOCX_SZ_TABLE - 1, 9),
                        verbose=True)
    shutil.copy2(tmp_work, OUT_PATH)
    try:
        os.remove(tmp_work)
    except Exception:
        pass

    # 修正 python-docx 已知 Bug
    fix_docx(OUT_PATH)
    print(f"✅ 品質計畫產出完成：{OUT_PATH}")


if __name__ == '__main__':
    build_doc()
