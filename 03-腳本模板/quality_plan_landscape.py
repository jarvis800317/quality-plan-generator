#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
整體品質計畫產生腳本
工程：桃園市第47期桃園區中平市地重劃工程
章節制：11章制（壹～拾壹）
產出：DOCX（含 DrawML 向量流程圖）
"""

import os, sys, io, zipfile, re, copy
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
from fontTools.ttLib import TTFont as _TTFont

from drawml_flowchart import (
    add_flowchart_placeholder, add_three_col_flow, finalize_flowcharts,
    preset_material_approval, preset_work_inspection,
    preset_nc_material, preset_nc_work,
    preset_corrective, preset_incoming_doc, preset_outgoing_doc,
    preset_equipment_test, preset_audit_flow,
    build_quality_org_chart,
)



# ════════════════════════════════════════════════════════════
# ⭐ 工程設定區（每次新工程必須更新以下所有欄位）
# ════════════════════════════════════════════════════════════

ENG_NAME    = "桃園市第47期桃園區中平市地重劃工程"
ENG_OWNER   = "桃園市政府地政局"
ENG_DESIGN  = "杜風工程服務股份有限公司（設計技師 趙厚任）"
ENG_SUPER   = "杜風工程服務股份有限公司（監造技師 高堉霖）"
ENG_CONTR   = "（承攬廠商待補）"
ENG_SITE    = "桃園市桃園區"
ENG_DAYS    = "270日曆天（B-2、B-3區同時開工）"
ENG_AMOUNT  = "新臺幣壹億零柒佰玖拾肆萬柒仟元整（107,947,000元）"
ENG_START   = "民國115年4月"
ENG_END     = "民國116年5月"
ENG_FOREMAN = "（工地主任待補）"
ENG_QA      = "（品管人員待補）（結業證書字號：待補）"
ENG_TECH    = "（主任技師待補）（技師證號：待補）"
ENG_SAFETY  = "（職業安全衛生管理人員待補）"

# ── 章節選擇（全11章）──
CHAPTER_FLAGS = {
    'scope':      True,
    'mgmt':       True,
    'method':     True,
    'standard':   True,
    'inspection': True,
    'equipment':  True,
    'selfcheck':  True,
    'nc':         True,
    'ca':         True,
    'audit':      True,
    'doc':        True,
}

# ── 字型設定（依監造計畫分析）──
DOCX_FONT_CJK   = "標楷體"      # DFKai-SB
DOCX_FONT_LATIN = "Arial"

DOCX_SZ_BODY       = 12
DOCX_SZ_H1         = 20
DOCX_SZ_H2         = 14
DOCX_SZ_H3         = 12
DOCX_SZ_TOC        = 12
DOCX_SZ_TABLE      = 11
DOCX_SZ_COVER_MAIN = 36
DOCX_SZ_COVER_PROJ = 22
DOCX_SZ_COVER_INFO = 14
DOCX_SZ_NOTE       = 9

# ── 輸出路徑（動態計算絕對路徑）──
_WORKSPACE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH = os.path.join(_WORKSPACE, f"{ENG_NAME}-整體品質計畫.docx")

# ── PIL 字體路徑（請依執行環境調整，以下為 Linux sandbox 路徑）──
FONT_CJK   = "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf"
FONT_ASCII = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"

# ════════════════════════════════════════════════════════════
# ⚠️ 工程主要施工項目及數量表（100%完整複製監造計畫表1.3-1）
# ════════════════════════════════════════════════════════════
CONSTRUCTION_ITEMS = [
    # (項次, 項目及說明, 單位, 數量)
    ('壹', '直接工程費', '', ''),
    ('壹.一', '整地拆除工程', '', ''),
    ('壹.一.1', '清除及掘除', 'M2', '15,148.00'),
    ('壹.一.2', '地坪打除，未含運費', 'M2', '11,545.00'),
    ('壹.一.3', '機械拆除，鋼筋混凝土，未含運費', 'M3', '951.00'),
    ('壹.一.4', '基地及路幅開挖(不含近運利用)', 'M3', '1,266.00'),
    ('壹.一.5', '基地及路堤填築', 'M3', '1,295.00'),
    ('壹.一.6', '餘方近運利用', 'M3', '1,266.00'),
    ('壹.一.7', 'PVC披覆鐵絲網柵欄，鍍鋅鋼管支柱', 'M', '624.00'),
    ('壹.一.8', '路幅整修，土方層整修', 'M2', '29,236.00'),
    ('壹.一.9', '草籽撒播', 'M2', '9,812.00'),
    ('壹.一.10', '鋼筋混凝土管(B型)，D=300mm，三級管，含管配件', 'M', '14.00'),
    ('壹.一.11', '禁倒廢土警告標示牌及安裝', '座', '7.00'),
    ('壹.一.12', '環境保護，空氣污染防制，防塵網', 'M2', '9,812.00'),
    ('壹.一.13', '機械切割', 'M', '161.00'),
    ('壹.一.14', '機械拆除，建築物拆除，鐵皮屋，未含運費', '式', '1.00'),
    ('壹.一.15', '籃球場地上物拆除', '式', '1.00'),
    ('壹.一.16', '廢棄物處理，營建廢棄物清除(含集中、運棄、處理)', 'M3', '3,260.00'),
    ('壹.一.17', '一般廢棄物清除(含集中、運棄、處理)', 'T', '200.00'),
    ('壹.一.18', '廢方處理(含運費)', 'M3', '100.00'),
    ('壹.一.19', '環境保護，基地開挖3M及檢視', '處', '7.00'),
    ('壹.二', '景觀工程', '', ''),
    ('壹.二.1', '景觀設施工程', '', ''),
    ('壹.二.1.1', '基地及路幅開挖(不含近運利用)', 'M3', '2,112.00'),
    ('壹.二.1.2', '餘方近運利用', 'M3', '2,115.00'),
    ('壹.二.1.3', '水泥混凝土舖面，停車場鋪面', 'M2', '1,557.00'),
    ('壹.二.1.4', '水泥混凝土舖面，路邊停車鋪面', 'M2', '294.00'),
    ('壹.二.1.5', '人行道面層，鋪面磚(高壓)，警示磚', 'M2', '29.00'),
    ('壹.二.1.6', '人行道面層，鋪面磚(高壓)，高壓混凝土磚', 'M2', '4,144.00'),
    ('壹.二.1.7', '緣石，預鑄，路緣石，A型', 'M', '572.00'),
    ('壹.二.1.9', '緣石，預鑄，路緣石，C型', 'M', '1,290.00'),
    ('壹.二.1.10', '緣石，鋼板緣石', 'M', '24.00'),
    ('壹.二.1.11', '室外簡易籃球場', '式', '1.00'),
    ('壹.二.1.12', '室外簡易籃球場，鍍鋅鋼網圍籬，H=3.08m', 'M', '119.00'),
    ('壹.二.1.13', '室外簡易籃球場，鍍鋅鋼網圍籬雙開門', '組', '1.00'),
    ('壹.二.1.14', '兒童遊戲場設備，樓梯', '組', '2.00'),
    ('壹.二.1.15', '兒童遊戲場設備，磨石子滑梯', '式', '1.00'),
    ('壹.二.1.16', '兒童遊戲場設備，洗手台結構及面層', '組', '1.00'),
    ('壹.二.1.17', '兒童遊戲場設備，洗手台水電', '組', '1.00'),
    ('壹.二.1.18', '兒童遊戲場設備，彩色噴花地坪', 'M2', '233.00'),
    ('壹.二.1.19', '兒童遊戲場設備，彈性無縫鋪面 t=7.5cm', 'M2', '37.00'),
    ('壹.二.1.20', '兒童遊戲場設備，彈性無縫鋪面 t=4.5cm', 'M2', '67.00'),
    ('壹.二.1.21', '兒童遊戲場設備，彈性橡膠地墊 t=9cm', 'M2', '5.00'),
    ('壹.二.1.22', '兒童遊戲場設備，彈性橡膠地墊 t=4.5cm', 'M2', '93.00'),
    ('壹.二.1.23', '兒童遊戲場設備，鋪面收邊', 'M', '34.00'),
    ('壹.二.1.24', '兒童遊戲場設備，2.5cm人工草皮', 'M2', '83.00'),
    ('壹.二.1.25', '兒童遊戲場設備，樹皮塊地坪', 'M2', '130.00'),
    ('壹.二.1.26', '兒童遊戲場設備，礫石地坪', 'M2', '341.00'),
    ('壹.二.1.27', '兒童遊戲場設備，黑色隔草板', 'M', '95.00'),
    ('壹.二.1.28', '標誌，無障礙停車位指示牌', '組', '2.00'),
    ('壹.二.1.29', '標誌，婦幼車位指示牌', '組', '2.00'),
    ('壹.二.1.30', '標誌，公園告示牌', '組', '1.00'),
    ('壹.二.1.31', '標誌，公園管理自治條例告示牌', '組', '1.00'),
    ('壹.二.1.32', '標誌，熱拌標線，厚2mm', 'M2', '304.00'),
    ('壹.二.1.33', '公園廁所，地面基礎及水電', '式', '1.00'),
    ('壹.二.1.34', '公園廁所，男廁單元', '式', '1.00'),
    ('壹.二.1.35', '公園廁所，女廁單元', '式', '1.00'),
    ('壹.二.1.36', '公園廁所，親子廁單元', '式', '1.00'),
    ('壹.二.1.37', '公園廁所，連接鋼架與頂板', '式', '1.00'),
    ('壹.二.2', '木作工程', '', ''),
    ('壹.二.2.1', '粗木作，砂坑塑木花架', '座', '1.00'),
    ('壹.二.2.2', '粗木作，砂坑塑木花架雨遮', '座', '1.00'),
    ('壹.二.2.3', '粗木作，砂坑塑木木平台', '座', '1.00'),
    ('壹.二.2.4', '粗木作，塑木花廊架鋼構', '座', '1.00'),
    ('壹.二.2.5', '粗木作，塑木花廊架', '座', '1.00'),
    ('壹.二.2.6', '粗木作，塑木花廊架雙面座椅', '組', '5.00'),
    ('壹.二.2.7', '粗木作，塑木靠背座椅', '座', '5.00'),
    ('壹.二.2.8', '粗木作，塑木座椅A', '座', '10.00'),
    ('壹.二.2.9', '粗木作，塑木座椅B', '座', '1.00'),
    ('壹.二.2.10', '粗木作，洗手台塑木踏板', '座', '1.00'),
    ('壹.二.2.11', '粗木作，攀爬架A', '座', '1.00'),
    ('壹.二.2.12', '粗木作，攀爬架B', '座', '1.00'),
    ('壹.二.2.13', '粗木作，圓柱攀爬架', '座', '1.00'),
    ('壹.二.3', '遊戲場設施工程', '', ''),
    ('壹.二.3.1', '兒童遊戲場設備，划船器', '組', '1.00'),
    ('壹.二.3.2', '兒童遊戲場設備，雙人浪板器', '組', '1.00'),
    ('壹.二.3.3', '兒童遊戲場設備，雙人漫步機', '組', '1.00'),
    ('壹.二.3.4', '兒童遊戲場設備，高中低單槓 230-200-170cm', '組', '1.00'),
    ('壹.二.3.5', '兒童遊戲場設備，複合式鞦韆', '組', '1.00'),
    ('壹.二.3.6', '兒童遊戲場設備，彈跳床 L150*W150cm', '組', '1.00'),
    ('壹.二.3.7', '兒童遊戲場設備，極限旋轉飛輪', '組', '1.00'),
    ('壹.二.3.8', '兒童遊戲場設備，擺盪大索', '組', '1.00'),
    ('壹.二.3.9', '兒童遊戲場設備，攀爬繩網', '組', '1.00'),
    ('壹.二.3.10', '兒童遊戲場設備，遊戲場告示牌', '組', '6.00'),
    ('壹.二.4', '公園排水工程', '', ''),
    ('壹.二.4.1', '排水管溝，污水下水道連接、開挖及修復', '式', '1.00'),
    ('壹.二.4.2', '排水管溝，PE地下儲水槽', '組', '3.00'),
    ('壹.二.4.3', '排水管溝，新設陰井90x90x135cm', '組', '1.00'),
    ('壹.二.4.4', '排水管溝，新設陰井70x70x75cm', '組', '36.00'),
    ('壹.二.4.5', '排水管溝，自設污水陰井內徑60cm', '組', '2.00'),
    ('壹.二.4.6', '排水管溝，HDPE透水網管(含埋設)', 'M', '842.00'),
    ('壹.二.4.7', '排水管溝，3"PVC加強埋設', 'M', '81.00'),
    ('壹.二.4.8', '排水管溝，2"PVC加強埋設', 'M', '117.00'),
    ('壹.二.4.9', '排水管溝，4"污水用PVC管', 'M', '77.00'),
    ('壹.二.4.10', '排水管溝，既有緣石打除及運棄', 'M', '286.00'),
    ('壹.二.4.11', '排水管溝，既有預鑄溝蓋打除及運棄', 'M', '104.00'),
    ('壹.二.4.12', '排水管溝，新作溝蓋板', 'M', '104.00'),
    ('壹.二.4.13', '排水管溝，鍍鋅隔柵板(w=50cm)', '組', '24.00'),
    ('壹.三', '植栽工程', '', ''),
    ('壹.三.1', '植栽維護，喬木類，枯木移除(喬木類)，樹徑≦30cm', '株', '40.00'),
    ('壹.三.2', '植栽維護，含斷根及定植作業', '株', '11.00'),
    ('壹.三.3', '植樹，杜英，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '24.00'),
    ('壹.三.4', '植樹，樟樹，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '56.00'),
    ('壹.三.5', '植樹，烏臼，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '13.00'),
    ('壹.三.6', '植樹，青剛櫟，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '11.00'),
    ('壹.三.7', '植樹，洋紅風鈴木，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '5.00'),
    ('壹.三.8', '植樹，黃花風鈴木，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '9.00'),
    ('壹.三.9', '植樹，台灣櫸木，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '7.00'),
    ('壹.三.10', '植樹，茄苳，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '32.00'),
    ('壹.三.11', '植樹，青楓，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '6.00'),
    ('壹.三.12', '植樹，光臘樹，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '20.00'),
    ('壹.三.13', '植樹，雀榕，350cm≦樹高，130cm≦樹幅，9cm≦米高直徑', '株', '3.00'),
    ('壹.三.14', '植樹，台灣欒樹，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '6.00'),
    ('壹.三.15', '植樹，流蘇，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '5.00'),
    ('壹.三.16', '植樹，藍花楹，270cm≦樹高，100cm≦樹幅，6cm≦米高直徑', '株', '11.00'),
    ('壹.三.17', '植樹，羅漢松，150cm≦樹高，90cm≦樹幅，5cm≦米高直徑', '株', '38.00'),
    ('壹.三.18', '植樹，灌木類，月橘，35cm≦高度，30cm≦寬度', '株', '361.00'),
    ('壹.三.19', '植樹，灌木類，厚葉石斑木，35cm≦高度，30cm≦寬度', '株', '400.00'),
    ('壹.三.20', '植草，假儉草(含施工養護)', 'M2', '11,107.00'),
    ('壹.四', '景觀照明工程', '', ''),
    ('壹.四.1', '照明燈具，電源開關箱，不銹鋼板，板厚2.0mm', '只', '3.00'),
    ('壹.四.2', '照明燈具，停車場單燈', '組', '8.00'),
    ('壹.四.3', '照明燈具，景觀高燈', '組', '27.00'),
    ('壹.四.4', '照明燈具，景觀柱燈', '組', '30.00'),
    ('壹.四.5', '照明燈具，籃球場投光燈', '組', '6.00'),
    ('壹.四.6', '照明燈具，多向性投光燈', '組', '8.00'),
    ('壹.四.7', '照明燈具，投光燈', '組', '9.00'),
    ('壹.四.8', 'XLPE 電纜線 5.5mm2*2C 600V(含損耗)', 'M', '1,834.00'),
    ('壹.四.9', 'PVC WIRE 接地線 3.5mm2 600V (含損耗)', 'M', '1,834.00'),
    ('壹.四.10', 'PVC PIPE ∮41mm(含損耗)', 'M', '1,834.00'),
    ('壹.四.11', 'PVC 電纜線 3.5mm2* 3/C(燈桿內)', 'M', '461.00'),
    ('壹.四.12', '平面式塑膠警示帶', 'M', '1,834.00'),
    ('壹.四.13', '開關箱至TPC管線材料含安裝', '處', '2.00'),
    ('壹.四.14', '申請用電文件費', '式', '1.00'),
    ('貳', '間接工程費', '', ''),
    ('貳.一', '雜項工程', '', ''),
    ('貳.一.1', '工程告示牌及工地標誌，工程告示牌，長300x寬170cm，基礎及現場裝拆(含維護)', '面', '2.00'),
    ('貳.一.2', '施工照相及攝(錄)影', '式', '1.00'),
    ('貳.一.3', '工地臨時建築設施，工地辦公室及設備', '式', '1.00'),
    ('貳.一.4', '工地臨時建築設施，守衛亭，鋁製', '式', '1.00'),
    ('貳.一.5', '施工圍籬，全阻隔式固定(第一級營建工程)，含(頂部)警告燈，折舊', 'M', '966.00'),
    ('貳.一.6', '施工圍籬，半阻隔式固定(第一級營建工程)，含(頂部)警告燈，折舊', 'M', '257.00'),
    ('貳.一.7', '工地臨時建築設施，出入口拉開式大門', '座', '4.00'),
    ('貳.一.8', '臨時施工便道', 'M2', '144.00'),
    ('貳.一.9', '臨時施工便道，鋼板使用費', '式', '1.00'),
    ('貳.一.10', '施工測量', '式', '1.00'),
    ('貳.一.11', '公共管線系統之保護', '式', '1.00'),
    ('貳.一.12', '袪水，擋抽排水', '式', '1.00'),
    ('貳.一.13', '工地即時監控系統', '式', '1.00'),
    ('貳.一.14', '環境保護，空氣污染防制，防塵網', 'M2', '10,000.00'),
    ('貳.一.15', '臨時排水攔砂及排水設施', 'M', '80.00'),
    ('貳.一.16', '臨時性沉砂池(32.4mX10m)', '處', '1.00'),
    ('貳.一.18', '既有設施維護及管理費', '式', '1.00'),
    ('貳.二', '工程材料檢(試)驗費', '', ''),
    ('貳.二.1', '品質管理，試驗規範及標準，A3045混凝土圓柱試體抗壓強度之檢驗法', '組', '18.00'),
    ('貳.二.2', '品質管理，試驗規範及標準，工地密度試驗', '組', '4.00'),
    ('貳.二.3', '品質管理，試驗規範及標準，土壤夯實試驗', '次', '4.00'),
    ('貳.二.4', '品質管理，試驗規範及標準，竹節鋼筋拉伸試驗', '隻', '3.00'),
    ('貳.二.5', '品質管理，試驗規範及標準，竹節鋼筋彎曲試驗', '件', '2.00'),
    ('貳.二.6', '品質管理，試驗規範及標準，砂漿抗壓強度試驗(2015)', '件', '1.00'),
    ('貳.二.7', '品質管理，試驗規範及標準，銲接鋼線網試驗', '件', '6.00'),
    ('貳.二.8', '品質管理，試驗規範及標準，高壓混凝土磚抗壓強度試驗', '組', '5.00'),
    ('貳.二.9', '品質管理，試驗規範及標準，高壓混凝土磚吸水率試驗', '組', '5.00'),
    ('貳.二.10', '品質管理，試驗規範及標準，路緣石抗彎強度試驗符合CNS13295規定', '組', '3.00'),
    ('貳.二.11', '品質管理，試驗規範及標準，瓷磚抗壓強度試驗-符合CNS453規定', '片', '3.00'),
    ('貳.二.12', '品質管理，試驗規範及標準，瓷磚抗彎強度試驗-符合CNS4392規定', '片', '3.00'),
    ('貳.二.13', '品質管理，試驗規範及標準，瓷磚止滑係數試驗-符合CNS8907規定', '片', '3.00'),
    ('貳.二.14', '其它材料設備檢驗及抽驗費', '式', '1.00'),
    ('貳.三', '交通維持費', '', ''),
    ('貳.三.1', '活動型拒馬(含警示燈)', '座', '20.00'),
    ('貳.三.2', '施工標誌(附掛式)', '座', '20.00'),
    ('貳.三.3', '施工警告標示牌(含基礎)', '座', '10.00'),
    ('貳.三.4', '施工警告燈號(租用)', '個', '30.00'),
    ('貳.三.5', '紐澤西護欄，預鑄混凝土', '座', '5.00'),
    ('貳.三.6', '交通錐(含連桿)', '個', '50.00'),
    ('貳.三.7', '人工旗手(含服裝、指揮棒等)', '工', '30.00'),
    ('貳.三.8', '義交(日間，含指揮棒等)', '時', '20.00'),
    ('貳.三.9', '其他交通設施及管理維護費', '式', '1.00'),
    ('貳.四', '環境保護費', '', ''),
    ('貳.四.1', '洗車台設備及沉澱池', '座', '1.00'),
    ('貳.四.2', '環境保護，施工便道灑水(含運輸道路)', '式', '1.00'),
    ('貳.四.3', '環境保護，工區臨近道路維護清理', '式', '1.00'),
    ('貳.四.4', '環境保護，環境監測', '式', '1.00'),
    ('貳.五', '職業安全衛生', '', ''),
    ('貳.五.1', '安衛組織費', '式', '1.00'),
    ('貳.五.2', '職業安全衛生設備', '式', '1.00'),
    ('貳.五.3', '安衛管理及其他安衛措施', '式', '1.00'),
    ('貳.六', '工程品管費', '', ''),
    ('貳.六.1', '竣工文件', '式', '1.00'),
    ('貳.六.2', '品質管理', '式', '1.00'),
    ('貳.七', '包商利潤及雜項費', '式', '1.00'),
    ('貳.八', '營造綜合保險費', '式', '1.00'),
]
print(f"✅ 施工項目共 {len(CONSTRUCTION_ITEMS)} 列（含大項標題）")

# ════════════════════════════════════════════════════════════
# 品質管理主要工項（用於參/肆/柒章表格）
# ════════════════════════════════════════════════════════════
MAIN_WORK_ITEMS = [
    "整地拆除工程",
    "混凝土鋪面工程",
    "鋪面磚（高壓混凝土磚）工程",
    "路緣石工程",
    "籃球場工程",
    "塑木（木作）工程",
    "兒童遊戲場設備工程",
    "公園排水工程",
    "植栽工程",
    "景觀照明工程",
]

# ════════════════════════════════════════════════════════════
# 輔助函式
# ════════════════════════════════════════════════════════════

def set_run_font(run, size_pt, bold=False, color=(0, 0, 0)):
    run.font.name = DOCX_FONT_LATIN
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    run._element.rPr.rFonts.set(qn('w:ascii'), DOCX_FONT_LATIN)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), DOCX_FONT_LATIN)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.page_break_before = True
    for run in p.runs:
        run.font.name = DOCX_FONT_LATIN
        run.font.color.rgb = RGBColor(0, 0, 0)
        if run._element.rPr is not None:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    return p


def body(doc, text, indent=0, bold=False):
    p = doc.add_paragraph()
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, DOCX_SZ_BODY, bold=bold)
    return p


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    illegal_tags = {'w:noWrap', 'w:tcMar', 'w:vAlign', 'w:hideMark'}
    first_illegal = next(
        (child for child in tcPr if child.tag.split('}')[-1] in
         {t.split(':')[-1] for t in illegal_tags}), None)
    if first_illegal is not None:
        tcPr.insertBefore(shd, first_illegal)
    else:
        tcPr.append(shd)


def make_table(doc, headers, rows, col_widths, header_bg='F2F2F2', header_fg=(0, 0, 0), font_size=None):
    if font_size is None:
        font_size = DOCX_SZ_TABLE
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row = table.rows[0]
    for ci, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[ci]
        cell.width = Cm(w)
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(*header_fg)
        run.font.size = Pt(font_size)
        run.font.name = DOCX_FONT_LATIN
        if run._element.rPr is not None:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, (cell_txt, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[ci]
            cell.width = Cm(w)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_txt))
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = DOCX_FONT_LATIN
            if run._element.rPr is not None:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    return table


def add_table_caption(doc, chapter, seq, name):
    caption_text = f"表{chapter}-{seq} {name}"
    p = doc.add_paragraph(caption_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(DOCX_SZ_TABLE)
    run.font.name = DOCX_FONT_LATIN
    run.font.color.rgb = RGBColor(0, 0, 0)
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    return p


def add_figure_caption(doc, chapter, seq, name):
    caption_text = f"圖{chapter}-{seq} {name}"
    p = doc.add_paragraph(caption_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(DOCX_SZ_BODY)
    run.font.name = DOCX_FONT_LATIN
    run.font.color.rgb = RGBColor(0, 0, 0)
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    return p


def _add_right_tab(paragraph, pos_cm=15.0):
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:tabs')):
        pPr.remove(old)
    tabs_el = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), str(int(pos_cm * 567)))
    tabs_el.append(tab)
    pPr.append(tabs_el)


def make_toc(doc, entries, right_pos_cm=15.0):
    for entry in entries:
        p = doc.add_paragraph()
        _add_right_tab(p, right_pos_cm)
        if entry['level'] == 2:
            p.paragraph_format.left_indent = Cm(1.0)
        run_title = p.add_run(entry['title'])
        run_title.bold = (entry['level'] == 1)
        run_title.font.size = Pt(DOCX_SZ_TOC)
        run_title.font.name = DOCX_FONT_LATIN
        run_title.font.color.rgb = RGBColor(0, 0, 0)
        if run_title._element.rPr is not None:
            run_title._element.rPr.rFonts.set(qn('w:eastAsia'), DOCX_FONT_CJK)
        run_page = p.add_run(f'\t{entry["page"]}')
        run_page.bold = (entry['level'] == 1)
        run_page.font.size = Pt(DOCX_SZ_TOC)
        run_page.font.name = DOCX_FONT_LATIN
        run_page.font.color.rgb = RGBColor(0, 0, 0)
        if run_page._element.rPr is not None:
            run_page._element.rPr.rFonts.set(qn('w:eastAsia'), DOCX_FONT_CJK)


def _add_page_num_field(paragraph):
    run = paragraph.add_run()
    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    run._r.extend([fc_begin, instr, fc_sep, fc_end])
    return paragraph


def _set_pg_num_type(sectPr, fmt, start=1):
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)


def setup_section_footer(section, fmt):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        for r in p.runs:
            r.text = ''
    if fmt == 'none':
        return
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_num_field(p)
    _set_pg_num_type(section._sectPr, fmt, start=1)


def insert_section_break(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')
    pgSz.set(qn('w:h'), '16838')
    sectPr.append(pgSz)
    pgType = OxmlElement('w:type')
    pgType.set(qn('w:val'), 'nextPage')
    sectPr.append(pgType)
    pPr.append(sectPr)
    return p


def fix_docx(out_path):
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path, 'r') as zi, \
         zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zo:
        for item in zi.infolist():
            data = zi.read(item.filename)
            if item.filename == 'word/settings.xml':
                data = re.sub(b'<w:zoom[^/]*/>', b'<w:zoom w:percent="100"/>', data)
            zo.writestr(item, data)
    os.replace(tmp, out_path)
    print(f"✅ DOCX 已儲存並修正：{out_path}")


def selfcheck_footer(doc):
    body(doc, "缺失複查結果：")
    body(doc, "□已完成改善")
    body(doc, "□未完成改善，填至「不合格管制總表」第○項進行追蹤改善")
    doc.add_paragraph()
    body(doc, "複查日期：　年　月　日　　複查人員職稱：　　　　　簽名：")
    doc.add_paragraph()
    body(doc, "備註：")
    body(doc, "1. 檢查標準應具體明確或量化尺寸。", indent=0.5)
    body(doc, "2. 檢查結果合格「○」，不合格「╳」，不需檢查「／」。", indent=0.5)
    body(doc, "3. 嚴重缺失未及時完成改善，應填具「不合格品管制總表」追蹤改善。", indent=0.5)
    body(doc, "4. 本表由工地現場施工人員實地檢查後覈實記載簽認。", indent=0.5)
    doc.add_paragraph()
    body(doc, "工地主任（工地負責人）簽名：＿＿＿＿　　現場施工人員（檢查人員）簽名：＿＿＿＿")


# ════════════════════════════════════════════════════════════
# 主程式：建立文件
# ════════════════════════════════════════════════════════════

def build_doc():
    doc = Document()

    # ── 設定頁面邊界（A4，上下左右各2.5cm）──
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)

    # ── 設定全域字型樣式 ──
    normal_style = doc.styles['Normal']
    normal_style.font.name = DOCX_FONT_LATIN
    normal_style.font.size = Pt(DOCX_SZ_BODY)
    _nr = normal_style.element.get_or_add_rPr()
    _nf = _nr.get_or_add_rFonts()
    _nf.set(qn('w:eastAsia'), DOCX_FONT_CJK)
    _nf.set(qn('w:ascii'), DOCX_FONT_LATIN)
    _nf.set(qn('w:hAnsi'), DOCX_FONT_LATIN)

    for lvl, sz in [(1, DOCX_SZ_H1), (2, DOCX_SZ_H2), (3, DOCX_SZ_H3)]:
        s = doc.styles[f'Heading {lvl}']
        s.font.name = DOCX_FONT_LATIN
        s.font.size = Pt(sz)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after  = Pt(6)
        if lvl == 1:
            s.paragraph_format.page_break_before = True
        _hr = s.element.get_or_add_rPr()
        _hf = _hr.get_or_add_rFonts()
        _hf.set(qn('w:eastAsia'), DOCX_FONT_CJK)
        _hf.set(qn('w:ascii'), DOCX_FONT_LATIN)
        _hf.set(qn('w:hAnsi'), DOCX_FONT_LATIN)

    # ════════════════════════════════
    # 封面
    # ════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p_main = doc.add_paragraph()
    p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_main.add_run("整體品質計畫")
    set_run_font(run, DOCX_SZ_COVER_MAIN, bold=True)

    doc.add_paragraph()

    p_proj = doc.add_paragraph()
    p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_proj.add_run(ENG_NAME)
    set_run_font(run, DOCX_SZ_COVER_PROJ, bold=True)

    for _ in range(3):
        doc.add_paragraph()

    for label, value in [
        ("主辦機關：", ENG_OWNER),
        ("監造單位：", ENG_SUPER),
        ("承攬廠商：", ENG_CONTR),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label + value)
        set_run_font(run, DOCX_SZ_COVER_INFO)

    doc.add_paragraph()
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_date.add_run("中華民國　　　年　　月")
    set_run_font(run, DOCX_SZ_COVER_INFO)

    insert_section_break(doc)

    # ════════════════════════════════
    # 目錄
    # ════════════════════════════════
    p_toc = doc.add_paragraph()
    p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_toc.add_run("目　　錄")
    set_run_font(run, DOCX_SZ_H1, bold=True)
    doc.add_paragraph()

    toc_entries = [
        {'level': 1, 'title': '壹、計畫範圍',                 'page': 1},
        {'level': 2, 'title': '  一、前言',                   'page': 1},
        {'level': 2, 'title': '  二、依據',                   'page': 1},
        {'level': 2, 'title': '  三、工程概要',               'page': 2},
        {'level': 2, 'title': '  四、工程主要施工項目及數量', 'page': 3},
        {'level': 2, 'title': '  五、適用對象',               'page': 6},
        {'level': 2, 'title': '  六、名詞定義',               'page': 6},
        {'level': 1, 'title': '貳、管理責任及權責分工',       'page': 8},
        {'level': 2, 'title': '  一、品管組織架構',           'page': 8},
        {'level': 2, 'title': '  二、工作職掌',               'page': 9},
        {'level': 2, 'title': '  三、人員輪值及請休假事項',   'page': 11},
        {'level': 2, 'title': '  四、管理審查',               'page': 12},
        {'level': 1, 'title': '參、施工要領',                 'page': 12},
        {'level': 2, 'title': '  一、分項計畫管制表',         'page': 12},
        {'level': 2, 'title': '  二、各工項施工要領',         'page': 13},
        {'level': 1, 'title': '肆、品質管理標準',             'page': 30},
        {'level': 2, 'title': '  一、各工項品質管理標準表',   'page': 30},
        {'level': 1, 'title': '伍、材料及施工檢驗程序',       'page': 55},
        {'level': 2, 'title': '  一、材料設備選定送審流程',   'page': 55},
        {'level': 2, 'title': '  二、施工檢驗流程',           'page': 56},
        {'level': 2, 'title': '  三、各工項施工作業流程及檢驗圖', 'page': 57},
        {'level': 2, 'title': '  四、材料設備送審管制總表',   'page': 80},
        {'level': 1, 'title': '陸、設備功能運轉測試',         'page': 85},
        {'level': 2, 'title': '  一、設備功能運轉檢測程序',   'page': 85},
        {'level': 2, 'title': '  二、設備功能運轉檢測標準',   'page': 86},
        {'level': 1, 'title': '柒、自主檢查表',               'page': 90},
        {'level': 1, 'title': '捌、不合格品之管制',           'page': 120},
        {'level': 2, 'title': '  一、不合格品管制流程',       'page': 120},
        {'level': 2, 'title': '  二、不合格品管制表單',       'page': 122},
        {'level': 1, 'title': '玖、矯正與預防措施',           'page': 125},
        {'level': 1, 'title': '拾、內部品質稽核',             'page': 128},
        {'level': 2, 'title': '  一、品質稽核流程',           'page': 128},
        {'level': 2, 'title': '  二、品質稽核計畫表',         'page': 129},
        {'level': 1, 'title': '拾壹、文件紀錄管理系統',       'page': 132},
        {'level': 2, 'title': '  一、文件收發管理流程',       'page': 132},
        {'level': 2, 'title': '  二、文件分類代碼表',         'page': 134},
    ]
    make_toc(doc, toc_entries)

    insert_section_break(doc)

    # ════════════════════════════════
    # 壹、計畫範圍
    # ════════════════════════════════
    heading(doc, "壹、計畫範圍", level=1)

    heading(doc, "一、前言", level=2)
    body(doc, f"本公司承攬「{ENG_NAME}」，為確保工程品質達到契約規範要求，特依據「公共工程施工品質管理作業要點」及相關法令，訂定本整體品質計畫，作為本工程施工品質管理之依據。")
    body(doc, "本公司將依本品質計畫確實執行施工品質管理，並積極配合主辦機關及監造單位之品質抽查作業，以確保工程如期如質完成。")

    heading(doc, "二、依據", level=2)
    for item in [
        "（一）工程契約（含圖說、投標須知、補充規定及施工說明書）",
        "（二）公共工程委員會頒布之「公共工程施工品質管理作業要點」",
        "（三）公共工程施工綱要規範",
        "（四）公共工程專業技師簽證規則",
        "（五）技師法",
        "（六）營造業法",
        "（七）職業安全衛生法及相關子法",
        "（八）環境保護相關法規",
        "（九）政府採購法及相關函示",
        "（十）其他相關法令規定",
    ]:
        body(doc, item, indent=0.5)

    heading(doc, "三、工程概要", level=2)
    overview_items = [
        ("工程名稱", ENG_NAME),
        ("主辦機關", ENG_OWNER),
        ("設計單位", ENG_DESIGN),
        ("監造單位", ENG_SUPER),
        ("承攬廠商", ENG_CONTR),
        ("工程地點", ENG_SITE),
        ("工程期限", f"預定開工：{ENG_START}；預定完工：{ENG_END}"),
        ("契約工期", ENG_DAYS),
        ("契約金額", ENG_AMOUNT),
        ("工地主任", ENG_FOREMAN),
        ("品管人員", ENG_QA),
        ("主任技師", ENG_TECH),
    ]
    for label, value in overview_items:
        body(doc, f"{label}：{value}", indent=0.5)

    body(doc, "\n工程規模概述：")
    for item in ["一、整地拆除工程", "二、景觀工程（含景觀設施、木作、遊戲場設施、公園排水）", "三、植栽工程", "四、景觀照明工程"]:
        body(doc, item, indent=0.5)

    heading(doc, "四、工程主要施工項目及數量", level=2)
    add_table_caption(doc, 1, 2, "工程主要施工項目及數量表")
    make_table(doc,
               headers=["項次", "項目及說明", "單位", "數量"],
               rows=CONSTRUCTION_ITEMS,
               col_widths=[2.0, 10.5, 2.0, 2.0])

    heading(doc, "五、適用對象", level=2)
    body(doc, f"本品質計畫適用於「{ENG_NAME}」之施工品質管理，舉凡承攬廠商、材料供應商、設備製造商及分包廠商等，均應依本品質計畫執行各項作業，以確保工程品質符合契約規範要求。")

    heading(doc, "六、名詞定義", level=2)
    terms = [
        ("主辦機關", f"{ENG_OWNER}。"),
        ("監造單位", f"{ENG_SUPER}。"),
        ("承攬廠商（本公司）", f"本工程之得標廠商。"),
        ("工地主任（工地負責人）", "指由承攬廠商指派，負責本工程現場施工管理之人員。"),
        ("品管人員", "指由承攬廠商依法設置，負責本工程施工品質管理之人員。"),
        ("主任技師", "指由承攬廠商指派，依法簽證並負責本工程技術監督之技師。"),
        ("協力廠商", "指由本公司依工程需要，分包各項施工項目之廠商。"),
        ("整體品質計畫", "指本公司就本工程整體施工品質所訂定之管理計畫書。"),
        ("分項品質計畫", "指本公司就各分項工程施工品質所訂定之計畫書。"),
        ("品質管理標準", "指各施工項目之品質管理基準，含施工標準、允許差值及試驗頻率。"),
        ("自主檢查", "指承攬廠商於施工前、中、後，自行對施工品質進行查驗之作業。"),
        ("檢驗停留點（▽）", "指施工作業中，必須經自主檢查合格後方可繼續施工之管制點。"),
        ("材料/設備送審", "指承攬廠商將擬用材料或設備之規格、型錄等送請監造單位審核。"),
        ("不合格品", "指不符合品質要求之材料、設備或施工成果。"),
        ("矯正措施", "指針對不合格品或品質異常，採取改善以防止再發之措施。"),
        ("預防措施", "指針對可能發生之品質問題，預先採取之防範措施。"),
        ("品質稽核", "指對品質管理系統及施工品質進行有系統的查核作業。"),
    ]
    for term, definition in terms:
        p = doc.add_paragraph()
        run1 = p.add_run(f"「{term}」：")
        set_run_font(run1, DOCX_SZ_BODY, bold=True)
        run2 = p.add_run(definition)
        set_run_font(run2, DOCX_SZ_BODY)

    # ════════════════════════════════
    # 貳、管理責任及權責分工
    # ════════════════════════════════
    heading(doc, "貳、管理責任及權責分工", level=1)

    heading(doc, "一、品管組織架構", level=2)
    body(doc, "本公司依本工程規模及複雜度，設置品管組織如下圖所示：")

    # 品管組織架構圖（v8.5：無主辦機關/監造單位，三線型圖例）
    _org_para = build_quality_org_chart(
        label_tech    = f'主任技師\n{ENG_TECH}',
        label_foreman = f'工地主任\n{ENG_FOREMAN}',
        label_qc      = f'品管人員\n{ENG_QA}',
        label_safety  = f'職業安全衛生\n管理人員\n{ENG_SAFETY}',
        font_cjk      = DOCX_FONT_CJK,
        font_latin    = DOCX_FONT_LATIN,
        font_pt       = DOCX_SZ_TABLE,
        doc_id        = 999,
    )
    _ref_p = doc.add_paragraph()
    _ref_p._p.getparent().replace(_ref_p._p, _org_para)
    add_figure_caption(doc, 2, 1, "品管組織架構圖")

    heading(doc, "二、工作職掌", level=2)
    body(doc, "　　依據工程契約及「公共工程施工品質管理作業要點」之規定，指派經訓練合格且具工程實務經驗之人員從事工程品質管理。又依「營造業法」及「職業安全衛生管理辦法」之規定，指派本工程除專任工程人員外，配置工地主任一人、品管人員一人、職業安全衛生業務主管（營造業）一人；工作職掌如下表2-1：")
    add_table_caption(doc, 2, 1, "工作職掌表")
    make_table(doc,
               headers=["職稱", "工作職掌"],
               rows=[
                   ("專任工程人員\n（主任技師）",
                    "1.查核施工計畫書，並於認可後簽名或蓋章。\n"
                    "2.於開工、竣工報告文件及工程查報表簽名或蓋章。\n"
                    "3.督察按圖施工、解決施工技術問題。\n"
                    "4.依工地負責人之通報，處理工地緊急異常狀況。\n"
                    "5.查驗工程時到場說明，並於工程查驗文件簽名或蓋章。\n"
                    "6.營繕工程必須勘驗部分赴現場履勘，並於申報勘驗文件簽名或蓋章。\n"
                    "7.主管機關勘驗工程時，在場說明，並於相關文件簽名或蓋章。\n"
                    "8.其他依法令規定應辦理之事項。"),
                   ("工地負責人\n（工地主任）",
                    "1.依施工計畫書執行按圖施工。\n"
                    "2.按日填報施工日誌。\n"
                    "3.工地之人員、機具及材料等管理。\n"
                    "4.工地職業安全衛生事項之督導、公共環境與安全之維護及其他工地行政事務。\n"
                    "5.工地遇緊急異常狀況之通報。\n"
                    "6.其他依法令規定應辦理之事項。"),
                   ("品管人員",
                    "1.品管負責人負責綜理乙方品管業務及與甲方之聯繫工作。\n"
                    "2.依據工程契約、設計圖說、規範、相關技術法規及參考品質計畫製作綱要等，訂定品質計畫，並據以推動實施。\n"
                    "3.執行內部品質稽核，如稽核自主檢查表之檢查項目、檢查結果是否詳實記錄等。\n"
                    "4.品管統計分析、矯正與預防措施之提出及追蹤改善。\n"
                    "5.依合約規定負責品質文件、紀錄之管理。\n"
                    "6.現場施工作業之監督，及其他提升工程品質事宜。\n"
                    "7.督導施工組確實執行品檢工作，所有品管文件及報表之填寫。\n"
                    "8.辦理自驗及會同甲方試驗人員等辦理品管檢驗及試驗作業。\n"
                    "9.執行材料檢驗與工地試驗、檢驗工作之協調。\n"
                    "10.各種品管文件（如證明、報告、紀錄、施工計劃等圖說）之建立與審查。\n"
                    "11.各單項工程之檢驗及會同業主辦理完工審查。\n"
                    "12.派員查驗分包商、供應商與製造商之品管流程。\n"
                    "13.經授權後得代理品管負責人執行指定工作。\n"
                    "14.輔助與督導各協力廠商、材料供應商與製造商品管工作。\n"
                    "15.整理各種文件提送作業。\n"
                    "16.配合品質控制作業、執行材料檢驗、工地檢驗及自主檢查工作。\n"
                    "17.針對工程現況與施工品質作持續性查驗與評估，填寫各項作業之查驗與評估文件，呈報品管負責人。\n"
                    "18.負責一切品管文件檔案整理及建檔。\n"
                    "19.其他臨時交辦事項。"),
                   ("職業安全衛生\n管理人員",
                    "1.執行督導及記錄整理。\n"
                    "2.職業安全衛生協議組織。\n"
                    "3.自主檢查及記錄。\n"
                    "4.其他工地安全衛生管理業務之執行。"),
                   ("現場工程師",
                    "1.負責施工計畫撰寫及工程施工。\n"
                    "2.簽報工程協力廠商工程付款。\n"
                    "3.相關日報表之填寫。\n"
                    "4.負責工程施工中協力廠商施工品質之控管與自主品管之查驗。\n"
                    "5.會同甲方辦理現場施工品質查驗。\n"
                    "6.每日安全衛生自動檢查執行與紀錄填寫。\n"
                    "7.負責工程材料之管理。\n"
                    "8.進場材料之存放、標示管理。"),
                   ("行政人員",
                    "1.綜合施工處理人事、出納、總務等業務。\n"
                    "2.負責處理非工程技術性之外涉事宜。\n"
                    "3.宣達公司常態性之通知與公告。\n"
                    "4.安排增進施工處同仁情誼相關事宜。"),
               ],
               col_widths=[3.5, 13.0])

    heading(doc, "三、人員輪值及請休假事項", level=2)
    body(doc, "　　工地人員依照勞基法施行細則第23條之1規定：「本法第37條所定休假遇本法第36條所定例假及休息日者，應予補假。但不包括本法第37條指定應放假之日。前項補假期日，由勞雇雙方協商排定之。」，遇有趕工必要時，工地人員會輪派留守。若工地人員請休假無法駐守工地時，會事先覓妥本公司合格人員代理。")

    heading(doc, "四、管理審查", level=2)
    body(doc, "　　依每季審查品質管理系統以確保其持續適切、正確及有效性。審查內容包括前次審查紀錄執行情形、內部品質稽核、矯正措施、預防措施、外部驗證及工程評鑑結果。")

    # ════════════════════════════════
    # 參、施工要領
    # ════════════════════════════════
    heading(doc, "參、施工要領", level=1)

    heading(doc, "一、分項計畫管制表", level=2)
    add_table_caption(doc, 3, 1, "各分項計畫管制表")
    make_table(doc,
               headers=["項次", "提送項目", "預定送審日期", "備註"],
               rows=[
                   ("1", "整體品質計畫", "開工前10日", ""),
                   ("2", "整體施工計畫", "開工前10日", ""),
                   ("3", "步道鋪面分項施工計畫（含排磚計畫）", "施工前7日", ""),
                   ("4", "籃球場分項施工計畫", "施工前7日", ""),
                   ("5", "遊戲場彈性地坪分項施工計畫", "施工前7日", ""),
                   ("6", "遊具設備安裝分項施工計畫", "施工前7日", ""),
                   ("7", "塑木花架分項施工計畫（含施工圖）", "施工前7日", ""),
                   ("8", "公園廁所分項施工計畫（含施工圖）", "施工前7日", ""),
                   ("9", "公園排水分項施工計畫", "施工前7日", ""),
                   ("10", "植栽分項施工計畫", "施工前7日", ""),
                   ("11", "景觀照明分項施工計畫", "施工前7日", ""),
                   ("12", "停車場及圍籬分項施工計畫（含施工圖）", "施工前7日", ""),
               ],
               col_widths=[1.5, 8.5, 3.0, 3.5])

    heading(doc, "二、各工項施工要領", level=2)

    # 各工項施工要領表
    work_items_methods = [
        ("整地拆除工程", [
            ("清除及掘除", "植被、廢棄物等", "推土機、怪手", "依環保規定處理廢棄物，避免二次污染"),
            ("地坪打除", "無", "怪手、破碎機", "注意地下管線位置，避免損壞既有設施"),
            ("廢棄物清除", "無", "垃圾車", "廢棄物依分類清運，不得任意棄置"),
            ("路幅整修", "無", "怪手、壓路機", "整修後土層應達規定壓實度"),
        ]),
        ("混凝土鋪面工程", [
            ("基地開挖", "無", "怪手", "開挖至規定深度，基底整平夯實"),
            ("級配底層鋪設", "碎石級配", "壓路機", "壓實度≧95%，厚度符合設計圖說"),
            ("混凝土澆置", "預拌混凝土（fc\'=210kgf/cm²以上）", "預拌車、振動棒", "坍度15±4cm，溫度13~32°C，氯離子≦0.15kg/m³"),
            ("養護", "養護膜或養護劑", "灑水設備", "養護期間不得受外力破壞，至少養護7天"),
        ]),
        ("高壓混凝土磚鋪面工程", [
            ("基底整平", "無", "整平工具", "基底壓實度≧95%"),
            ("砂漿打底", "水泥砂漿1:3", "砂漿攪拌機", "砂漿厚度均勻，鋪設後立即鋪磚"),
            ("鋪設高壓磚", "高壓混凝土磚（抗壓強度≧350kgf/cm²）", "橡皮槌、水準儀", "坡度符合圖說，縫隙均勻，垂直度確認"),
            ("填縫", "細砂或填縫劑", "掃把、振動板", "填縫完整，表面清潔"),
        ]),
        ("路緣石工程", [
            ("開挖溝槽", "無", "怪手", "溝槽尺寸符合設計圖說"),
            ("基礎混凝土澆置", "預拌混凝土", "模板、振動棒", "基礎厚度及強度符合規定"),
            ("路緣石安裝", "預鑄路緣石", "橡皮槌、水平尺", "線形順暢，高程符合設計"),
            ("填縫及養護", "水泥砂漿", "刮刀", "填縫飽滿，養護至強度達要求"),
        ]),
        ("塑木工程", [
            ("基礎施工", "鋼構或混凝土基礎", "怪手、模板", "基礎強度及尺寸符合結構計算"),
            ("鋼構安裝", "型鋼、鍍鋅螺栓", "吊車、焊接設備", "鋼構組裝精準，焊接品質符合規範"),
            ("塑木安裝", "塑木板材", "電鑽、螺絲", "塑木間距均勻，固定牢固，螺絲不外露"),
            ("防腐防蝕處理", "防蝕塗料", "噴塗設備", "金屬件均應熱浸鍍鋅或塗防蝕漆"),
        ]),
        ("兒童遊戲場設備工程", [
            ("基礎施工", "混凝土基礎", "模板、混凝土", "基礎尺寸及強度符合設備廠商規定"),
            ("設備安裝", "遊戲場設備", "吊車、工具組", "依廠商施工圖安裝，確認鎖固扭力"),
            ("地坪鋪設", "彈性地坪、橡膠地墊、人工草皮", "整平工具", "厚度及面積符合設計，落高防護符合CNS規定"),
            ("安全性確認", "無", "無", "確認無尖銳物、夾縫、吞噬點，符合CNS12642遊戲場安全規範"),
        ]),
        ("公園排水工程", [
            ("溝槽開挖", "無", "怪手", "開挖位置及深度符合圖說"),
            ("管材安裝", "HDPE透水網管、PVC管", "無", "管材坡度符合設計，接頭密合"),
            ("陰井施工", "預鑄陰井或現場灌漿", "模板、混凝土", "陰井高程符合設計，蓋板牢固"),
            ("回填夯實", "原土或砂礫", "夯實機", "分層回填，壓實度≧90%"),
        ]),
        ("植栽工程", [
            ("植穴開挖", "無", "怪手、人工", "植穴尺寸符合樹種要求"),
            ("底土改良", "有機肥、客土", "無", "土壤改良充分，排水良好"),
            ("植樹", "各類喬木、灌木", "吊車（大型喬木）", "種植深度適當，支撐架設置牢固"),
            ("植草", "假儉草", "無", "舖植密實，確認存活"),
            ("養護", "水、肥料", "澆水設備", "養護期間至少6個月，確保存活率≧90%"),
        ]),
        ("景觀照明工程", [
            ("管線埋設", "PVC管、XLPE電纜", "無", "管線埋設深度符合電工法規，轉角處留餘裕"),
            ("燈座基礎", "混凝土基礎", "模板、混凝土", "基礎強度及尺寸符合燈具廠商規定"),
            ("燈具安裝", "各類景觀燈具", "工具組", "燈具方向、角度符合設計，鎖固確實"),
            ("電氣配線", "電纜、接線盒", "電工工具", "線徑符合設計，接頭絕緣確實"),
            ("功能測試", "無", "照度計、電錶", "照度符合設計規定，絕緣電阻≧1MΩ"),
        ]),
    ]

    for idx, (work_name, steps) in enumerate(work_items_methods, 1):
        add_table_caption(doc, 3, idx + 1, f"{work_name}施工要領")
        make_table(doc,
                   headers=["施工步驟", "使用材料", "施工機具", "注意事項"],
                   rows=steps,
                   col_widths=[3.0, 4.0, 3.5, 6.0])
        doc.add_paragraph()

    # ════════════════════════════════
    # 肆、品質管理標準
    # ════════════════════════════════
    heading(doc, "肆、品質管理標準", level=1)
    heading(doc, "一、各工項品質管理標準表", level=2)
    body(doc, "各工項品質管理標準如下表所示，所有管理項目均應量化，並依規定頻率進行自主檢查。")
    body(doc, "符號說明：▽ 為檢驗停留點，施工成果須自主檢查合格後方可進行下一步驟。")

    qm_items = [
        ("整地拆除工程", [
            ("廢棄物清除", "廢棄物清除完整度", "廢棄物完全清除，不得有殘留", "施工中", "目視", "全面", "廢棄物清除紀錄表 ▽", "重新清除"),
            ("路幅整修", "路基壓實度", "壓實度≧95%（修正夯實試驗）", "施工中", "工地密度試驗", "每200M2一點", "壓實度試驗報告 ▽", "重新夯實後再試驗"),
        ]),
        ("混凝土鋪面工程", [
            ("混凝土澆置", "坍度", "15±4 cm", "澆置前", "坍度試驗", "每50M3一次", "混凝土品管紀錄 ▽", "退回不合格混凝土"),
            ("混凝土澆置", "溫度", "13°C～32°C", "澆置中", "溫度計量測", "每車", "混凝土品管紀錄", "高溫加冰低溫加熱"),
            ("混凝土澆置", "氯離子含量", "≦0.15 kg/m³", "澆置前", "氯離子試驗", "每50M3一次", "氯離子試驗報告 ▽", "退回不合格混凝土"),
            ("混凝土澆置", "抗壓強度", "≧210 kgf/cm²（28天）", "養護後", "圓柱試體抗壓", "每50M3一組（3支）", "試體強度報告 ▽", "鑽心取樣複驗"),
            ("混凝土養護", "養護時間", "至少7天", "澆置後", "目視計時", "全數", "養護紀錄表", "延長養護"),
            ("鋪面平整度", "3m直尺量測", "凹凸差≦3mm", "施工後", "3m直尺", "每50M2一點 ▽", "自主檢查紀錄", "磨平或補鋪"),
        ]),
        ("高壓混凝土磚鋪面工程", [
            ("磚材進場", "外觀", "無裂縫、缺角，色澤均勻", "進場時", "目視 ▽", "全數抽樣", "進場檢驗紀錄", "不合格品退場"),
            ("磚材進場", "抗壓強度", "≧350 kgf/cm²", "進場時", "CNS13295試驗", "每批500塊抽1組", "強度試驗報告 ▽", "整批退場"),
            ("鋪設", "平整度", "凹凸差≦3mm（3m直尺）", "施工後", "3m直尺 ▽", "每50M2一點", "自主檢查表", "調整或重舖"),
            ("鋪設", "坡度", "符合設計圖說（橫坡≦2%）", "施工後", "水準儀", "每50M2一點 ▽", "坡度量測紀錄", "重新調整"),
        ]),
        ("路緣石工程", [
            ("路緣石進場", "抗彎強度", "符合CNS13295規定", "進場時", "抗彎試驗", "每批抽1組", "試驗報告 ▽", "整批退場"),
            ("安裝", "高程", "誤差±5mm以內", "施工後", "水準儀 ▽", "每10M一點", "自主檢查表", "調整高程"),
            ("安裝", "線形", "目視順暢，無明顯偏差", "施工後", "目視", "全線", "自主檢查表", "重新調整"),
        ]),
        ("塑木工程", [
            ("材料進場", "外觀", "無裂縫、變形，顏色均勻", "進場時", "目視 ▽", "全數", "進場紀錄", "不合格品退場"),
            ("鋼構安裝", "垂直度", "誤差±5mm/m以內", "施工中", "水平尺、鉛錘 ▽", "全數量測", "自主檢查表", "重新調整"),
            ("塑木安裝", "間距", "符合設計圖說±3mm", "施工後", "量尺 ▽", "抽查20%", "自主檢查表", "調整間距"),
        ]),
        ("植栽工程", [
            ("苗木進場", "規格", "符合設計圖說規格（樹高、樹幅、米高直徑）", "進場時 ▽", "目視量測", "全數", "苗木進場紀錄", "不合格品退場"),
            ("種植", "植穴尺寸", "符合規定（喬木：寬≧1.5倍根球直徑）", "施工中", "量尺 ▽", "抽查20%", "自主檢查表", "重挖植穴"),
            ("養護", "存活率", "≧90%（竣工後6個月）", "養護期 ▽", "目視清點", "全數", "存活率統計表", "補植至達標"),
            ("植草", "覆蓋率", "≧95%（竣工後3個月）", "養護期", "目視", "全面", "自主檢查表", "補植"),
        ]),
        ("景觀照明工程", [
            ("電纜進場", "規格", "符合設計圖說，線徑及絕緣等級符合要求", "進場時 ▽", "目視、型錄確認", "全數", "進場紀錄", "不合格品退場"),
            ("管線埋設", "埋設深度", "≧60cm（車行道）、≧30cm（步行區）", "施工中", "量尺 ▽", "每20M抽查一點", "自主檢查表", "重新埋設"),
            ("電氣安裝", "絕緣電阻", "≧1 MΩ（500V搖表量測）", "完工後 ▽", "搖表量測", "全數", "電氣試驗紀錄", "重新接線"),
            ("照明效果", "照度", "符合設計圖說", "試燈時 ▽", "照度計", "抽查30%", "照度量測紀錄", "調整燈具角度"),
        ]),
    ]

    for idx, (work_name, rows) in enumerate(qm_items, 1):
        add_table_caption(doc, 4, idx, f"{work_name}品質管理標準表")
        make_table(doc,
                   headers=["施工流程", "管理項目", "檢查標準", "檢查時機", "檢查方法", "檢查頻率", "管理紀錄", "不符合標準之處置方法"],
                   rows=rows,
                   col_widths=[2.0, 2.0, 3.5, 1.8, 2.0, 2.0, 2.0, 3.2])
        body(doc, "＊為檢驗停留點　▽為自主檢查停留點", indent=0)
        doc.add_paragraph()

    # ════════════════════════════════
    # 伍、材料及施工檢驗程序
    # ════════════════════════════════
    heading(doc, "伍、材料及施工檢驗程序", level=1)

    heading(doc, "一、材料設備選定送審流程", level=2)
    body(doc, "本工程材料及設備之選定及送審，依下列流程辦理：")
    add_flowchart_placeholder(doc, 'fig_5_1', preset_material_approval())
    add_figure_caption(doc, 5, 1, "材料/設備選定送審流程圖")

    heading(doc, "二、施工檢驗流程", level=2)
    body(doc, "本工程施工檢驗依下列流程辦理：")
    add_flowchart_placeholder(doc, 'fig_5_2', preset_work_inspection())
    add_figure_caption(doc, 5, 2, "施工檢驗流程圖")

    heading(doc, "三、各工項施工作業流程及檢驗圖", level=2)

    # 各工項三欄式流程圖
    three_col_items = [
        ("整地拆除工程", "dc_work", [
            {'flow': '施工前現況調查', 'check': '地下管線確認，現況照相', 'doc': '現況調查表', 'type': 'process'},
            {'flow': '施工範圍放樣', 'check': '範圍確認符合圖說', 'doc': '放樣紀錄', 'mark': '▽', 'type': 'process'},
            {'flow': '地坪打除拆除', 'check': '拆除完整，廢棄物分類', 'doc': '拆除紀錄', 'type': 'process'},
            {'flow': '廢棄物清運', 'check': '依環保規定清運，清運單保存', 'doc': '廢棄物清運單', 'mark': '▽', 'type': 'process'},
            {'flow': '完成', 'check': '', 'doc': '', 'type': 'terminator'},
        ]),
        ("混凝土鋪面工程", "conc_pave", [
            {'flow': '基底開挖整平', 'check': '基底高程，壓實度≧95%', 'doc': '壓實度試驗', 'mark': '▽', 'type': 'process'},
            {'flow': '級配底層鋪設', 'check': '厚度符合設計，壓實度≧95%', 'doc': '底層紀錄', 'mark': '▽', 'type': 'process'},
            {'flow': '模板組立', 'check': '尺寸、線形、水密性', 'doc': '模板自主檢查表', 'type': 'process'},
            {'flow': '混凝土試體取樣', 'check': '坍度、溫度、氯離子試驗', 'doc': '混凝土品管紀錄', 'mark': '▽', 'type': 'process'},
            {'flow': '混凝土澆置', 'check': '振動確實，分層≦30cm', 'doc': '澆置紀錄', 'type': 'process'},
            {'flow': '混凝土養護', 'check': '養護7天以上', 'doc': '養護紀錄', 'mark': '▽', 'type': 'process'},
            {'flow': '完成驗收', 'check': '平整度≦3mm、抗壓強度報告', 'doc': '驗收紀錄', 'type': 'terminator'},
        ]),
        ("植栽工程", "plant_work", [
            {'flow': '苗木進場檢驗', 'check': '規格符合設計，外觀健康', 'doc': '苗木進場紀錄', 'mark': '▽', 'type': 'process'},
            {'flow': '植穴開挖', 'check': '尺寸符合規定，底土改良', 'doc': '植穴自主檢查表', 'mark': '▽', 'type': 'process'},
            {'flow': '植樹', 'check': '種植深度、支撐架設置', 'doc': '植栽紀錄', 'type': 'process'},
            {'flow': '初期養護', 'check': '澆水、施肥、病蟲害防治', 'doc': '養護日誌', 'type': 'process'},
            {'flow': '養護期確認', 'check': '存活率≧90%（6個月後）', 'doc': '存活率統計表', 'mark': '▽', 'type': 'process'},
            {'flow': '完成驗收', 'check': '符合設計規格及存活率', 'doc': '驗收紀錄', 'type': 'terminator'},
        ]),
        ("景觀照明工程", "lighting", [
            {'flow': '材料進場檢驗', 'check': '規格、型號符合核准型錄', 'doc': '進場紀錄', 'mark': '▽', 'type': 'process'},
            {'flow': '管線埋設', 'check': '深度、坡向符合規定', 'doc': '管線自主檢查表', 'mark': '▽', 'type': 'process'},
            {'flow': '燈座基礎施作', 'check': '尺寸、強度符合設計', 'doc': '基礎紀錄', 'type': 'process'},
            {'flow': '燈具安裝', 'check': '方向、角度符合設計', 'doc': '安裝紀錄', 'mark': '▽', 'type': 'process'},
            {'flow': '電氣接線', 'check': '線徑正確，接頭絕緣', 'doc': '電氣紀錄', 'type': 'process'},
            {'flow': '功能測試', 'check': '照度、絕緣電阻符合規定', 'doc': '功能測試紀錄', 'mark': '▽', 'type': 'process'},
            {'flow': '完成驗收', 'check': '全數點亮確認正常', 'doc': '驗收紀錄', 'type': 'terminator'},
        ]),
    ]

    for idx, (work_name, prefix, items) in enumerate(three_col_items, 1):
        add_three_col_flow(doc, f"圖5-{idx + 2} {work_name}施工作業流程及檢驗圖",
                           items, marker_prefix=prefix,
                           font_cjk=DOCX_FONT_CJK,
                           font_latin=DOCX_FONT_LATIN,
                           font_pt=DOCX_SZ_TABLE)
        doc.add_paragraph()

    heading(doc, "四、材料設備送審管制總表", level=2)
    add_table_caption(doc, 5, 7, "材料設備送審管制總表")
    make_table(doc,
               headers=["項次", "材料/設備名稱", "送審依據", "預計送審日期", "審查期限", "備註"],
               rows=[
                   ("1", "預拌混凝土（配比設計資料）", "施工規範03050", "施工前15日", "10日", ""),
                   ("2", "高壓混凝土磚", "CNS13295", "施工前15日", "10日", ""),
                   ("3", "預鑄路緣石", "CNS13295", "施工前15日", "10日", ""),
                   ("4", "塑木材料", "設計圖說規格", "施工前15日", "10日", ""),
                   ("5", "各類遊戲場設備", "CNS12642", "施工前15日", "10日", ""),
                   ("6", "PE地下儲水槽", "設計圖說規格", "施工前15日", "10日", ""),
                   ("7", "HDPE透水網管", "設計圖說規格", "施工前15日", "10日", ""),
                   ("8", "各類苗木（植物材料）", "設計圖說規格", "施工前15日", "10日", ""),
                   ("9", "景觀燈具（各類型）", "設計圖說規格", "施工前15日", "10日", ""),
                   ("10", "XLPE電纜線", "CNS規定", "施工前15日", "10日", ""),
                   ("11", "電源開關箱", "電工法規", "施工前15日", "10日", ""),
                   ("12", "公園廁所各單元", "設計圖說規格", "施工前15日", "10日", ""),
               ],
               col_widths=[1.5, 4.0, 3.0, 2.5, 2.0, 3.5])

    # ════════════════════════════════
    # 陸、設備功能運轉測試
    # ════════════════════════════════
    heading(doc, "陸、設備功能運轉測試", level=1)

    heading(doc, "一、設備功能運轉檢測程序", level=2)
    body(doc, "本工程景觀照明系統設備功能運轉測試，依下列流程辦理：")
    add_flowchart_placeholder(doc, 'fig_6_1', preset_equipment_test())
    add_figure_caption(doc, 6, 1, "設備功能運轉檢測程序流程圖")

    heading(doc, "二、設備功能運轉檢測標準", level=2)
    add_table_caption(doc, 6, 1, "景觀照明系統設備功能運轉檢測標準一覽表")
    make_table(doc,
               headers=["項次", "測試項目", "測試方法", "合格標準", "測試頻率", "紀錄表單"],
               rows=[
                   ("1", "燈具點亮測試", "逐一開關測試", "全數正常點亮，無閃爍", "全數", "功能測試紀錄"),
                   ("2", "照度測試", "照度計量測", "符合設計照度值±10%", "抽查30%", "照度量測紀錄"),
                   ("3", "絕緣電阻測試", "500V搖表量測", "≧1 MΩ", "全數迴路", "電氣試驗紀錄"),
                   ("4", "接地電阻測試", "接地電阻計量測", "≦10 Ω", "全數接地點", "接地試驗紀錄"),
                   ("5", "開關箱功能測試", "人工操作測試", "開關動作正常，保護裝置正常", "全數", "開關箱測試紀錄"),
                   ("6", "漏電保護器測試", "測試鈕測試", "動作正常，跳脫時間符合規定", "全數", "漏電保護測試紀錄"),
                   ("7", "連續運轉測試", "連續通電24小時", "無異常發熱、異味或故障", "抽查", "連續運轉紀錄"),
               ],
               col_widths=[1.5, 3.0, 3.0, 3.5, 2.0, 3.5])

    body(doc, "\n設備功能運轉測試完成後，承攬廠商應提送測試報告，經監造單位確認合格後，方視為完成驗收。")

    # ════════════════════════════════
    # 柒、自主檢查表
    # ════════════════════════════════
    heading(doc, "柒、自主檢查表", level=1)
    body(doc, "各工項施工自主檢查表如下所示，每次施工前後均應確實填寫，並由工地主任及現場施工人員簽認。")

    selfcheck_items = [
        ("整地拆除工程", "EW-01", [
            ("廢棄物清除範圍", "施工範圍內廢棄物全部清除", "○╳/"),
            ("廢棄物分類清運", "依環保法規分類清運，保存清運單", "○╳/"),
            ("地坪打除完整性", "打除面積符合設計圖說，無遺漏", "○╳/"),
            ("路基壓實度", "壓實度≧95%（工地密度試驗）", "量測值：____%"),
            ("現況照相", "施工前後照相，存檔備查", "○╳/"),
        ]),
        ("混凝土鋪面工程", "EW-02", [
            ("基底整平及壓實", "基底高程符合設計，壓實度≧95%", "量測值：____%"),
            ("模板組立", "模板尺寸、線形、水密性符合要求", "○╳/"),
            ("混凝土坍度", "15±4 cm", "量測值：____cm"),
            ("混凝土溫度", "13°C～32°C", "量測值：____°C"),
            ("氯離子試驗", "≦0.15 kg/m³", "試驗報告編號：____"),
            ("試體取樣", "每50M3取樣一組（3支）", "試體編號：____"),
            ("養護時間", "至少7天，記錄養護方式", "養護方式：____"),
            ("鋪面平整度", "凹凸差≦3mm（3m直尺）", "量測值：____mm"),
        ]),
        ("植栽工程", "EW-03", [
            ("苗木規格確認", "樹高、樹幅、米高直徑符合設計圖說", "○╳/"),
            ("植穴尺寸", "寬度≧根球直徑1.5倍，深度適當", "量測值：____cm"),
            ("底土改良", "有機肥施加量及攪拌均勻", "○╳/"),
            ("種植深度", "根頸露出地面適當，不過深過淺", "○╳/"),
            ("支撐架設置", "竹竿或鋼管支撐確實，方向適當", "○╳/"),
            ("澆水量", "初期澆水充足，土壤濕潤至根球底部", "○╳/"),
            ("存活確認（6個月）", "存活率≧90%，死株補植", "存活株數：____"),
        ]),
        ("景觀照明工程", "EW-04", [
            ("燈具規格確認", "型號、瓦數符合設計圖說及核准型錄", "○╳/"),
            ("管線埋設深度", "車行道≧60cm，步行區≧30cm", "量測值：____cm"),
            ("燈座基礎尺寸", "符合設計圖說及廠商規定", "○╳/"),
            ("燈具安裝方向", "投光方向符合設計", "○╳/"),
            ("電氣接線", "線徑正確，接頭絕緣良好", "○╳/"),
            ("絕緣電阻", "≧1 MΩ（500V搖表）", "量測值：____MΩ"),
            ("燈具點亮測試", "全數正常點亮，無閃爍", "○╳/"),
            ("照度量測", "符合設計照度值", "量測值：____Lux"),
        ]),
    ]

    for idx, (work_name, ew_no, checks) in enumerate(selfcheck_items, 1):
        add_table_caption(doc, 7, idx, f"{work_name}施工自主檢查表（{ew_no}）")
        make_table(doc,
                   headers=["檢查項目", "設計圖說、規範之檢查標準（定量/定性）", "實際檢查情形（量化數據）", "檢查結果"],
                   rows=[(c[0], c[1], c[2], "○╳/") for c in checks],
                   col_widths=[3.5, 6.0, 4.0, 2.0])
        selfcheck_footer(doc)
        doc.add_paragraph()

    # ════════════════════════════════
    # 捌、不合格品之管制
    # ════════════════════════════════
    heading(doc, "捌、不合格品之管制", level=1)

    heading(doc, "一、不合格品管制流程", level=2)
    body(doc, "材料不合格品管制流程：")
    add_flowchart_placeholder(doc, 'fig_8_1', preset_nc_material())
    add_figure_caption(doc, 8, 1, "材料自主檢查不合格管制流程圖")

    body(doc, "施工不合格品管制流程：")
    add_flowchart_placeholder(doc, 'fig_8_2', preset_nc_work())
    add_figure_caption(doc, 8, 2, "施工自主檢查不合格管制流程圖")

    heading(doc, "二、不合格品管制表單", level=2)
    add_table_caption(doc, 8, 1, "不合格品改正紀錄表")
    make_table(doc,
               headers=["編號", "發現日期", "工項", "不合格項目", "不合格原因", "改正措施", "改正期限", "改正結果", "確認人員"],
               rows=[("    ", "    ", "    ", "    ", "    ", "    ", "    ", "    ", "    ")] * 5,
               col_widths=[1.2, 1.8, 2.0, 2.0, 2.5, 2.5, 1.8, 1.8, 1.9])

    doc.add_paragraph()
    add_table_caption(doc, 8, 2, "不合格品管制總表")
    make_table(doc,
               headers=["項次", "發現日期", "施工工項", "不合格內容", "處置方式", "完成日期", "複查結果", "備註"],
               rows=[("    ", "    ", "    ", "    ", "    ", "    ", "    ", "    ")] * 8,
               col_widths=[1.2, 1.8, 2.5, 3.0, 2.5, 1.8, 1.8, 2.9])

    # ════════════════════════════════
    # 玖、矯正與預防措施
    # ════════════════════════════════
    heading(doc, "玖、矯正與預防措施", level=1)

    body(doc, "本公司矯正與預防措施流程如下：")
    add_flowchart_placeholder(doc, 'fig_9_1', preset_corrective())
    add_figure_caption(doc, 9, 1, "矯正措施流程圖")

    body(doc, "一、矯正措施")
    for item in [
        "（一）當自主檢查發現不合格品或施工品質未達要求時，品管人員應立即填寫「不合格品改正紀錄表」。",
        "（二）品管人員應分析不合格原因，擬訂矯正措施，並追蹤改善結果。",
        "（三）矯正措施完成後，應重新進行自主檢查，確認品質達到要求。",
        "（四）所有矯正措施均應記錄於「不合格品管制總表」，作為品質改善依據。",
    ]:
        body(doc, item, indent=0.5)

    body(doc, "\n二、預防措施")
    for item in [
        "（一）定期辦理品質管理審查，分析品質問題趨勢，擬訂預防措施。",
        "（二）施工前辦理施工要領說明，確保施工人員了解品質要求。",
        "（三）材料進場前確認規格，避免不合格材料使用於工程中。",
        "（四）定期辦理品質教育訓練，提升施工人員品質意識。",
    ]:
        body(doc, item, indent=0.5)

    # ════════════════════════════════
    # 拾、內部品質稽核
    # ════════════════════════════════
    heading(doc, "拾、內部品質稽核", level=1)

    heading(doc, "一、品質稽核流程", level=2)
    body(doc, "本公司內部品質稽核流程如下：")
    add_flowchart_placeholder(doc, 'fig_10_1', preset_audit_flow())
    add_figure_caption(doc, 10, 1, "品質稽核流程圖")

    heading(doc, "二、品質稽核計畫表", level=2)
    body(doc, "本工程品質稽核計畫如下表所示，稽核頻率為每季一次：")
    add_table_caption(doc, 10, 1, "品質稽核計畫表")
    make_table(doc,
               headers=["稽核項目", "稽核內容", "稽核頻率", "稽核人員", "紀錄表單"],
               rows=[
                   ("品質計畫執行", "品質計畫各項措施執行情形", "每季", "主任技師/品管人員", "稽核紀錄表"),
                   ("材料管理", "材料送審、進場檢驗、試驗執行情形", "每季", "品管人員", "材料稽核表"),
                   ("施工品質", "自主檢查執行情形及品質紀錄", "每季", "品管人員", "施工稽核表"),
                   ("不合格品管制", "不合格品處置及矯正措施執行", "每季", "品管人員", "NC稽核表"),
                   ("文件管理", "品質紀錄文件完整性及保存", "每季", "品管人員", "文件稽核表"),
                   ("安全衛生", "安全衛生管理執行情形", "每季", "品管人員", "安衛稽核表"),
               ],
               col_widths=[2.5, 5.0, 2.0, 3.0, 4.0])

    # ════════════════════════════════
    # 拾壹、文件紀錄管理系統
    # ════════════════════════════════
    heading(doc, "拾壹、文件紀錄管理系統", level=1)

    heading(doc, "一、文件收發管理流程", level=2)
    body(doc, "收文傳送及歸檔流程：")
    add_flowchart_placeholder(doc, 'fig_11_1', preset_incoming_doc())
    add_figure_caption(doc, 11, 1, "收文傳送及歸檔流程圖")

    body(doc, "發文傳送及歸檔流程：")
    add_flowchart_placeholder(doc, 'fig_11_2', preset_outgoing_doc())
    add_figure_caption(doc, 11, 2, "發文傳送及歸檔流程圖")

    heading(doc, "二、文件分類代碼表", level=2)
    add_table_caption(doc, 11, 1, "文件分類代碼表")
    make_table(doc,
               headers=["代碼", "文件類別", "說明", "保存期限"],
               rows=[
                   ("QP", "品質計畫類", "整體品質計畫、分項品質計畫", "工程完工後5年"),
                   ("SP", "施工計畫類", "整體/分項施工計畫", "工程完工後5年"),
                   ("MR", "材料送審類", "材料型錄、試驗報告、送審表單", "工程完工後5年"),
                   ("SI", "自主檢查類", "各工項施工自主檢查表", "工程完工後5年"),
                   ("NC", "不合格品類", "不合格品改正紀錄、管制總表", "工程完工後5年"),
                   ("CA", "矯正預防類", "矯正/預防措施記錄", "工程完工後5年"),
                   ("AU", "稽核紀錄類", "品質稽核計畫、稽核紀錄", "工程完工後5年"),
                   ("LT", "往來文件類", "主辦機關、監造單位往來公文", "工程完工後5年"),
                   ("PH", "工程照相類", "施工前、中、後照片", "工程完工後5年"),
                   ("DR", "圖說類", "設計圖、竣工圖", "永久保存"),
               ],
               col_widths=[2.0, 3.5, 6.0, 4.0])

    body(doc, "\n文件編碼格式：【代碼】-【工項代號】-【流水號】-【日期】")
    body(doc, "例：QP-ALL-001-1150415（整體品質計畫，民國115年4月15日）")
    body(doc, "    SI-CONC-005-1151020（混凝土鋪面自主檢查，第5次，115年10月20日）")

    # ════════════════════════════════
    # 設定頁尾（三段式）
    # ════════════════════════════════
    sections = doc.sections
    if len(sections) >= 3:
        setup_section_footer(sections[0], 'none')
        setup_section_footer(sections[1], 'lowerRoman')
        setup_section_footer(sections[2], 'decimal')
    elif len(sections) == 2:
        setup_section_footer(sections[0], 'none')
        setup_section_footer(sections[1], 'decimal')

    # ════════════════════════════════
    # 儲存 & 後處理
    # ════════════════════════════════
    doc.save(OUT_PATH)
    print(f"✅ DOCX 已暫存：{OUT_PATH}")

    finalize_flowcharts(OUT_PATH,
                        theme_name='quality_plan',
                        font_cjk=DOCX_FONT_CJK,
                        font_latin=DOCX_FONT_LATIN,
                        font_pt=max(DOCX_SZ_TABLE - 1, 9))
    print("✅ DrawML 流程圖已替換完成")

    fix_docx(OUT_PATH)
    print(f"Done: {OUT_PATH}")


if __name__ == "__main__":
    build_doc()
